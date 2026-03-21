# =============================================================================
#  backend.py — SkillForge v14  |  Fixed edition
#  FIXES:
#  - _groq_call: handle 400 "Failed to generate JSON" with retry (simpler prompt)
#  - mega_call: split into two smaller calls (candidate + audit) to stay under limits
#  - mega_call: schema uses shorter field names in prompt, full names in parser
#  - _groq_call_vision: handle 400 errors
#  - run_analysis: surface json_too_complex error clearly
# =============================================================================

import os, json, io, re, time, hashlib, shelve, base64, threading
from typing import Tuple, Optional, List, Dict, Any
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor

import plotly.graph_objects as go
import networkx as nx
import pdfplumber
from docx import Document
from groq import Groq

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors as rl_colors
    REPORTLAB = True
except Exception:
    REPORTLAB = False

MODEL_FAST   = "llama-3.3-70b-versatile"
MODEL_VISION = "meta-llama/llama-4-scout-17b-16e-instruct"
CURRENT_YEAR = datetime.now().year
_CACHE_PATH  = "/tmp/skillforge_v14"

_GROQ_KEY   = os.getenv("GROQ_API_KEY", "")
GROQ_CLIENT = Groq(api_key=_GROQ_KEY) if _GROQ_KEY else None

_DDG_ERROR: str = ""
_DDG_LOCK       = threading.Lock()

_MAX_IMAGE_BYTES = 4 * 1024 * 1024  # 4 MB
_MIN_PDF_WORDS   = 30


def get_ddg_error() -> str:
    return _DDG_ERROR


def ddg_search(query: str, max_results: int = 5) -> List[dict]:
    global _DDG_ERROR
    try:
        from ddgs import DDGS
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=max_results * 2, region="wt-wt"))
            blocked = ["zhihu.com", "baidu.com", "weibo.com", "163.com", "csdn.net"]
            results = [r for r in results
                       if not any(b in r.get("href", "") for b in blocked)][:max_results]
            with _DDG_LOCK:
                _DDG_ERROR = ""
            return results
    except ImportError:
        with _DDG_LOCK:
            _DDG_ERROR = "web_search_unavailable"
        return []
    except Exception as e:
        with _DDG_LOCK:
            _DDG_ERROR = f"web_search_unavailable: {str(e)[:60]}"
        return []


def _is_english(text: str) -> bool:
    if not text:
        return True
    ascii_chars = sum(1 for c in text if ord(c) < 128)
    return ascii_chars / max(len(text), 1) > 0.75


def _is_meaningful_text(text: str) -> bool:
    if not text or not text.strip():
        return False
    words = [w for w in text.split() if len(w) > 1 and re.search(r'[a-zA-Z]', w)]
    if len(words) < _MIN_PDF_WORDS:
        return False
    garbage_ratio = text.count('\ufffd') / max(len(text), 1)
    if garbage_ratio > 0.05:
        return False
    return True


SEMANTIC: bool = False
_ST             = None
_CEMBS          = None


def _load_semantic_bg() -> None:
    global SEMANTIC, _ST, _CEMBS
    try:
        from sentence_transformers import SentenceTransformer
        _ST    = SentenceTransformer("all-MiniLM-L6-v2")
        _CEMBS = _ST.encode([c["skill"].lower() for c in CATALOG])
        SEMANTIC = True
    except Exception:
        pass


# =============================================================================
#  SKILL ALIASES
# =============================================================================
SKILL_ALIASES: Dict[str, str] = {
    "rest api design": "rest apis", "rest api": "rest apis", "restful api": "rest apis",
    "restful apis": "rest apis", "api design": "rest apis", "node.js": "rest apis",
    "nodejs": "rest apis", "express": "rest apis", "express.js": "rest apis",
    "expressjs": "rest apis", "graphql": "rest apis", "spring boot": "rest apis",
    "spring": "rest apis",
    "react.js": "react", "reactjs": "react", "next.js": "react",
    "nextjs": "react", "next": "react",
    "js": "javascript", "es6": "javascript", "es2015": "javascript",
    "typescript": "javascript", "ts": "javascript", "vue": "javascript",
    "vue.js": "javascript", "vuejs": "javascript", "angular": "javascript",
    "jest": "javascript",
    "python3": "python", "flask": "python", "django": "python",
    "streamlit": "python", "pytest": "python", "celery": "python",
    "html5": "html/css", "css3": "html/css", "html": "html/css",
    "css": "html/css", "sass": "html/css", "tailwind": "html/css",
    "tailwindcss": "html/css",
    "containerization": "docker", "containers": "docker",
    "dockerfile": "docker", "docker-compose": "docker", "docker compose": "docker",
    "k8s": "kubernetes", "helm": "kubernetes", "helm charts": "kubernetes",
    "mysql": "sql", "postgresql": "sql", "postgres": "sql", "sqlite": "sql",
    "nosql": "databases", "mongodb": "databases", "redis": "databases",
    "dynamodb": "databases", "cassandra": "databases", "elasticsearch": "databases",
    "pinecone": "databases", "vector database": "databases", "vector db": "databases",
    "kafka": "databases", "rabbitmq": "databases",
    "pytorch": "deep learning", "tensorflow": "deep learning",
    "huggingface": "deep learning", "hugging face": "deep learning",
    "transformers": "deep learning", "neural network": "deep learning",
    "cnn": "deep learning", "rnn": "deep learning", "lstm": "deep learning",
    "scikit-learn": "machine learning", "sklearn": "machine learning",
    "supervised learning": "machine learning", "unsupervised learning": "machine learning",
    "regression": "machine learning", "classification": "machine learning",
    "random forest": "machine learning", "xgboost": "machine learning",
    "nlp": "nlp", "langchain": "nlp", "llamaindex": "nlp", "llama index": "nlp",
    "rag": "nlp", "retrieval augmented": "nlp", "retrieval-augmented": "nlp",
    "llm": "nlp", "large language model": "nlp", "large language models": "nlp",
    "openai api": "nlp", "gpt": "nlp", "bert": "nlp", "sentence transformers": "nlp",
    "apache airflow": "mlops", "airflow": "mlops", "mlflow": "mlops",
    "model deployment": "mlops", "model monitoring": "mlops",
    "pandas": "data analysis", "numpy": "data analysis", "pyspark": "data analysis",
    "apache spark": "data analysis", "spark": "data analysis", "dbt": "data analysis",
    "matplotlib": "data visualization", "seaborn": "data visualization",
    "plotly": "data visualization", "tableau": "data visualization",
    "power bi": "data visualization", "powerbi": "data visualization",
    "looker": "data visualization", "grafana": "data visualization",
    "hypothesis testing": "statistics", "a/b testing": "statistics",
    "a/b test": "statistics", "statistical analysis": "statistics",
    "bash": "linux", "bash scripting": "linux", "shell scripting": "linux",
    "shell script": "linux", "unix": "linux", "nginx": "linux",
    "github actions": "ci/cd", "jenkins": "ci/cd", "gitlab ci": "ci/cd",
    "gitlab-ci": "ci/cd", "circleci": "ci/cd", "travis ci": "ci/cd", "argocd": "ci/cd",
    "aws sagemaker": "aws", "ec2": "aws", "s3": "aws", "lambda": "aws",
    "rds": "aws", "aws lambda": "aws",
    "gcp": "gcp", "google cloud": "gcp", "bigquery": "gcp",
    "terraform": "cloud computing", "infrastructure as code": "cloud computing",
    "iac": "cloud computing", "ansible": "cloud computing", "pulumi": "cloud computing",
    "owasp": "application security", "penetration testing": "application security",
    "pen testing": "application security", "devsecops": "application security",
    "scrum": "agile", "kanban": "agile", "sprint": "agile", "jira": "agile",
    # ── MERN ─────────────────────────────────────────────────────────────────
    "mern": "react", "mern stack": "react",
}

MERN_IMPLIES: List[str] = ["react", "javascript", "rest apis"]

_SKILL_REGEX_MAP: List[Tuple[str, str, int]] = [
    (r"\bSQL\b", "SQL", 5), (r"\bMySQL\b", "SQL", 6), (r"\bPostgreSQL\b", "SQL", 6),
    (r"\bSQLite\b", "SQL", 4), (r"\bMongoDB\b", "Databases", 5),
    (r"\bRedis\b", "Databases", 5), (r"\bDynamoDB\b", "Databases", 4),
    (r"\bCassandra\b", "Databases", 4), (r"\bElasticsearch\b", "Databases", 4),
    (r"\bPinecone\b", "Databases", 4), (r"\bKafka\b", "Databases", 4),
    (r"\bDockerfile\b", "Docker", 4), (r"\bdocker[- ]compose\b", "Docker", 5),
    (r"\bKubernetes\b", "Kubernetes", 3), (r"\bk8s\b", "Kubernetes", 3),
    (r"\bHelm\b", "Kubernetes", 4), (r"\bFastAPI\b", "FastAPI", 5),
    (r"\bflask\b", "Python", 6), (r"\bdjango\b", "Python", 6),
    (r"\bStreamlit\b", "Python", 5), (r"\bpytest\b", "Python", 5),
    (r"\bCelery\b", "Python", 5), (r"\bTypeScript\b", "JavaScript", 6),
    (r"\bNext\.?js\b", "React", 5), (r"\bVue\.?js\b", "JavaScript", 5),
    (r"\bAngular\b", "JavaScript", 5), (r"\bExpress\.?js\b", "REST APIs", 5),
    (r"\bGraphQL\b", "REST APIs", 5), (r"\bJest\b", "JavaScript", 4),
    (r"\bCI/CD\b", "CI/CD", 4), (r"\bGitHub Actions\b", "CI/CD", 5),
    (r"\bJenkins\b", "CI/CD", 5), (r"\bGitLab\s*CI\b", "CI/CD", 5),
    (r"\bCircleCI\b", "CI/CD", 5), (r"\bArgoCD\b", "CI/CD", 4),
    (r"\bAWS\b", "AWS", 4), (r"\bGCP\b", "GCP", 4),
    (r"\bAzure\b", "Cloud Computing", 4), (r"\bTerraform\b", "Cloud Computing", 4),
    (r"\bAnsible\b", "Cloud Computing", 4), (r"\bBigQuery\b", "GCP", 5),
    (r"\bLinux\b", "Linux", 4), (r"\bBash\b", "Linux", 5),
    (r"\bshell\s+script", "Linux", 4), (r"\bnginx\b", "Linux", 4),
    (r"\bPyTorch\b", "Deep Learning", 6), (r"\bTensorFlow\b", "Deep Learning", 6),
    (r"\bHugging\s*Face\b", "Deep Learning", 5),
    (r"\btransformer model\b", "Deep Learning", 5),
    (r"\bscikit[- ]learn\b", "Machine Learning", 6),
    (r"\bXGBoost\b", "Machine Learning", 5),
    (r"\bLangChain\b", "NLP", 5), (r"\bLlamaIndex\b", "NLP", 5),
    (r"\bRAG\b", "NLP", 4), (r"\bretrieval[- ]augmented", "NLP", 4),
    (r"\bLLM\b", "NLP", 4), (r"\bOpenAI\b", "NLP", 4), (r"\bBERT\b", "NLP", 5),
    (r"\bAirflow\b", "MLOps", 4), (r"\bMLflow\b", "MLOps", 5),
    (r"\bmodel\s+deploy", "MLOps", 4),
    (r"\bPySpark\b", "Data Analysis", 5), (r"\bApache\s+Spark\b", "Data Analysis", 5),
    (r"\bdbt\b", "Data Analysis", 4),
    (r"\bTableau\b", "Data Visualization", 5),
    (r"\bPower\s*BI\b", "Data Visualization", 5),
    (r"\bGrafana\b", "Data Visualization", 4),
    (r"\bRESTful?\b", "REST APIs", 5), (r"\bSpring\s*Boot\b", "REST APIs", 5),
    # ── MERN Stack ───────────────────────────────────────────────────────────
    (r"\bMERN\b",                     "React",           6),
    # ── Data Science basics ──────────────────────────────────────────────────
    (r"\bNumPy\b",                    "Data Analysis",   5),
    (r"\bPandas\b",                   "Data Analysis",   6),
    (r"\bMatplotlib\b",               "Data Visualization", 5),
    # ── Languages ────────────────────────────────────────────────────────────
    (r"\bJava\b",                     "Java",            5),
    (r"\bRust\b",                     "Rust",            5),
    (r"\bGo\b",                       "Go",              5),
    # ── Vector/AI databases ──────────────────────────────────────────────────
    (r"\bQdrant\b",                   "Databases",       5),
    (r"\bChromaDB\b",                 "Databases",       5),
    (r"\bSupabase\b",                 "Databases",       5),
    # ── Misc tools ───────────────────────────────────────────────────────────
    (r"\bGradio\b",                   "Python",          5),
    (r"\bLLaMA\b",                    "NLP",             5),
    (r"\bSymPy\b",                    "Python",          4),
]


def _apply_regex_skill_fallback(candidate: dict, resume_text: str) -> dict:
    # Guard: candidate must be a dict with a skills list of dicts
    if not isinstance(candidate, dict):
        candidate = {}
    skills = candidate.get("skills", [])
    if not isinstance(skills, list):
        skills = []
    # Filter to only dict entries that have a "skill" key
    candidate["skills"] = [s for s in skills if isinstance(s, dict) and "skill" in s]

    existing = {s["skill"].lower() for s in candidate["skills"]}
    alias_existing: set = set()
    for sk in list(existing):
        alias_existing.add(SKILL_ALIASES.get(sk, sk))
    existing |= alias_existing

    for pattern, skill_name, min_prof in _SKILL_REGEX_MAP:
        canonical = skill_name.lower()
        if canonical in existing:
            continue
        if SKILL_ALIASES.get(canonical, canonical) in existing:
            continue
        if re.search(pattern, resume_text, re.IGNORECASE):
            candidate.setdefault("skills", []).append({
                "skill":          skill_name,
                "proficiency":    min_prof,
                "year_last_used": CURRENT_YEAR,
                "context":        "Detected in resume text (regex scanner)",
            })
            existing.add(canonical)
            # MERN expands to React + JavaScript + REST APIs
            if skill_name == "React" and re.search(r"\bMERN\b", resume_text, re.IGNORECASE):
                for implied_skill, implied_prof in [("JavaScript", 6), ("REST APIs", 6)]:
                    if implied_skill.lower() not in existing:
                        candidate["skills"].append({
                            "skill": implied_skill, "proficiency": implied_prof,
                            "year_last_used": CURRENT_YEAR,
                            "context": "Derived from MERN Stack (regex scanner)",
                        })
                        existing.add(implied_skill.lower())
    return candidate


def _strip_mern_prefix(context: str) -> str:
    return re.sub(r"^Derived from \w[\w\s]+:\s*", "", context or "")


# =============================================================================
#  CATALOG — 47 courses
# =============================================================================
CATALOG: List[Dict] = [
    {"id":"PY01","title":"Python Fundamentals","skill":"Python","domain":"Tech","level":"Beginner","duration_hrs":6,"prereqs":[]},
    {"id":"PY02","title":"Python Intermediate: OOP & Modules","skill":"Python","domain":"Tech","level":"Intermediate","duration_hrs":8,"prereqs":["PY01"]},
    {"id":"PY03","title":"Python Advanced: Async & Decorators","skill":"Python","domain":"Tech","level":"Advanced","duration_hrs":10,"prereqs":["PY02"]},
    {"id":"DA01","title":"Data Analysis with Pandas","skill":"Data Analysis","domain":"Tech","level":"Beginner","duration_hrs":7,"prereqs":["PY01"]},
    {"id":"DA02","title":"Data Visualization (Matplotlib & Seaborn)","skill":"Data Visualization","domain":"Tech","level":"Intermediate","duration_hrs":5,"prereqs":["DA01"]},
    {"id":"DA03","title":"Statistical Analysis & Hypothesis Testing","skill":"Statistics","domain":"Tech","level":"Intermediate","duration_hrs":8,"prereqs":["DA01"]},
    {"id":"ML01","title":"Machine Learning Foundations","skill":"Machine Learning","domain":"Tech","level":"Beginner","duration_hrs":10,"prereqs":["DA01","DA03"]},
    {"id":"ML02","title":"Supervised Learning: Regression & Classification","skill":"Machine Learning","domain":"Tech","level":"Intermediate","duration_hrs":12,"prereqs":["ML01"]},
    {"id":"ML03","title":"Deep Learning with PyTorch","skill":"Deep Learning","domain":"Tech","level":"Advanced","duration_hrs":15,"prereqs":["ML02"]},
    {"id":"ML04","title":"NLP & Large Language Models","skill":"NLP","domain":"Tech","level":"Advanced","duration_hrs":14,"prereqs":["ML02"]},
    {"id":"ML05","title":"MLOps & Model Deployment","skill":"MLOps","domain":"Tech","level":"Advanced","duration_hrs":10,"prereqs":["ML02","DO02"]},
    {"id":"SQL01","title":"SQL Fundamentals","skill":"SQL","domain":"Tech","level":"Beginner","duration_hrs":5,"prereqs":[]},
    {"id":"SQL02","title":"Advanced SQL: Window Functions & Optimization","skill":"SQL","domain":"Tech","level":"Advanced","duration_hrs":7,"prereqs":["SQL01"]},
    {"id":"SQL03","title":"Database Design & NoSQL","skill":"Databases","domain":"Tech","level":"Intermediate","duration_hrs":6,"prereqs":["SQL01"]},
    {"id":"DO01","title":"Linux & Bash Scripting","skill":"Linux","domain":"Tech","level":"Beginner","duration_hrs":6,"prereqs":[]},
    {"id":"DO02","title":"Docker & Containerization","skill":"Docker","domain":"Tech","level":"Intermediate","duration_hrs":7,"prereqs":["DO01"]},
    {"id":"DO03","title":"Kubernetes Orchestration","skill":"Kubernetes","domain":"Tech","level":"Advanced","duration_hrs":10,"prereqs":["DO02"]},
    {"id":"DO04","title":"CI/CD Pipelines with GitHub Actions","skill":"CI/CD","domain":"Tech","level":"Intermediate","duration_hrs":6,"prereqs":["DO01"]},
    {"id":"CL01","title":"Cloud Computing Fundamentals","skill":"Cloud Computing","domain":"Tech","level":"Beginner","duration_hrs":6,"prereqs":[]},
    {"id":"CL02","title":"AWS Core Services Deep Dive","skill":"AWS","domain":"Tech","level":"Intermediate","duration_hrs":10,"prereqs":["CL01"]},
    {"id":"CL03","title":"GCP & BigQuery for Data Engineers","skill":"GCP","domain":"Tech","level":"Intermediate","duration_hrs":8,"prereqs":["CL01","DA01"]},
    {"id":"WE01","title":"HTML & CSS Foundations","skill":"HTML/CSS","domain":"Tech","level":"Beginner","duration_hrs":5,"prereqs":[]},
    {"id":"WE02","title":"JavaScript Essentials","skill":"JavaScript","domain":"Tech","level":"Beginner","duration_hrs":8,"prereqs":["WE01"]},
    {"id":"WE03","title":"React.js Fundamentals","skill":"React","domain":"Tech","level":"Intermediate","duration_hrs":10,"prereqs":["WE02"]},
    {"id":"WE04","title":"FastAPI Backend Development","skill":"FastAPI","domain":"Tech","level":"Intermediate","duration_hrs":8,"prereqs":["PY02"]},
    {"id":"WE05","title":"Full-Stack Integration & REST APIs","skill":"REST APIs","domain":"Tech","level":"Intermediate","duration_hrs":7,"prereqs":["WE03","WE04"]},
    {"id":"SE01","title":"Cybersecurity Fundamentals","skill":"Cybersecurity","domain":"Tech","level":"Beginner","duration_hrs":6,"prereqs":[]},
    {"id":"SE02","title":"Application Security & OWASP Top 10","skill":"Application Security","domain":"Tech","level":"Intermediate","duration_hrs":8,"prereqs":["SE01"]},
    {"id":"AG01","title":"Agile & Scrum Fundamentals","skill":"Agile","domain":"Tech","level":"Beginner","duration_hrs":4,"prereqs":[]},
    {"id":"AG02","title":"Advanced Scrum Master Certification Prep","skill":"Scrum","domain":"Tech","level":"Advanced","duration_hrs":6,"prereqs":["AG01"]},
    {"id":"PM01","title":"Project Management Essentials (PMI)","skill":"Project Management","domain":"Soft","level":"Intermediate","duration_hrs":8,"prereqs":["LD01"]},
    {"id":"HR01","title":"HR Fundamentals & Employment Law","skill":"Human Resources","domain":"Non-Tech","level":"Beginner","duration_hrs":5,"prereqs":[]},
    {"id":"HR02","title":"Talent Acquisition & Recruitment","skill":"Recruitment","domain":"Non-Tech","level":"Intermediate","duration_hrs":6,"prereqs":["HR01"]},
    {"id":"HR03","title":"Performance Management & Appraisals","skill":"Performance Management","domain":"Non-Tech","level":"Intermediate","duration_hrs":5,"prereqs":["HR01"]},
    {"id":"HR04","title":"Employee Relations & Conflict Resolution","skill":"Employee Relations","domain":"Non-Tech","level":"Advanced","duration_hrs":6,"prereqs":["HR01"]},
    {"id":"HR05","title":"Learning & Development Strategy","skill":"L&D Strategy","domain":"Non-Tech","level":"Advanced","duration_hrs":6,"prereqs":["HR03"]},
    {"id":"OP01","title":"Supply Chain & Logistics Fundamentals","skill":"Logistics","domain":"Non-Tech","level":"Beginner","duration_hrs":5,"prereqs":[]},
    {"id":"OP02","title":"Warehouse Management Systems","skill":"Warehouse Management","domain":"Non-Tech","level":"Intermediate","duration_hrs":6,"prereqs":["OP01"]},
    {"id":"OP03","title":"Inventory Control & Demand Planning","skill":"Inventory Management","domain":"Non-Tech","level":"Intermediate","duration_hrs":5,"prereqs":["OP01"]},
    {"id":"OP04","title":"Lean Manufacturing & Six Sigma Green Belt","skill":"Process Improvement","domain":"Non-Tech","level":"Advanced","duration_hrs":8,"prereqs":["OP01"]},
    {"id":"FI01","title":"Financial Accounting Basics","skill":"Accounting","domain":"Non-Tech","level":"Beginner","duration_hrs":6,"prereqs":[]},
    {"id":"FI02","title":"Financial Analysis & Modeling","skill":"Financial Analysis","domain":"Non-Tech","level":"Intermediate","duration_hrs":8,"prereqs":["FI01"]},
    {"id":"FI03","title":"Budgeting & Forecasting","skill":"Budgeting","domain":"Non-Tech","level":"Intermediate","duration_hrs":5,"prereqs":["FI01"]},
    {"id":"LD01","title":"Communication & Presentation Skills","skill":"Communication","domain":"Soft","level":"Beginner","duration_hrs":4,"prereqs":[]},
    {"id":"LD02","title":"Team Leadership & People Management","skill":"Leadership","domain":"Soft","level":"Intermediate","duration_hrs":6,"prereqs":["LD01"]},
    {"id":"LD03","title":"Strategic Thinking & Decision Making","skill":"Strategic Planning","domain":"Soft","level":"Advanced","duration_hrs":6,"prereqs":["LD02"]},
    {"id":"LD04","title":"Cross-Functional Collaboration","skill":"Collaboration","domain":"Soft","level":"Beginner","duration_hrs":3,"prereqs":["LD01"]},
]
CATALOG_BY_ID  = {c["id"]: c for c in CATALOG}
CATALOG_SKILLS = [c["skill"].lower() for c in CATALOG]

MARKET_DEMAND: Dict[str, int] = {
    "python":3,"machine learning":3,"deep learning":3,"aws":3,"docker":3,
    "sql":3,"react":3,"kubernetes":3,"mlops":3,"fastapi":3,"data analysis":3,
    "cloud computing":3,"ci/cd":3,"nlp":3,"rest apis":3,"javascript":3,
    "statistics":2,"data visualization":2,"gcp":2,"linux":2,"html/css":2,
    "agile":2,"databases":2,"cybersecurity":2,"communication":2,"leadership":2,
    "project management":2,"recruitment":2,"performance management":2,
    "financial analysis":2,"application security":2,"process improvement":2,
    "human resources":1,"accounting":1,"budgeting":1,"logistics":1,"scrum":1,
    "collaboration":1,"strategic planning":1,"warehouse management":1,
    "inventory management":1,"l&d strategy":1,
}


def demand_label(skill: str) -> str:
    d = MARKET_DEMAND.get(skill.lower(), 1)
    if d >= 3: return "🔥 Hot"
    if d >= 2: return "📈 Growing"
    return "✓ Stable"


OBSOLESCENCE_RISK: Dict[str, str] = {
    "jquery":         "Replaced by vanilla JS and React",
    "php":            "Declining; Python/Node dominant",
    "hadoop":         "Replaced by Spark + cloud-native",
    "excel vba":      "Power Query and Python replacing VBA",
    "manual testing": "AI-assisted automation replacing manual QA",
    "waterfall":      "Industry fully shifted to Agile/DevOps",
}

TRANSFER_MAP: Dict[str, Dict[str, int]] = {
    "python":            {"machine learning":40,"mlops":35,"fastapi":60,"data analysis":50,
                          "deep learning":30,"rest apis":45,"docker":25,"aws":20},
    "machine learning":  {"deep learning":50,"mlops":45,"nlp":40,"statistics":30},
    "javascript":        {"react":55,"rest apis":40,"html/css":35,"ci/cd":20},
    "react":             {"javascript":50,"html/css":60,"rest apis":30},
    "rest apis":         {"fastapi":45},
    "sql":               {"data analysis":35,"databases":60,"aws":15},
    "docker":            {"kubernetes":45,"ci/cd":35,"mlops":30,"aws":30,"linux":40},
    "linux":             {"docker":40,"ci/cd":30,"aws":20},
    "aws":               {"gcp":30,"cloud computing":70,"mlops":25,"docker":20,"kubernetes":25},
    "mern stack":        {"react":65,"javascript":60,"rest apis":55,"html/css":50},
    "html/css":          {"react":30,"javascript":20},
    "rest api design":   {"fastapi":50,"rest apis":55},
    "human resources":   {"recruitment":45,"performance management":40,"employee relations":35},
    "communication":     {"leadership":35,"project management":25},
    "leadership":        {"strategic planning":40},
    "financial analysis":{"budgeting":55,"accounting":40},
}

SENIORITY_MAP: Dict[str, int] = {"Junior":0,"Mid":1,"Senior":2,"Lead":3}

SAMPLES: Dict[str, Dict] = {
    "junior_swe": {
        "label":  "Junior SWE → Full Stack",
        "resume": """John Smith
Junior Software Developer | 1 year experience
Skills: Python (basic, 4/10), HTML/CSS, some JavaScript
Education: B.Tech Computer Science 2023
Projects: Built a todo app using Flask. Familiar with Git basics.
No professional cloud or DevOps experience.""",
        "jd": """Software Engineer Full Stack - Mid Level
Required: Python, React, FastAPI, Docker, SQL, REST APIs, AWS
Preferred: Kubernetes, CI/CD
Seniority: Mid | Domain: Tech""",
    },
    "senior_ds": {
        "label":  "Senior DS → Lead AI",
        "resume": """Priya Patel
Senior Data Scientist | 7 years experience
Skills: Python (expert, 9/10), Machine Learning (expert), Deep Learning (PyTorch, 8/10), SQL (advanced, 8/10), AWS SageMaker (7/10)
Last used NLP: 2022. Last used MLOps: 2021.
Led team of 5. Published 3 ML papers.""",
        "jd": """Lead Data Scientist - AI Products
Required: Python, Machine Learning, Deep Learning, NLP, MLOps, SQL, AWS
Preferred: GCP, Kubernetes, Leadership
Seniority: Lead | Domain: Tech""",
    },
    "hr_manager": {
        "label":  "HR Coordinator → Manager",
        "resume": """Amara Johnson
HR Coordinator | 3 years experience
Skills: Human Resources (intermediate, 6/10), Recruitment (good, 7/10), Microsoft Office
Some performance review experience. No formal L&D training.""",
        "jd": """HR Manager - People and Culture
Required: Human Resources, Recruitment, Performance Management, Employee Relations
Preferred: L&D Strategy, Communication, Leadership
Seniority: Senior | Domain: Non-Tech""",
    },
}


def _build_graph() -> nx.DiGraph:
    G = nx.DiGraph()
    for c in CATALOG:
        G.add_node(c["id"], **c)
        for p in c["prereqs"]:
            G.add_edge(p, c["id"])
    return G


SKILL_GRAPH = _build_graph()


def _parse_bytes(raw_bytes: bytes, filename: str) -> Tuple[str, Optional[str]]:
    name = filename.lower()

    if name.endswith(".pdf"):
        text = ""
        try:
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                parts = []
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        parts.append(extracted)
                text = "\n".join(parts)
        except Exception as e:
            text = f"[pdfplumber error: {e}]"

        if _is_meaningful_text(text):
            return text, None

        try:
            import fitz
            doc   = fitz.open(stream=raw_bytes, filetype="pdf")
            mat   = fitz.Matrix(2, 2)
            pixmaps = []
            for page_idx in range(min(3, doc.page_count)):
                pg  = doc.load_page(page_idx)
                pix = pg.get_pixmap(matrix=mat)
                pixmaps.append(pix)
            doc.close()

            if len(pixmaps) == 1:
                img_bytes = pixmaps[0].tobytes("png")
            else:
                try:
                    from PIL import Image as PILImage
                    pages_pil = [PILImage.open(io.BytesIO(p.tobytes("png"))) for p in pixmaps]
                    max_w     = max(i.width for i in pages_pil)
                    total_h   = sum(i.height for i in pages_pil)
                    canvas    = PILImage.new("RGB", (max_w, total_h), (255, 255, 255))
                    y_off     = 0
                    for img in pages_pil:
                        canvas.paste(img, (0, y_off))
                        y_off += img.height
                    buf = io.BytesIO()
                    canvas.save(buf, format="PNG")
                    img_bytes = buf.getvalue()
                except ImportError:
                    img_bytes = pixmaps[0].tobytes("png")

            if len(img_bytes) > _MAX_IMAGE_BYTES:
                try:
                    from PIL import Image as PILImage
                    pil_img = PILImage.open(io.BytesIO(img_bytes))
                    pil_img.thumbnail((1400, 3000), PILImage.LANCZOS)
                    buf = io.BytesIO()
                    pil_img.save(buf, format="PNG")
                    img_bytes = buf.getvalue()
                except Exception:
                    img_bytes = pixmaps[0].tobytes("png")

            if len(img_bytes) <= _MAX_IMAGE_BYTES:
                b64 = base64.b64encode(img_bytes).decode()
                return "", f"data:image/png;base64,{b64}"
        except ImportError:
            pass
        except Exception:
            pass

        return text if text and not text.startswith("[") else "[Scanned PDF — install PyMuPDF for vision OCR]", None

    if name.endswith(".docx"):
        try:
            doc  = Document(io.BytesIO(raw_bytes))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if _is_meaningful_text(text):
                return text, None
            text_full = "\n".join(
                r.text for p in doc.paragraphs for r in p.runs if r.text.strip()
            )
            return text_full if text_full.strip() else "[DOCX empty or unreadable]", None
        except Exception as e:
            return f"[DOCX error: {e}]", None

    if any(name.endswith(x) for x in [".jpg", ".jpeg", ".png", ".webp"]):
        media = (
            "image/jpeg" if name.endswith((".jpg", ".jpeg"))
            else "image/png"  if name.endswith(".png")
            else "image/webp"
        )
        if len(raw_bytes) > _MAX_IMAGE_BYTES:
            try:
                from PIL import Image as PILImage
                img = PILImage.open(io.BytesIO(raw_bytes))
                img.thumbnail((1600, 2200), PILImage.LANCZOS)
                buf = io.BytesIO()
                fmt = "JPEG" if media == "image/jpeg" else "PNG"
                img.save(buf, format=fmt, quality=85)
                raw_bytes = buf.getvalue()
                media     = "image/jpeg" if fmt == "JPEG" else "image/png"
            except Exception:
                pass

        b64 = base64.b64encode(raw_bytes).decode()
        return "", f"data:{media};base64,{b64}"

    return raw_bytes.decode("utf-8", errors="ignore"), None


_audit_log: List[dict] = []


def _repair_json(text: str) -> str:
    """Close a truncated JSON string by appending missing brackets/braces."""
    closers: list = []
    in_str = escape = False
    for ch in text:
        if escape:            escape = False; continue
        if ch == "\\" and in_str: escape = True; continue
        if ch == '"':         in_str = not in_str; continue
        if in_str:            continue
        if   ch == "{":       closers.append("}")
        elif ch == "[":       closers.append("]")
        elif ch == "}":
            if closers and closers[-1] == "}": closers.pop()
        elif ch == "]":
            if closers and closers[-1] == "]": closers.pop()
    if in_str: text += '"'
    text += "".join(reversed(closers))
    return text


def _extract_json(text: str) -> str:
    """
    Extract and, if necessary, repair a JSON object from model output.
    1. Strip markdown fences, try direct parse
    2. Find { … } by brace-matching
    3. If brace-match fails (truncation), repair and retry
    """
    text = re.sub(r"```(?:json)?\s*", "", text).replace("```", "").strip()
    try:
        json.loads(text); return text
    except (json.JSONDecodeError, ValueError):
        pass

    start = text.find("{")
    if start == -1:
        return "{}"

    depth, end, in_str, escape = 0, -1, False, False
    for i, ch in enumerate(text[start:], start):
        if escape:            escape = False; continue
        if ch == "\\" and in_str: escape = True; continue
        if ch == '"':         in_str = not in_str; continue
        if in_str:            continue
        if ch == "{":         depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0: end = i + 1; break

    if end > start:
        candidate = text[start:end]
        try:
            json.loads(candidate); return candidate
        except (json.JSONDecodeError, ValueError):
            pass

    # Truncated — repair
    fragment = text[start:]
    repaired = _repair_json(fragment)
    try:
        json.loads(repaired); return repaired
    except (json.JSONDecodeError, ValueError):
        return "{}"


def _groq_call_vision(prompt: str, system: str, image_b64: str, max_tokens: int = 3200) -> dict:
    if not GROQ_CLIENT:
        return {"error": "GROQ_API_KEY not set"}
    content  = [
        {"type": "image_url", "image_url": {"url": image_b64}},
        {"type": "text", "text": prompt},
    ]
    messages = [
        {"role": "system", "content": system},
        {"role": "user",   "content": content},
    ]
    t0 = time.time()
    try:
        r        = GROQ_CLIENT.chat.completions.create(
            model=MODEL_VISION, messages=messages, temperature=0.1, max_tokens=max_tokens,
        )
        raw_text = r.choices[0].message.content or "{}"
        usage    = r.usage
        in_tok   = usage.prompt_tokens     if usage else 0
        out_tok  = usage.completion_tokens if usage else 0
        cost     = round((in_tok * 0.00000011) + (out_tok * 0.00000034), 6)
        _audit_log.append({
            "ts": datetime.now().strftime("%H:%M:%S"),
            "model": MODEL_VISION.split("/")[-1][:22],
            "in": in_tok, "out": out_tok,
            "ms": round((time.time() - t0) * 1000), "cost": cost, "status": "ok",
        })
        return json.loads(_extract_json(raw_text))
    except json.JSONDecodeError as e:
        _audit_log.append({"ts": datetime.now().strftime("%H:%M:%S"),
                           "model": MODEL_VISION.split("/")[-1][:22],
                           "status": f"json_err:{str(e)[:40]}",
                           "in": 0, "out": 0, "ms": 0, "cost": 0})
        return {"error": f"vision_json_parse_failed: {e}"}
    except Exception as e:
        err = str(e)
        if "429" in err or "rate_limit" in err.lower():
            m = re.search(r"try again in (\d+)m([\d.]+)s", err)
            wait_s = (int(m.group(1)) * 60 + float(m.group(2))) if m else 60
            return {"error": "rate_limited", "wait_seconds": int(wait_s),
                    "message": f"Rate limited. Retry in {int(wait_s // 60)}m{int(wait_s % 60)}s."}
        if "400" in err:
            return {"error": "json_schema_too_complex",
                    "message": "Vision prompt too complex. Try uploading as PDF or pasting text."}
        _audit_log.append({"ts": datetime.now().strftime("%H:%M:%S"),
                           "model": MODEL_VISION.split("/")[-1][:22],
                           "status": f"err:{err[:40]}",
                           "in": 0, "out": 0, "ms": 0, "cost": 0})
        return {"error": err}


# =============================================================================
#  GROQ TEXT CALL — handles 400 "Failed to generate JSON" with clean retry
# =============================================================================
def _groq_call(prompt: str, system: str, model: str = MODEL_FAST,
               max_tokens: int = 2800, _retry: bool = True) -> dict:
    """
    Calls Groq with response_format=json_object.
    On 400 (schema too complex), strips the prompt and retries once without
    response_format forcing — then manually extracts JSON from the response.
    """
    if not GROQ_CLIENT:
        return {"error": "GROQ_API_KEY not set"}
    messages = [
        {"role": "system", "content": system},
        {"role": "user",   "content": prompt},
    ]
    t0 = time.time()
    try:
        r = GROQ_CLIENT.chat.completions.create(
            model=model, messages=messages, temperature=0.1,
            max_tokens=max_tokens,
            response_format={"type": "json_object"},
        )
        usage   = r.usage
        in_tok  = usage.prompt_tokens     if usage else 0
        out_tok = usage.completion_tokens if usage else 0
        cost    = round((in_tok * 0.00000011) + (out_tok * 0.00000034), 6)
        _audit_log.append({
            "ts": datetime.now().strftime("%H:%M:%S"),
            "model": model.split("/")[-1][:22],
            "in": in_tok, "out": out_tok,
            "ms": round((time.time() - t0) * 1000), "cost": cost, "status": "ok",
        })
        raw = r.choices[0].message.content or "{}"
        return json.loads(raw)

    except json.JSONDecodeError:
        return {"error": "json_parse_failed"}

    except Exception as e:
        err = str(e)

        # ── Rate limit ──────────────────────────────────────────────────────
        if "429" in err or "rate_limit" in err.lower():
            m = re.search(r"try again in (\d+)m([\d.]+)s", err)
            wait_s = (int(m.group(1)) * 60 + float(m.group(2))) if m else 60
            return {"error": "rate_limited", "wait_seconds": int(wait_s),
                    "message": f"Rate limited. Retry in {int(wait_s // 60)}m{int(wait_s % 60)}s."}

        # ── 400 "Failed to generate JSON" — retry without forced JSON mode ─
        if ("400" in err or "failed_generation" in err.lower()
                or "failed to generate json" in err.lower()):
            if _retry:
                return _groq_call_no_format(prompt, system, model, max_tokens)
            return {"error": "json_schema_too_complex",
                    "message": "Prompt too complex for JSON mode. Retried without format — still failed."}

        _audit_log.append({"ts": datetime.now().strftime("%H:%M:%S"),
                           "model": model.split("/")[-1][:22],
                           "status": f"err:{err[:40]}",
                           "in": 0, "out": 0, "ms": 0, "cost": 0})
        return {"error": err}


def _groq_call_no_format(prompt: str, system: str, model: str = MODEL_FAST,
                          max_tokens: int = 6000) -> dict:
    """
    Fallback: call without response_format=json_object, then manually extract JSON.
    Used when the model returns 400 in strict JSON mode.
    """
    if not GROQ_CLIENT:
        return {"error": "GROQ_API_KEY not set"}
    messages = [
        {"role": "system", "content": system + "\nIMPORTANT: Respond with ONLY a valid JSON object. No prose before or after."},
        {"role": "user",   "content": prompt},
    ]
    t0 = time.time()
    try:
        r = GROQ_CLIENT.chat.completions.create(
            model=model, messages=messages, temperature=0.1, max_tokens=max_tokens,
        )
        usage   = r.usage
        in_tok  = usage.prompt_tokens     if usage else 0
        out_tok = usage.completion_tokens if usage else 0
        cost    = round((in_tok * 0.00000011) + (out_tok * 0.00000034), 6)
        _audit_log.append({
            "ts": datetime.now().strftime("%H:%M:%S"),
            "model": (model.split("/")[-1][:22]) + "(nf)",
            "in": in_tok, "out": out_tok,
            "ms": round((time.time() - t0) * 1000), "cost": cost, "status": "ok-nf",
        })
        raw = r.choices[0].message.content or "{}"
        extracted = _extract_json(raw)
        return json.loads(extracted)
    except json.JSONDecodeError:
        return {"error": "json_parse_failed_fallback"}
    except Exception as e:
        err = str(e)
        if "429" in err or "rate_limit" in err.lower():
            m = re.search(r"try again in (\d+)m([\d.]+)s", err)
            wait_s = (int(m.group(1)) * 60 + float(m.group(2))) if m else 60
            return {"error": "rate_limited", "wait_seconds": int(wait_s),
                    "message": f"Rate limited. Retry in {int(wait_s // 60)}m{int(wait_s % 60)}s."}
        return {"error": err}


# =============================================================================
#  SYSTEM PROMPTS
# =============================================================================
_MEGA_SYS = (
    "You are a senior tech recruiter and L&D expert. "
    "Extract structured skill and job data from the resume and job description. "
    "RULES: Only extract facts explicitly present in the text. Do not infer or fabricate. "
    "Return ONLY valid JSON — no markdown, no prose."
)
_VISION_SYS = (
    "You are an expert resume parser reading an IMAGE of a resume. "
    "OCR all visible text and extract structured data. "
    "Only extract what is visibly present. Return ONLY valid JSON."
)
_AUDIT_SYS = (
    "You are an ATS expert and resume auditor. "
    "Analyze the resume against the job description and return structured audit data. "
    "Return ONLY valid JSON — no markdown, no prose."
)


# =============================================================================
#  MEGA CALL — bullet-proof extraction with normalizer + regex-only fallback
# =============================================================================

# NOTE: schemas use CONCRETE example values (not <placeholders>) so the model
# treats them as templates, not literal text to reproduce.
_CANDIDATE_SCHEMA = (
    '{\n'
    '  "candidate": {\n'
    '    "name": "Full Name",\n'
    '    "current_role": "Job Title",\n'
    '    "years_experience": 2,\n'
    '    "seniority": "Junior",\n'
    '    "domain": "Tech",\n'
    '    "education": "Degree Field",\n'
    '    "skills": [\n'
    '      {"skill": "Python", "proficiency": 8, "year_last_used": 2025, "context": "built FastAPI app"}\n'
    '    ],\n'
    '    "strengths": ["strength1", "strength2"],\n'
    '    "red_flags": []\n'
    '  },\n'
    '  "jd": {\n'
    '    "role_title": "Target Role",\n'
    '    "seniority_required": "Mid",\n'
    '    "domain": "Tech",\n'
    '    "required_skills": ["Python", "Docker"],\n'
    '    "preferred_skills": ["Kubernetes"],\n'
    '    "key_responsibilities": ["Build APIs", "Deploy services"]\n'
    '  }\n'
    '}'
)

_AUDIT_SCHEMA = (
    '{\n'
    '  "ats_score": 72,\n'
    '  "completeness_score": 68,\n'
    '  "clarity_score": 75,\n'
    '  "overall_grade": "B",\n'
    '  "ats_issues": ["Missing quantified achievements"],\n'
    '  "improvement_tips": ["Add metrics", "Use action verbs", "Add LinkedIn URL"],\n'
    '  "missing_keywords": ["Docker", "CI/CD"],\n'
    '  "interview_talking_points": ["Your Python projects", "ML deployment experience"]\n'
    '}'
)


def _parse_jd_from_text(jd_text: str) -> dict:
    """
    Pure-text JD parser — used as last resort when LLM returns no JD data.
    Extracts role title and skill keywords from raw JD text.
    """
    lines  = [l.strip() for l in jd_text.split("\n") if l.strip()]
    title  = lines[0][:60] if lines else "Target Role"

    # Simple seniority detection
    seniority = "Mid"
    for token, level in [("lead","Lead"),("senior","Senior"),("mid","Mid"),("junior","Junior")]:
        if token in jd_text.lower():
            seniority = level
            break

    # Skill keyword scan
    tech_keywords = [
        "Python","JavaScript","React","Docker","Kubernetes","AWS","GCP","Azure",
        "SQL","PostgreSQL","MongoDB","FastAPI","Django","Flask","Node","TypeScript",
        "Machine Learning","Deep Learning","NLP","MLOps","CI/CD","Linux","Git",
        "Terraform","Spark","Kafka","Redis","Elasticsearch","Java","Go","Rust",
    ]
    found = [k for k in tech_keywords if re.search(k, jd_text, re.IGNORECASE)]

    return {
        "role_title": title,
        "seniority_required": seniority,
        "domain": "Tech" if found else "Non-Tech",
        "required_skills": found[:8],
        "preferred_skills": found[8:12],
        "key_responsibilities": [],
    }


def _build_candidate_from_regex(resume_text: str, name: str = "Unknown") -> dict:
    """
    Build a synthetic candidate dict using only the regex scanner.
    Used when the LLM completely fails to extract skills.
    """
    stub = {"name": name, "current_role": "Professional",
            "years_experience": 2, "seniority": "Junior",
            "domain": "Tech", "education": "", "skills": [],
            "strengths": [], "red_flags": []}
    stub = _apply_regex_skill_fallback(stub, resume_text)
    return stub


def _normalize_llm_response(raw: dict, resume_text: str, jd_text: str) -> dict:
    """
    The LLM sometimes returns:
      - The correct {"candidate":{}, "jd":{}} wrapper  → use as-is
      - A flat dict with candidate fields at top level  → wrap it
      - Just {"skills":[...]} or {"name":"...",...}     → wrap it
      - An empty {} or partially filled dict             → patch missing parts
      - {"candidate": "some text", ...}                 → string value, ignore

    This function always returns {"candidate":{...}, "jd":{...}}.
    """
    if not isinstance(raw, dict):
        raw = {}

    # Extract candidate and jd, but only if they are dicts
    raw_c = raw.get("candidate")
    raw_j = raw.get("jd")
    candidate = raw_c if isinstance(raw_c, dict) else None
    jd        = raw_j if isinstance(raw_j, dict) else None

    # Already correct shape
    if candidate and jd:
        return {"candidate": candidate, "jd": jd}

    # Has candidate but no jd
    if candidate and not jd:
        return {"candidate": candidate, "jd": _parse_jd_from_text(jd_text)}

    # Has jd but no candidate
    if jd and not candidate:
        return {"candidate": _build_candidate_from_regex(resume_text), "jd": jd}

    # Flat structure: candidate fields at top level (has "name" or "skills")
    if "name" in raw or "skills" in raw:
        candidate_keys = {"name","current_role","years_experience","seniority",
                          "domain","education","skills","strengths","red_flags"}
        jd_keys        = {"role_title","seniority_required","required_skills",
                          "preferred_skills","key_responsibilities"}
        candidate_part = {k: raw[k] for k in candidate_keys if k in raw}
        jd_part        = {k: raw[k] for k in jd_keys        if k in raw}
        # Sanitize skills in candidate_part
        if "skills" in candidate_part and not isinstance(candidate_part["skills"], list):
            candidate_part["skills"] = []
        if not candidate_part.get("skills"):
            candidate_part = _build_candidate_from_regex(
                resume_text, candidate_part.get("name", "Unknown")
            )
        return {
            "candidate": candidate_part or _build_candidate_from_regex(resume_text),
            "jd":        jd_part        or _parse_jd_from_text(jd_text),
        }

    # Completely empty or unrecognised shape — full regex fallback
    return {
        "candidate": _build_candidate_from_regex(resume_text),
        "jd":        _parse_jd_from_text(jd_text),
    }


def _build_candidate_prompt(resume: str, jd: str, max_skills: int = 15) -> str:
    """
    Prompt that is explicit about the required wrapper keys.
    Concrete examples prevent the model from reproducing the schema literally.
    """
    return (
        f"You are a resume parser. Extract data and return a JSON object.\n\n"
        f"RESUME TEXT:\n{resume}\n\n"
        f"JOB DESCRIPTION:\n{jd}\n\n"
        f"Rules:\n"
        f"- The JSON MUST have exactly two top-level keys: \"candidate\" and \"jd\"\n"
        f"- List at most {max_skills} skills from the resume (most relevant first)\n"
        f"- Proficiency 0-10: expert=9, proficient=7, familiar=5, basic=3\n"
        f"- year_last_used: most recent year the skill appeared in the resume\n"
        f"- context: max 6 words of evidence from the resume\n"
        f"- Only include skills that are EXPLICITLY mentioned in the resume\n\n"
        f"Return ONLY the JSON object below, filled with real data (no extra text):\n"
        f"{_CANDIDATE_SCHEMA}"
    )


def mega_call(resume_text: str, jd_text: str,
              modules_hint=None,
              resume_image_b64=None) -> dict:

    _FATAL = {"rate_limited", "GROQ_API_KEY not set"}

    # ── VISION PATH ──────────────────────────────────────────────────────────
    if resume_image_b64:
        prompt = (
            "You are a resume parser. Read this resume image carefully.\n\n"
            f"JOB DESCRIPTION:\n{jd_text[:1200]}\n\n"
            f"Rules: List max 15 skills visible in the image. Context max 6 words.\n"
            f"The JSON MUST have top-level keys 'candidate' and 'jd'.\n\n"
            f"Return ONLY this JSON filled with real data:\n{_CANDIDATE_SCHEMA}"
        )
        raw = _groq_call_vision(prompt=prompt, system=_VISION_SYS,
                                 image_b64=resume_image_b64, max_tokens=4000)
        if "error" in raw:
            return raw
        result = _normalize_llm_response(raw, "", jd_text)

    # ── TEXT PATH — 3-tier LLM + regex final fallback ────────────────────────
    else:
        resume_trunc = resume_text[:3000]
        jd_trunc     = jd_text[:1000]

        # Tier 1: json_object mode, 3500 tokens (reliable on Groq free tier)
        p = _build_candidate_prompt(resume_trunc, jd_trunc, max_skills=15)
        r = _groq_call(prompt=p, system=_MEGA_SYS, model=MODEL_FAST, max_tokens=3500)

        # Tier 2: no forced format (catches 400 / json_parse_failed)
        if "error" in r and r["error"] not in _FATAL:
            r = _groq_call_no_format(prompt=p, system=_MEGA_SYS,
                                      model=MODEL_FAST, max_tokens=3500)

        # Tier 3: shorter resume, no format, 10 skill cap
        if "error" in r and r["error"] not in _FATAL:
            p3 = _build_candidate_prompt(resume_trunc[:1500], jd_trunc[:600], max_skills=10)
            r  = _groq_call_no_format(prompt=p3, system=_MEGA_SYS,
                                       model=MODEL_FAST, max_tokens=2500)

        # Fatal errors (rate limit / no key) — surface immediately
        if "error" in r and r["error"] in _FATAL:
            return r

        # Tier 4 (final): regex-only — never fails, always returns something useful
        # Used when LLM errored or returned empty/malformed JSON
        if "error" in r:
            r = {}  # will be caught by normalizer → full regex fallback

        # Normalize whatever shape the LLM returned
        result = _normalize_llm_response(r, resume_text, jd_text)

    # ── AUDIT (separate, non-fatal — small delay prevents rate limit burst) ──
    time.sleep(0.5)
    audit_raw = _groq_call(
        prompt=(
            f"Audit this resume against the job description.\n\n"
            f"RESUME:\n{resume_text[:1600] if resume_text else '[image resume]'}\n\n"
            f"JOB DESCRIPTION:\n{jd_text[:800]}\n\n"
            f"Return ONLY this JSON (fill with real values):\n{_AUDIT_SCHEMA}"
        ),
        system=_AUDIT_SYS, model=MODEL_FAST, max_tokens=600,
    )

    return {
        "candidate": result.get("candidate", {}),
        "jd":        result.get("jd", {}),
        "audit":     {} if "error" in audit_raw else audit_raw,
        "reasoning": {},
    }


# =============================================================================
#  RESUME REWRITE
# =============================================================================
def rewrite_resume(resume_text: str, jd: dict, missing_kw: List[str]) -> str:
    system = (
        "You are an expert ATS resume writer. RULES:\n"
        "1. NEVER add any skill not in the original resume.\n"
        "2. NEVER write gap admissions or cover letter language.\n"
        "3. ONLY restructure and rephrase EXISTING content.\n"
        "4. Add missing keywords ONLY where honestly supported.\n"
        'Return JSON: {"rewritten_resume": "<full text>"}'
    )
    prompt = (
        f"Rewrite this resume for ATS. Only use EXISTING facts.\n"
        f"Add these keywords where honestly supported: {missing_kw[:8]}\n\n"
        f"ORIGINAL:\n{resume_text[:2500]}\n\n"
        f"TARGET ROLE: {jd.get('role_title','--')}\n"
        f"REQUIRED SKILLS: {jd.get('required_skills', [])}\n\n"
        f'Return JSON: {{"rewritten_resume": "<full rewritten resume>"}}'
    )
    r = _groq_call(prompt=prompt, system=system, model=MODEL_FAST, max_tokens=2000)
    return r.get("rewritten_resume", "Could not rewrite resume.")


# =============================================================================
#  WEB SEARCH FEATURES
# =============================================================================
def generate_reasoning(path: List[dict], candidate: dict, jd: dict) -> Dict[str, str]:
    if not GROQ_CLIENT or not path:
        return {}
    cname  = candidate.get("name", "the candidate")
    crole  = candidate.get("current_role", "")
    known  = [s["skill"] for s in candidate.get("skills", [])
              if int(s.get("proficiency") or 0) >= 7]
    gaps   = [m["gap_skill"] for m in path if m.get("is_required")][:6]
    role   = jd.get("role_title", "")
    mods   = "\n".join(
        f'- {m["id"]}: {m["title"]} (skill:{m["skill"]},crit:{m.get("is_critical",False)})'
        for m in path[:12]
    )
    prompt = (
        f"Candidate: {cname}, {crole}. Target: {role}.\n"
        f"Known: {known}. Gaps: {gaps}.\n\n"
        f"For each module below, write exactly 2 sentences:\n"
        f"  Sentence 1: Why this candidate specifically needs it (reference their background).\n"
        f"  Sentence 2: How completing it moves them toward the target role.\n\n"
        f"Modules:\n{mods}\n\n"
        f"Return JSON like this example:\n"
        f'{{"reasoning": {{"PY01": "Candidate has basic Python from projects. Intermediate OOP is required to build production FastAPI services for this role.", "DO02": "No Docker experience found. Containerization is mandatory for deployment in this role."}}}}'
    )
    r = _groq_call(
        prompt=prompt,
        system="Expert L&D advisor. Return JSON only.",
        model=MODEL_FAST, max_tokens=1000,
    )
    return r.get("reasoning", {}) if "error" not in r else {}


def search_real_salary(role: str, location: str) -> dict:
    results = ddg_search(f"{role} salary {location} 2025 average annual", max_results=6)
    if not results:
        return {}
    snippets = "\n".join(
        f"- {r.get('title','')}: {r.get('body','')[:200]}" for r in results[:5]
    )
    r = _groq_call(
        f'Extract salary for "{role}" in {location} from:\n{snippets}\n\n'
        f'Return JSON: {{"min_lpa":<n>,"max_lpa":<n>,"median_lpa":<n>,'
        f'"currency":"INR or USD","source":"<site>","note":"<caveat>"}}',
        system="Salary extractor. JSON only.",
        model=MODEL_FAST, max_tokens=300,
    )
    if "error" in r:
        return {}
    try:
        median = float(r.get("median_lpa") or 0)
    except (TypeError, ValueError):
        median = 0.0
    return r if median > 0 else {}


def search_course_links(skill: str) -> List[dict]:
    results = ddg_search(
        f"{skill} online course english 2024 2025 "
        "site:coursera.org OR site:udemy.com OR site:youtube.com OR site:edx.org",
        max_results=8,
    )
    courses = []
    for r in results:
        url   = r.get("href", "")
        title = r.get("title", "")
        body  = r.get("body", "")
        if not url or not title or not _is_english(title):
            continue
        if   "coursera.org" in url: plat, icon = "Coursera", "🎓"
        elif "udemy.com"    in url: plat, icon = "Udemy",    "🎯"
        elif "youtube.com"  in url: plat, icon = "YouTube",  "▶"
        elif "edx.org"      in url: plat, icon = "edX",      "📘"
        elif "linkedin.com" in url: plat, icon = "LinkedIn", "💼"
        else:
            continue
        courses.append({"title": title[:65], "url": url,
                         "platform": plat, "icon": icon, "snippet": body[:120]})
    return courses[:4]


def search_skill_trends(skills: List[str]) -> Dict[str, str]:
    out = {skill: demand_label(skill) for skill in skills}
    if not skills:
        return out
    try:
        results = ddg_search(
            f"most in-demand tech skills 2025 hiring india {' '.join(skills[:6])}",
            max_results=5,
        )
        if results:
            text = " ".join(
                r.get("body", "") for r in results if _is_english(r.get("body", ""))
            ).lower()
            for skill in skills:
                if text.count(skill.lower()) >= 3 and out[skill] != "🔥 Hot":
                    out[skill] = "🔥 Hot"
    except Exception:
        pass
    return out


def search_job_market(role: str) -> List[str]:
    results = ddg_search(f"{role} job market hiring demand 2025 india", max_results=5)
    if not results:
        return []
    snippets = "\n".join(
        r.get("body", "")[:300] for r in results[:5] if _is_english(r.get("body", ""))
    )
    if not snippets:
        return []
    r = _groq_call(
        f'Give 3 short job market insights for "{role}" based on:\n{snippets}\n\n'
        f'Return JSON: {{"insights":["<i1>","<i2>","<i3>"]}}',
        system="Job market analyst. English only. JSON only.",
        model=MODEL_FAST, max_tokens=300,
    )
    return r.get("insights", []) if "error" not in r else []


# =============================================================================
#  CACHE
# =============================================================================
def _ckey(r: str, j: str) -> str:
    return hashlib.md5((r + "||" + j).encode()).hexdigest()


def cache_get(r: str, j: str) -> Any:
    try:
        with shelve.open(_CACHE_PATH) as db:
            return db.get(_ckey(r, j))
    except Exception:
        return None


def cache_set(r: str, j: str, v: Any) -> None:
    try:
        with shelve.open(_CACHE_PATH) as db:
            db[_ckey(r, j)] = v
    except Exception:
        pass


def cache_bust(resume_text: str, resume_img: Optional[str], jd_text: str) -> None:
    try:
        ck = (
            "txt:" + hashlib.md5(resume_text.encode()).hexdigest()
            if resume_text.strip()
            else "img:" + hashlib.md5((resume_img or "").encode()).hexdigest()
        )
        with shelve.open(_CACHE_PATH) as db:
            kk = _ckey(ck, jd_text)
            if kk in db:
                del db[kk]
    except Exception:
        pass


# =============================================================================
#  ANALYSIS ENGINE
# =============================================================================
def _normalize_skill(skill: str) -> str:
    sl = (
        skill.lower()
        .replace(".js", "").replace(".ts", "")
        .replace("(", "").replace(")", "").strip()
    )
    return SKILL_ALIASES.get(sl, sl)


def _match_skill(skill: str) -> int:
    sl = _normalize_skill(skill)
    for i, cs in enumerate(CATALOG_SKILLS):
        if sl == cs or sl in cs or cs in sl:
            return i
    if SEMANTIC and _ST is not None and _CEMBS is not None:
        try:
            from sklearn.metrics.pairwise import cosine_similarity
            import numpy as np
            sims = cosine_similarity(_ST.encode([sl]), _CEMBS)[0]
            best = int(np.argmax(sims))
            if sims[best] >= 0.52:
                return best
        except Exception:
            pass
    tokens         = set(sl.split())
    best_s, best_i = 0.0, -1
    for i, cs in enumerate(CATALOG_SKILLS):
        ov = len(tokens & set(cs.split())) / max(len(tokens), 1)
        if ov > best_s:
            best_s, best_i = ov, i
    return best_i if best_s >= 0.4 else -1


def skill_decay(p: Any, yr: Any) -> Tuple[int, bool]:
    try:
        p = int(p or 0)
    except (TypeError, ValueError):
        p = 0
    try:
        yr = int(yr or 0)
    except (TypeError, ValueError):
        yr = 0
    if yr <= 0 or yr >= CURRENT_YEAR - 1:
        return p, False
    yrs = CURRENT_YEAR - yr
    if yrs <= 2:
        return p, False
    adjusted = round(p * max(0.5, 1 - yrs / 5))
    return adjusted, adjusted < p


def _build_candidate_skill_lookup(candidate: dict) -> Dict[str, dict]:
    rs: Dict[str, dict] = {}
    for s in candidate.get("skills", []):
        raw = s["skill"]
        sl  = raw.lower().replace(".js","").replace(".ts","").replace("(","").replace(")","").strip()
        rs[sl] = s
        alias = SKILL_ALIASES.get(sl)
        if alias:
            rs[alias] = s
        if "mern" in sl:
            for implied in MERN_IMPLIES:
                if implied not in rs:
                    rs[implied] = {
                        **s, "skill": implied,
                        "context": f"Derived from MERN Stack: {s.get('context','')}",
                    }
        rs[sl.replace(".", "")] = s
    return rs


def analyze_gap(candidate: dict, jd: dict) -> List[dict]:
    rs    = _build_candidate_skill_lookup(candidate)
    all_s = (
        [(s, True)  for s in jd.get("required_skills",  [])] +
        [(s, False) for s in jd.get("preferred_skills", [])]
    )
    out = []
    for skill, req in all_s:
        sl_raw = skill.lower().replace(".js","").replace(".ts","").replace("(","").replace(")","").strip()
        sl     = SKILL_ALIASES.get(sl_raw, sl_raw)
        status, prof, ctx, dec, orig = "Missing", 0, "", False, 0
        src = (
            rs.get(sl) or rs.get(sl_raw) or
            next((v for k, v in rs.items()
                  if sl in k or k in sl or sl_raw in k or k in sl_raw), None)
        )
        if src:
            try:
                raw_p = int(src.get("proficiency") or 0)
            except (TypeError, ValueError):
                raw_p = 0
            prof, dec = skill_decay(raw_p, src.get("year_last_used") or 0)
            orig      = raw_p
            ctx       = _strip_mern_prefix(src.get("context", ""))
            status    = "Known" if prof >= 7 else ("Partial" if prof > 0 else "Missing")
        idx    = _match_skill(skill)
        demand = MARKET_DEMAND.get(sl, MARKET_DEMAND.get(skill.lower(), 1))
        obs    = OBSOLESCENCE_RISK.get(sl)
        out.append({
            "skill": skill, "status": status, "proficiency": prof,
            "original_prof": orig, "decayed": dec, "is_required": req,
            "context": ctx,
            "catalog_course": CATALOG[idx] if idx >= 0 else None,
            "demand": demand, "obsolescence_risk": obs,
        })
    return out


def seniority_check(c: dict, jd: dict) -> dict:
    cs  = c.get("seniority", "Mid")
    rs  = jd.get("seniority_required", "Mid")
    gap = SENIORITY_MAP.get(rs, 1) - SENIORITY_MAP.get(cs, 1)
    return {
        "has_mismatch":   gap > 0,
        "gap_levels":     gap,
        "candidate":      cs,
        "required":       rs,
        "add_leadership": gap >= 1,
        "add_strategic":  gap >= 2,
    }


def build_path(gp: List[dict], c: dict, jd: Optional[dict] = None) -> List[dict]:
    needed:  set  = set()
    id2gap:  dict = {}

    known_prof: Dict[str, int] = {}
    for skill_entry in c.get("skills", []):
        sl   = skill_entry["skill"].lower()
        p, _ = skill_decay(skill_entry.get("proficiency", 0), skill_entry.get("year_last_used", 0))
        known_prof[sl] = p
        alias = SKILL_ALIASES.get(sl)
        if alias:
            known_prof[alias] = p

    for g in gp:
        if g["status"] == "Known":
            continue
        co = g.get("catalog_course")
        if not co:
            continue
        candidate_prof = g.get("proficiency", 0)
        if candidate_prof >= 5 and co["level"] == "Beginner":
            better = next(
                (cat for cat in CATALOG
                 if cat["skill"].lower() == co["skill"].lower()
                 and cat["level"] == "Intermediate"),
                None,
            )
            if better:
                co = better
        needed.add(co["id"])
        id2gap[co["id"]] = g
        try:
            for anc in nx.ancestors(SKILL_GRAPH, co["id"]):
                ad = CATALOG_BY_ID.get(anc)
                if not ad:
                    continue
                anc_sl = ad["skill"].lower()
                if known_prof.get(anc_sl, 0) >= 6:
                    continue
                alias_known = any(
                    known_prof.get(av, 0) >= 6
                    for ak, av in SKILL_ALIASES.items()
                    if av == anc_sl or ak == anc_sl
                )
                if alias_known:
                    continue
                if not any(x["status"] == "Known" and x["skill"].lower() in anc_sl for x in gp):
                    needed.add(anc)
        except Exception:
            pass

    if jd:
        sm               = seniority_check(c, jd)
        responsibilities = " ".join(jd.get("key_responsibilities", [])).lower()
        domain           = jd.get("domain", "Tech")
        needs_leadership = sm["add_leadership"] and (
            domain != "Tech"
            or "team"  in responsibilities
            or "lead"  in responsibilities
            or "manag" in responsibilities
            or sm["gap_levels"] >= 2
        )
        if needs_leadership:
            needed.update(["LD01", "LD02"])
        if sm["add_strategic"]:
            needed.add("LD03")

    sub = SKILL_GRAPH.subgraph(needed)
    try:
        ordered = list(nx.topological_sort(sub))
    except Exception:
        ordered = list(needed)

    crit: set = set()
    try:
        if sub.nodes:
            required_gap_skills = {g["skill"].lower() for g in gp if g["is_required"]}
            required_course_ids = {
                node for node in sub.nodes
                if CATALOG_BY_ID.get(node, {}).get("skill", "").lower() in required_gap_skills
            }
            for req_id in required_course_ids:
                crit.add(req_id)
                try:
                    crit.update(nx.ancestors(sub, req_id))
                except Exception:
                    pass
            for nid in list(crit):
                co_skill = CATALOG_BY_ID.get(nid, {}).get("skill", "").lower()
                if co_skill in required_gap_skills:
                    continue
                try:
                    descs   = nx.descendants(sub, nid)
                    has_req = any(
                        CATALOG_BY_ID.get(d, {}).get("skill", "").lower() in required_gap_skills
                        for d in descs
                    )
                    if not has_req:
                        crit.discard(nid)
                except Exception:
                    pass
    except Exception:
        pass

    path, seen = [], set()
    for cid in ordered:
        if cid in seen:
            continue
        seen.add(cid)
        co = CATALOG_BY_ID.get(cid)
        if not co:
            continue
        g = id2gap.get(cid, {})
        try:
            prio_prof = int(g.get("proficiency") or 0)
        except (TypeError, ValueError):
            prio_prof = 0
        path.append({
            **co,
            "gap_skill":   g.get("skill", co["skill"]),
            "gap_status":  g.get("status", "Prereq"),
            "priority":    (0 if g.get("is_required") else 1, prio_prof),
            "reasoning":   "",
            "is_critical": cid in crit,
            "demand":      g.get("demand", 1),
            "is_required": g.get("is_required", False),
        })
    path.sort(key=lambda x: x["priority"])
    return path


def calc_impact(gp: List[dict], path: List[dict]) -> dict:
    tot   = len(gp)
    known = sum(1 for g in gp if g["status"] == "Known")
    jd_skills_lower = {g["skill"].lower() for g in gp}
    covered = len({
        m["gap_skill"].lower() for m in path
        if m["gap_skill"].lower() in jd_skills_lower
    })
    rhrs = sum(int(m.get("duration_hrs") or 0) for m in path)

    if rhrs <= 40:
        completion_prob = 0.85
    elif rhrs <= 80:
        completion_prob = 0.72
    else:
        completion_prob = 0.58

    cur  = min(100, round(known / max(tot, 1) * 100))
    proj = min(92, round(
        (known / max(tot, 1) * 100) +
        (covered / max(tot, 1) * 100 * completion_prob)
    ))
    hours_saved = max(0, 60 - rhrs)
    return {
        "total_skills":   tot,
        "known_skills":   known,
        "gaps_addressed": covered,
        "roadmap_hours":  rhrs,
        "hours_saved":    hours_saved,
        "roadmap_longer": rhrs > 60,
        "current_fit":    cur,
        "projected_fit":  proj,
        "fit_delta":      proj - cur,
        "modules_count":  len(path),
        "critical_count": sum(1 for m in path if m.get("is_critical")),
        "decayed_skills": sum(1 for g in gp if g.get("decayed")),
    }


def interview_readiness(gp: List[dict], c: dict) -> dict:
    rk  = [g for g in gp if g["status"] == "Known"   and g["is_required"]]
    rp  = [g for g in gp if g["status"] == "Partial"  and g["is_required"]
           and int(g.get("proficiency") or 0) > 0]
    rm  = [g for g in gp if (
        g["status"] == "Missing" or
        (g["status"] == "Partial" and int(g.get("proficiency") or 0) == 0)
    ) and g["is_required"]]
    tot = max(len(rk) + len(rp) + len(rm), 1)
    adj = {"Junior": 5, "Mid": 0, "Senior": -5, "Lead": -10}.get(c.get("seniority", "Mid"), 0)
    sc  = max(0, min(100, round((len(rk) + len(rp) * 0.4) / tot * 100) + adj))
    if   sc >= 75: v = ("Strong",    "#4ade80", "Ready for most rounds")
    elif sc >= 50: v = ("Moderate",  "#fbbf24", "Pass screening; prep gaps")
    elif sc >= 30: v = ("Weak",      "#fb923c", "Gap work before applying")
    else:          v = ("Not Ready", "#f87171", "Significant prep needed")
    return {
        "score": sc, "label": v[0], "color": v[1], "advice": v[2],
        "req_known": len(rk), "req_partial": len(rp), "req_missing": len(rm),
    }


def weekly_plan(path: List[dict], hpd: float = 2.0) -> List[dict]:
    try:
        hpd = max(float(hpd), 0.5)
    except (TypeError, ValueError):
        hpd = 2.0
    cap               = hpd * 5
    weeks, cur, hrs, wn = [], [], 0.0, 1
    for m in path:
        try:
            rem = float(int(m.get("duration_hrs") or 0))
        except (TypeError, ValueError):
            rem = 0.0
        while rem > 0:
            avail = cap - hrs
            if avail <= 0:
                weeks.append({"week": wn, "modules": cur, "total_hrs": hrs})
                cur, hrs, wn = [], 0.0, wn + 1
                avail = cap
            chunk = min(rem, avail)
            ex    = next((x for x in cur if x["id"] == m["id"]), None)
            if ex:
                ex["hrs_this_week"] += chunk
            else:
                cur.append({
                    "id": m["id"], "title": m["title"], "level": m["level"],
                    "domain": m["domain"], "is_critical": m.get("is_critical", False),
                    "hrs_this_week": chunk,
                    "total_hrs": int(m.get("duration_hrs") or 0),
                })
            hrs += chunk
            rem -= chunk
    if cur:
        weeks.append({"week": wn, "modules": cur, "total_hrs": hrs})
    return weeks


def transfer_map_calc(c: dict, gp: List[dict]) -> List[dict]:
    known: set = set()
    for s in c.get("skills", []):
        if int(s.get("proficiency") or 0) >= 6:
            sl = s["skill"].lower().replace(".js","").replace("(","").replace(")","").strip()
            known.add(sl)
            alias = SKILL_ALIASES.get(sl)
            if alias:
                known.add(alias)
            if "mern" in sl:
                known.update(["react", "javascript", "rest apis"])
    out = []
    seen_destinations: set = set()
    for g in gp:
        if g["status"] == "Known":
            continue
        sl = g["skill"].lower()
        best_pct, best_k = 0, None
        for k in known:
            pct = TRANSFER_MAP.get(k, {}).get(sl, 0)
            if pct > best_pct:
                best_pct, best_k = pct, k
        if best_k and best_pct > 0 and sl not in seen_destinations:
            seen_destinations.add(sl)
            out.append({
                "gap_skill":    g["skill"],
                "known_skill":  best_k.title(),
                "transfer_pct": best_pct,
                "label": f"Your {best_k.title()} → {best_pct}% head start on {g['skill']}",
                "strength": "Strong" if best_pct >= 50 else "Moderate" if best_pct >= 30 else "Partial",
            })
    return sorted(out, key=lambda x: x["transfer_pct"], reverse=True)[:4]


def roi_rank(gp: List[dict], path: List[dict]) -> List[dict]:
    out = []
    for m in path:
        g   = next((x for x in gp if x["skill"] == m.get("gap_skill")), {})
        hrs = max(int(m.get("duration_hrs") or 1), 1)
        roi = round((g.get("demand", 1) * (1.5 if g.get("is_required") else 1) * 10) / hrs, 2)
        out.append({
            "id": m["id"], "title": m["title"], "skill": m["skill"],
            "roi": roi, "hrs": int(m.get("duration_hrs") or 0),
            "is_required": g.get("is_required", False),
        })
    return sorted(out, key=lambda x: x["roi"], reverse=True)


def weeks_ready(hrs: float, hpd: float) -> str:
    try:
        hpd = max(float(hpd), 0.5)
        hrs = float(hrs)
    except (TypeError, ValueError):
        return "–"
    if hrs <= 0:
        return "0d"
    w = (hrs / hpd) / 5
    if w < 1:  return f"{int(hrs / hpd)}d"
    if w < 4:  return f"{w:.1f}w"
    return f"{(w / 4):.1f}mo"


# =============================================================================
#  ICS CALENDAR
# =============================================================================
def build_ics_calendar(path: List[dict], hpd: float = 2.0,
                       start_date: Optional[Any] = None) -> str:
    from datetime import datetime as dt, timedelta

    if start_date is None:
        base       = dt.now().replace(hour=19, minute=0, second=0, microsecond=0)
        days_ahead = (7 - base.weekday()) % 7 or 7
        start_date = base + timedelta(days=days_ahead)

    try:
        session_hrs = min(max(float(hpd), 0.5), 2.0)
    except (TypeError, ValueError):
        session_hrs = 2.0

    start_hour   = int(start_date.hour)
    start_minute = int(start_date.minute)

    lines = [
        "BEGIN:VCALENDAR", "VERSION:2.0",
        "PRODID:-//SkillForge//AI Adaptive Onboarding//EN",
        "CALSCALE:GREGORIAN", "METHOD:PUBLISH",
        "X-WR-CALNAME:SkillForge Learning Roadmap",
        "X-WR-TIMEZONE:Asia/Kolkata",
    ]

    current_day = start_date.date()
    uid_counter = 0

    for m in path:
        total_hrs = int(m.get("duration_hrs") or 0)
        sessions  = max(1, round(total_hrs / session_hrs))
        title     = m["title"]
        skill     = m["skill"]
        crit_tag  = "  ★" if m.get("is_critical") else ""

        for sess in range(sessions):
            while current_day.weekday() >= 5:
                current_day += timedelta(days=1)

            s_start = dt.combine(
                current_day,
                dt.min.time().replace(hour=start_hour, minute=start_minute, second=0),
            )
            s_end       = s_start + timedelta(hours=session_hrs)
            uid_counter += 1
            dtstamp     = dt.utcnow().strftime("%Y%m%dT%H%M%SZ")

            lines += [
                "BEGIN:VEVENT",
                f"UID:skillforge-{uid_counter}@skillforge.ai",
                f"DTSTAMP:{dtstamp}",
                f"DTSTART;TZID=Asia/Kolkata:{s_start.strftime('%Y%m%dT%H%M%S')}",
                f"DTEND;TZID=Asia/Kolkata:{s_end.strftime('%Y%m%dT%H%M%S')}",
                f"SUMMARY:📚 {title}{crit_tag}",
                f"DESCRIPTION:Skill: {skill}\\nLevel: {m['level']}"
                f"\\nSession {sess+1}/{sessions}"
                f"\\n{(m.get('reasoning') or '')[:100]}",
                f"CATEGORIES:SkillForge,{skill}",
                "STATUS:CONFIRMED",
                "END:VEVENT",
            ]
            current_day += timedelta(days=1)

    lines.append("END:VCALENDAR")
    return "\r\n".join(lines)


# =============================================================================
#  FULL PIPELINE
# =============================================================================
def run_analysis(resume_text: str, jd_text: str,
                 resume_image_b64: Optional[str] = None) -> dict:
    if resume_text and resume_text.strip():
        cache_k = "txt:" + hashlib.md5(resume_text.encode()).hexdigest()
    elif resume_image_b64:
        cache_k = "img:" + hashlib.md5(resume_image_b64.encode()).hexdigest()
    else:
        cache_k = "empty"

    cached = cache_get(cache_k, jd_text)
    if cached:
        cached["_cache_hit"] = True
        return cached

    raw = mega_call(resume_text=resume_text, jd_text=jd_text,
                    resume_image_b64=resume_image_b64)
    if "error" in raw:
        return raw

    candidate = raw.get("candidate", {})
    jd_data   = raw.get("jd",        {})
    quality   = raw.get("audit",     {})

    # ── Type-safety: LLM sometimes returns a string/list instead of dict ─────
    if not isinstance(candidate, dict): candidate = {}
    if not isinstance(jd_data,   dict): jd_data   = {}
    if not isinstance(quality,   dict): quality   = {}

    # ── Patch missing candidate/jd rather than hard-failing ──────────────────
    if not candidate:
        candidate = _build_candidate_from_regex(resume_text or "")
    if not jd_data:
        jd_data = _parse_jd_from_text(jd_text)

    # Sanitize skills list — must be a list of dicts
    skills = candidate.get("skills", [])
    if not isinstance(skills, list):
        skills = []
    clean_skills = []
    for s in skills:
        if isinstance(s, dict) and "skill" in s:
            clean_skills.append(s)
    candidate["skills"] = clean_skills

    # Always apply regex scanner to catch skills the LLM missed
    if resume_text and resume_text.strip():
        candidate = _apply_regex_skill_fallback(candidate, resume_text)

    # If still no skills at all, that's a genuine failure
    if not candidate.get("skills"):
        return {"error": "analysis_quality_failure — no skills could be extracted. "
                         "Please paste the resume as plain text and try again."}

    gp   = analyze_gap(candidate, jd_data)
    path = build_path(gp, candidate, jd_data)

    rsn_map = generate_reasoning(path, candidate, jd_data)
    for m in path:
        llm_rsn    = rsn_map.get(m["id"])
        m["reasoning"] = (
            llm_rsn if llm_rsn
            else f"Addresses gap in {m['gap_skill']} — required for {jd_data.get('role_title','the role')}."
        )

    im  = calc_impact(gp, path)
    sm  = seniority_check(candidate, jd_data)
    iv  = interview_readiness(gp, candidate)
    wp  = weekly_plan(path)
    tf  = transfer_map_calc(candidate, gp)
    roi = roi_rank(gp, path)
    obs = [
        {"skill": g["skill"], "status": g["status"],
         "reason": OBSOLESCENCE_RISK[g["skill"].lower()]}
        for g in gp if OBSOLESCENCE_RISK.get(g["skill"].lower())
    ]
    cgm = max(
        0,
        (SENIORITY_MAP.get(jd_data.get("seniority_required", "Mid"), 1) -
         SENIORITY_MAP.get(candidate.get("seniority", "Mid"), 1)) * 18,
    )
    result = {
        "candidate": candidate, "jd": jd_data, "gap_profile": gp, "path": path,
        "impact": im, "seniority": sm, "quality": quality, "interview": iv,
        "weekly_plan": wp, "transfers": tf, "roi": roi, "obsolescence": obs,
        "career_months": cgm, "_cache_hit": False,
        "_is_image": bool(resume_image_b64),
    }
    cache_set(cache_k, jd_text, result)
    return result


def run_analysis_with_web(resume_text: str, jd_text: str,
                          resume_image_b64: Optional[str] = None,
                          location: str = "India") -> dict:
    result = run_analysis(resume_text, jd_text, resume_image_b64)
    if "error" in result:
        return result
    role       = result["jd"].get("role_title", "")
    gap_skills = [g["skill"] for g in result["gap_profile"] if g["status"] != "Known"][:6]
    with ThreadPoolExecutor(max_workers=3) as ex:
        sal_f   = ex.submit(search_real_salary,  role, location)
        trend_f = ex.submit(search_skill_trends, gap_skills)
        mkt_f   = ex.submit(search_job_market,   role)
    result["salary"]          = sal_f.result()
    result["skill_trends"]    = trend_f.result()
    result["market_insights"] = mkt_f.result()
    return result


# =============================================================================
#  PDF EXPORT
# =============================================================================
def build_pdf(c: dict, jd: dict, gp: List[dict], path: List[dict],
              im: dict, ql: Optional[dict] = None, iv: Optional[dict] = None) -> io.BytesIO:
    buf = io.BytesIO()
    if not REPORTLAB:
        return buf
    doc    = SimpleDocTemplate(buf, pagesize=letter,
                                topMargin=48, bottomMargin=48,
                                leftMargin=48, rightMargin=48)
    styles = getSampleStyleSheet()
    TEAL   = rl_colors.HexColor("#2dd4bf")
    BD = ParagraphStyle("BD", parent=styles["Normal"],   fontSize=10, spaceAfter=5)
    H1 = ParagraphStyle("H1", parent=styles["Title"],    fontSize=20, spaceAfter=4,  textColor=TEAL)
    H2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13, spaceAfter=6,  spaceBefore=14)
    IT = ParagraphStyle("IT", parent=styles["Normal"],   fontSize=9,  spaceAfter=4,
                         leftIndent=18, textColor=rl_colors.HexColor("#555"))
    story = [
        Paragraph("SkillForge — AI Adaptive Onboarding Report", H1),
        Paragraph(
            f"Candidate: <b>{c.get('name','--')}</b>  |  "
            f"Role: <b>{jd.get('role_title','--')}</b>  |  "
            f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')}", BD),
        Spacer(1, 14),
    ]
    if ql or iv:
        story.append(Paragraph("Scores", H2))
        rows: List[List[str]] = []
        if ql:
            rows += [
                ["ATS Score",    f"{ql.get('ats_score','--')}%"],
                ["Grade",        str(ql.get("overall_grade", "--"))],
                ["Completeness", f"{ql.get('completeness_score','--')}%"],
                ["Clarity",      f"{ql.get('clarity_score','--')}%"],
            ]
        if iv:
            rows += [
                ["Interview Ready", f"{iv['score']}% — {iv['label']}"],
                ["Known skills",    str(iv["req_known"])],
                ["Partial skills",  str(iv["req_partial"])],
                ["Missing skills",  str(iv["req_missing"])],
            ]
        t = Table([["Metric", "Value"]] + rows, colWidths=[200, 260])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), TEAL),
            ("TEXTCOLOR",  (0, 0), (-1, 0), rl_colors.white),
            ("FONTSIZE",   (0, 0), (-1, -1), 10),
            ("GRID",       (0, 0), (-1, -1), 0.4, rl_colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [rl_colors.whitesmoke, rl_colors.white]),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ]))
        story += [t, Spacer(1, 14)]
    story.append(Paragraph("Learning Roadmap", H2))
    for i, m in enumerate(path):
        crit_label = "[CRITICAL] " if m.get("is_critical") else ""
        story.append(Paragraph(
            f"<b>{i+1}. {crit_label}{m['title']}</b>"
            f" — {m['level']} / {m['duration_hrs']}h", BD))
        if m.get("reasoning"):
            story.append(Paragraph(f"→ {m['reasoning']}", IT))
    doc.build(story)
    buf.seek(0)
    return buf


# =============================================================================
#  CHARTS
# =============================================================================
_BG    = "rgba(0,0,0,0)"
_GRID  = "rgba(255,255,255,0.06)"
_TEAL  = "#2dd4bf"
_AMBER = "#f59e0b"
_RED   = "#ef4444"
_GREEN = "#4ade80"
_FONT  = dict(color="#94a3b8", family="'DM Mono', 'IBM Plex Mono', monospace")


def _bl(**kw) -> dict:
    return dict(paper_bgcolor=_BG, plot_bgcolor=_BG, font=_FONT,
                margin=dict(l=8, r=8, t=8, b=36), **kw)


def _gap_bar_chart(gp: List[dict]) -> go.Figure:
    items    = [g for g in gp if g["is_required"]][:8] or gp[:8]
    skills   = [g["skill"][:16] for g in items]
    current  = [int(g.get("proficiency") or 0) for g in items]
    required = [10] * len(items)
    fig = go.Figure()
    fig.add_trace(go.Bar(name="Required", y=skills, x=required, orientation="h",
                         marker_color=_RED,  opacity=0.15))
    fig.add_trace(go.Bar(name="Current",  y=skills, x=current,  orientation="h",
                         marker_color=_TEAL, opacity=0.85))
    fig.update_layout(
        **_bl(height=320), barmode="overlay",
        xaxis=dict(range=[0, 10], gridcolor=_GRID, tickfont=dict(size=10)),
        yaxis=dict(gridcolor="rgba(0,0,0,0)", tickfont=dict(size=10), autorange="reversed"),
        legend=dict(bgcolor=_BG, orientation="h", y=1.1, font=dict(size=10)),
    )
    return fig


def radar_chart(gp: List[dict]) -> go.Figure:
    items    = [g for g in gp if g["is_required"]][:10] or gp[:10]
    if not items:
        return go.Figure()
    all_zero = all(int(g.get("proficiency") or 0) == 0 for g in items)
    if all_zero or len([g for g in items if int(g.get("proficiency") or 0) > 0]) < 2:
        return _gap_bar_chart(gp)
    theta = [g["skill"][:14] for g in items]
    fig   = go.Figure(data=[
        go.Scatterpolar(r=[10]*len(items), theta=theta, fill="toself",
                        name="Required", line=dict(color=_RED, width=1), opacity=0.08),
        go.Scatterpolar(r=[int(g.get("proficiency") or 0) for g in items], theta=theta,
                        fill="toself", name="Current",
                        line=dict(color=_TEAL, width=2.5), opacity=0.7),
    ])
    fig.update_layout(
        **_bl(height=320),
        polar=dict(
            bgcolor=_BG,
            radialaxis=dict(visible=True, range=[0,10], gridcolor=_GRID,
                            tickfont=dict(size=9, color="#475569")),
            angularaxis=dict(gridcolor=_GRID, tickfont=dict(size=11)),
        ),
        showlegend=True, legend=dict(bgcolor=_BG, x=0.72, y=1.22, font=dict(size=10)),
    )
    return fig


def animated_radar_chart(gp: List[dict]) -> go.Figure:
    items    = [g for g in gp if g["is_required"]][:10] or gp[:10]
    if not items:
        return go.Figure()
    all_zero = all(int(g.get("proficiency") or 0) == 0 for g in items)
    if all_zero or len([g for g in items if int(g.get("proficiency") or 0) > 0]) < 2:
        return _gap_bar_chart(gp)
    theta  = [g["skill"][:14] for g in items]
    target = [int(g.get("proficiency") or 0) for g in items]
    frames = []
    for step in range(13):
        frac   = step / 12
        t_ease = 1 - (1 - frac) ** 3
        frames.append(go.Frame(
            data=[
                go.Scatterpolar(r=[10]*len(items), theta=theta, fill="toself",
                                name="Required", line=dict(color=_RED, width=1), opacity=0.08),
                go.Scatterpolar(r=[round(v * t_ease, 1) for v in target], theta=theta,
                                fill="toself", name="Current",
                                line=dict(color=_TEAL, width=2.5), opacity=0.8),
            ],
            name=str(step),
        ))
    fig = go.Figure(
        data=[
            go.Scatterpolar(r=[10]*len(items), theta=theta, fill="toself",
                            name="Required", line=dict(color=_RED, width=1), opacity=0.08),
            go.Scatterpolar(r=[0]*len(items), theta=theta, fill="toself",
                            name="Current", line=dict(color=_TEAL, width=2.5), opacity=0.8),
        ],
        frames=frames,
        layout=go.Layout(
            **_bl(height=320),
            polar=dict(
                bgcolor=_BG,
                radialaxis=dict(visible=True, range=[0,10], gridcolor=_GRID,
                                tickfont=dict(size=9, color="#475569")),
                angularaxis=dict(gridcolor=_GRID, tickfont=dict(size=11)),
            ),
            showlegend=True, legend=dict(bgcolor=_BG, x=0.72, y=1.22, font=dict(size=10)),
            updatemenus=[dict(
                type="buttons", showactive=False, y=1.28, x=0.58, xanchor="left",
                buttons=[dict(
                    label="▶ Animate",
                    method="animate",
                    args=[None, {"frame": {"duration": 60, "redraw": True},
                                 "fromcurrent": True, "transition": {"duration": 0}}],
                )],
            )],
        ),
    )
    return fig


def timeline_chart(path: List[dict]) -> go.Figure:
    if not path:
        return go.Figure()
    lc    = {"Critical": _RED, "Beginner": _TEAL, "Intermediate": _AMBER, "Advanced": "#f97316"}
    shown: set = set()
    fig   = go.Figure()
    for m in path:
        k    = "Critical" if m.get("is_critical") else m["level"]
        show = k not in shown
        shown.add(k)
        fig.add_trace(go.Bar(
            x=[int(m.get("duration_hrs") or 0)], y=[m["title"][:30]], orientation="h",
            marker=dict(color=lc.get(k, "#64748b"), opacity=0.85, line=dict(width=0)),
            name=k, legendgroup=k, showlegend=show,
            hovertemplate=f"<b>{m['title']}</b><br>{m['level']} · {m['duration_hrs']}h<extra></extra>",
        ))
    fig.update_layout(
        **_bl(height=max(260, len(path) * 38)),
        xaxis=dict(title="Hours", gridcolor=_GRID, zeroline=False, tickfont=dict(size=11)),
        yaxis=dict(gridcolor="rgba(0,0,0,0)", tickfont=dict(size=11), autorange="reversed"),
        legend=dict(bgcolor=_BG, orientation="h", y=1.05, font=dict(size=11)),
        barmode="overlay",
    )
    return fig


def salary_chart(s: dict) -> go.Figure:
    if not s:
        return go.Figure()
    try:
        med = float(s.get("median_lpa") or 0)
    except (TypeError, ValueError):
        med = 0.0
    if med <= 0:
        return go.Figure()

    def _n(v: Any) -> float:
        try:
            return float(v or 0)
        except (TypeError, ValueError):
            return 0.0

    curr      = s.get("currency", "INR")
    sym       = "₹" if curr == "INR" else "$"
    unit      = "L/yr" if curr == "INR" else "k/yr"
    raw_vals  = [_n(s.get("min_lpa")), _n(s.get("median_lpa")), _n(s.get("max_lpa"))]
    vals = [round(v / 1000, 1) for v in raw_vals] if (curr == "USD" and any(v > 500 for v in raw_vals)) else raw_vals
    lbls = [f"{sym}{v}{unit}" for v in vals]
    fig  = go.Figure(go.Bar(
        x=["Min", "Median", "Max"], y=vals,
        marker_color=[_TEAL, _AMBER, _RED], opacity=0.82,
        text=lbls, textposition="outside",
        textfont=dict(size=13, family="'DM Mono','IBM Plex Mono',monospace"),
    ))
    fig.update_layout(
        **_bl(height=230),
        yaxis=dict(title=unit, gridcolor=_GRID, zeroline=False, tickfont=dict(size=11)),
        xaxis=dict(gridcolor="rgba(0,0,0,0)", tickfont=dict(size=12)),
    )
    return fig


def roi_bar(roi_list: List[dict]) -> go.Figure:
    if not roi_list:
        return go.Figure()
    top = roi_list[:10]
    fig = go.Figure(go.Bar(
        x=[m["roi"] for m in top], y=[m["title"][:28] for m in top], orientation="h",
        marker=dict(color=[_RED if m["is_required"] else _TEAL for m in top],
                    opacity=0.85, line=dict(width=0)),
        hovertemplate="<b>%{y}</b><br>ROI Index: %{x}<extra></extra>",
    ))
    fig.update_layout(
        **_bl(height=max(200, len(top) * 36)),
        xaxis=dict(title="ROI Index", gridcolor=_GRID, zeroline=False, tickfont=dict(size=11)),
        yaxis=dict(gridcolor="rgba(0,0,0,0)", autorange="reversed", tickfont=dict(size=11)),
    )
    return fig


# =============================================================================
#  INTERVIEW QUESTIONS
# =============================================================================
def generate_interview_questions(gp: List[dict], candidate: dict, jd: dict) -> Dict[str, List[str]]:
    if not GROQ_CLIENT:
        return {}
    relevant = [g for g in gp if g["status"] in ("Known", "Partial") and g["is_required"]][:5]
    if not relevant:
        return {}
    role      = jd.get("role_title", "the role")
    seniority = candidate.get("seniority", "Junior")
    skills_list = "\n".join(
        f"- {g['skill']} ({g['proficiency']}/10, {g['status']})"
        for g in relevant
    )
    prompt = (
        f"Generate interview questions for a {seniority}-level candidate targeting {role}.\n\n"
        f"Skills to cover:\n{skills_list}\n\n"
        f"Write exactly 3 questions per skill, calibrated to {seniority} level.\n"
        f"Questions should test real depth — practical, specific, not generic.\n\n"
        f"Return JSON like this example (use actual skill names as keys, not placeholders):\n"
        f'{{"questions": {{"Python": ["Explain how Python\'s GIL affects multithreaded code.", "How would you optimize a slow pandas DataFrame operation?", "What is the difference between __init__ and __new__?"], "SQL": ["Write a query to find the second highest salary.", "When would you use a CTE vs a subquery?", "Explain the difference between INNER JOIN and LEFT JOIN."]}}}}'
    )
    r = _groq_call(
        prompt=prompt,
        system=f"Expert technical interviewer for {seniority}-level engineers. JSON only.",
        model=MODEL_FAST, max_tokens=900,
    )
    return r.get("questions", {}) if "error" not in r else {}