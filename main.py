# =============================================================================
#  main.py — SkillForge v10  |  All bugs fixed · PDF upload · None-safe · Production ready
#  Run: streamlit run main.py
# =============================================================================

import os, sys, json, io, re, time, hashlib, shelve, threading, argparse, base64, urllib.parse
from typing import Tuple, Optional, List, Dict, Any
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor

from dotenv import load_dotenv
load_dotenv()

import streamlit as st

st.set_page_config(
    page_title="SkillForge",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Inject hide-toolbar CSS immediately — must be first st call after set_page_config ──
st.markdown("""<style>
footer,#MainMenu,
header[data-testid="stHeader"],
[data-testid="stToolbar"],
[data-testid="stDecoration"],
[data-testid="stStatusWidget"],
[data-testid="stAppDeployButton"],
[data-testid="stBaseButton-header"],
[data-testid="stBaseButton-minimal"],
[data-testid="stBaseButton-borderless"],
.stDeployButton,button[kind="header"],button[kind="minimal"],
[data-testid="stHeader"] button,[data-testid="stHeader"] a,
[data-testid="stHeader"]>div,#stDecoration
{display:none!important;visibility:hidden!important;height:0!important;
 width:0!important;overflow:hidden!important;opacity:0!important}
[data-testid="stHeader"]{height:0!important;min-height:0!important;padding:0!important}
</style>""", unsafe_allow_html=True)

import plotly.graph_objects as go
import networkx as nx
import pdfplumber
from docx import Document
from groq import Groq

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors as rl_colors
    REPORTLAB = True
except Exception:
    REPORTLAB = False

# =============================================================================
#  WEB SEARCH
# =============================================================================
_DDG_ERROR: str = ""

def ddg_search(query: str, max_results: int = 5) -> List[dict]:
    global _DDG_ERROR
    try:
        from duckduckgo_search import DDGS
        with DDGS() as ddgs:
            return list(ddgs.text(query, max_results=max_results))
    except ImportError:
        _DDG_ERROR = "duckduckgo-search not installed — run `pip install duckduckgo-search`"
        return []
    except Exception as e:
        _DDG_ERROR = f"Web search unavailable: {e}"
        return []

# =============================================================================
#  SEMANTIC MATCHING
# =============================================================================
SEMANTIC, _ST, _CEMBS = False, None, None

def _load_semantic_bg():
    global SEMANTIC, _ST, _CEMBS
    try:
        from sentence_transformers import SentenceTransformer
        _ST    = SentenceTransformer("all-MiniLM-L6-v2")
        _CEMBS = _ST.encode([c["skill"].lower() for c in CATALOG])
        SEMANTIC = True
    except Exception:
        pass

# =============================================================================
#  GROQ CLIENT
# =============================================================================
_GROQ_KEY = os.getenv("GROQ_API_KEY", "")
if not _GROQ_KEY:
    st.error("**GROQ_API_KEY missing** — add it to `.env`  →  [console.groq.com](https://console.groq.com)")
    st.stop()

GROQ_CLIENT   = Groq(api_key=_GROQ_KEY)
MODEL_FAST    = "llama-3.3-70b-versatile"
# FIX #1: Use a vision-capable model for image resumes
MODEL_VISION  = "meta-llama/llama-4-scout-17b-16e-instruct"
CURRENT_YEAR  = datetime.now().year

# =============================================================================
#  CATALOG — 47 courses
# =============================================================================
CATALOG: List[Dict] = [
    {"id":"PY01","title":"Python Fundamentals","skill":"Python","domain":"Tech","level":"Beginner","duration_hrs":6,"prereqs":[]},
    {"id":"PY02","title":"Python Intermediate: OOP & Modules","skill":"Python","domain":"Tech","level":"Intermediate","duration_hrs":8,"prereqs":["PY01"]},
    {"id":"PY03","title":"Python Advanced: Async & Decorators","skill":"Python","domain":"Tech","level":"Advanced","duration_hrs":10,"prereqs":["PY02"]},
    {"id":"DA01","title":"Data Analysis with Pandas","skill":"Data Analysis","domain":"Tech","level":"Beginner","duration_hrs":7,"prereqs":["PY01"]},
    {"id":"DA02","title":"Data Visualization (Matplotlib & Seaborn)","skill":"Data Visualization","domain":"Tech","level":"Intermediate","duration_hrs":5,"prereqs":["DA01"]},
    {"id":"DA03","title":"Statistical Analysis & Hypothesis Testing","skill":"Statistics","domain":"Tech","level":"Intermediate","duration_hrs":8,"prereqs":["DA01"]},
    {"id":"ML01","title":"Machine Learning Foundations","skill":"Machine Learning","domain":"Tech","level":"Beginner","duration_hrs":10,"prereqs":["DA01","DA03"]},
    {"id":"ML02","title":"Supervised Learning: Regression & Classification","skill":"Machine Learning","domain":"Tech","level":"Intermediate","duration_hrs":12,"prereqs":["ML01"]},
    {"id":"ML03","title":"Deep Learning with PyTorch","skill":"Deep Learning","domain":"Tech","level":"Advanced","duration_hrs":15,"prereqs":["ML02"]},
    {"id":"ML04","title":"NLP & Large Language Models","skill":"NLP","domain":"Tech","level":"Advanced","duration_hrs":14,"prereqs":["ML02"]},
    {"id":"ML05","title":"MLOps & Model Deployment","skill":"MLOps","domain":"Tech","level":"Advanced","duration_hrs":10,"prereqs":["ML02","DO02"]},
    {"id":"SQL01","title":"SQL Fundamentals","skill":"SQL","domain":"Tech","level":"Beginner","duration_hrs":5,"prereqs":[]},
    {"id":"SQL02","title":"Advanced SQL: Window Functions & Optimization","skill":"SQL","domain":"Tech","level":"Advanced","duration_hrs":7,"prereqs":["SQL01"]},
    {"id":"SQL03","title":"Database Design & NoSQL","skill":"Databases","domain":"Tech","level":"Intermediate","duration_hrs":6,"prereqs":["SQL01"]},
    {"id":"DO01","title":"Linux & Bash Scripting","skill":"Linux","domain":"Tech","level":"Beginner","duration_hrs":6,"prereqs":[]},
    {"id":"DO02","title":"Docker & Containerization","skill":"Docker","domain":"Tech","level":"Intermediate","duration_hrs":7,"prereqs":["DO01"]},
    {"id":"DO03","title":"Kubernetes Orchestration","skill":"Kubernetes","domain":"Tech","level":"Advanced","duration_hrs":10,"prereqs":["DO02"]},
    {"id":"DO04","title":"CI/CD Pipelines with GitHub Actions","skill":"CI/CD","domain":"Tech","level":"Intermediate","duration_hrs":6,"prereqs":["DO01"]},
    {"id":"CL01","title":"Cloud Computing Fundamentals","skill":"Cloud Computing","domain":"Tech","level":"Beginner","duration_hrs":6,"prereqs":[]},
    {"id":"CL02","title":"AWS Core Services Deep Dive","skill":"AWS","domain":"Tech","level":"Intermediate","duration_hrs":10,"prereqs":["CL01"]},
    {"id":"CL03","title":"GCP & BigQuery for Data Engineers","skill":"GCP","domain":"Tech","level":"Intermediate","duration_hrs":8,"prereqs":["CL01","DA01"]},
    {"id":"WE01","title":"HTML & CSS Foundations","skill":"HTML/CSS","domain":"Tech","level":"Beginner","duration_hrs":5,"prereqs":[]},
    {"id":"WE02","title":"JavaScript Essentials","skill":"JavaScript","domain":"Tech","level":"Beginner","duration_hrs":8,"prereqs":["WE01"]},
    {"id":"WE03","title":"React.js Fundamentals","skill":"React","domain":"Tech","level":"Intermediate","duration_hrs":10,"prereqs":["WE02"]},
    {"id":"WE04","title":"FastAPI Backend Development","skill":"FastAPI","domain":"Tech","level":"Intermediate","duration_hrs":8,"prereqs":["PY02"]},
    {"id":"WE05","title":"Full-Stack Integration & REST APIs","skill":"REST APIs","domain":"Tech","level":"Intermediate","duration_hrs":7,"prereqs":["WE03","WE04"]},
    {"id":"SE01","title":"Cybersecurity Fundamentals","skill":"Cybersecurity","domain":"Tech","level":"Beginner","duration_hrs":6,"prereqs":[]},
    {"id":"SE02","title":"Application Security & OWASP Top 10","skill":"Application Security","domain":"Tech","level":"Intermediate","duration_hrs":8,"prereqs":["SE01"]},
    {"id":"AG01","title":"Agile & Scrum Fundamentals","skill":"Agile","domain":"Tech","level":"Beginner","duration_hrs":4,"prereqs":[]},
    {"id":"AG02","title":"Advanced Scrum Master Certification Prep","skill":"Scrum","domain":"Tech","level":"Advanced","duration_hrs":6,"prereqs":["AG01"]},
    {"id":"PM01","title":"Project Management Essentials (PMI)","skill":"Project Management","domain":"Soft","level":"Intermediate","duration_hrs":8,"prereqs":["LD01"]},
    {"id":"HR01","title":"HR Fundamentals & Employment Law","skill":"Human Resources","domain":"Non-Tech","level":"Beginner","duration_hrs":5,"prereqs":[]},
    {"id":"HR02","title":"Talent Acquisition & Recruitment","skill":"Recruitment","domain":"Non-Tech","level":"Intermediate","duration_hrs":6,"prereqs":["HR01"]},
    {"id":"HR03","title":"Performance Management & Appraisals","skill":"Performance Management","domain":"Non-Tech","level":"Intermediate","duration_hrs":5,"prereqs":["HR01"]},
    {"id":"HR04","title":"Employee Relations & Conflict Resolution","skill":"Employee Relations","domain":"Non-Tech","level":"Advanced","duration_hrs":6,"prereqs":["HR01"]},
    {"id":"HR05","title":"Learning & Development Strategy","skill":"L&D Strategy","domain":"Non-Tech","level":"Advanced","duration_hrs":6,"prereqs":["HR03"]},
    {"id":"OP01","title":"Supply Chain & Logistics Fundamentals","skill":"Logistics","domain":"Non-Tech","level":"Beginner","duration_hrs":5,"prereqs":[]},
    {"id":"OP02","title":"Warehouse Management Systems","skill":"Warehouse Management","domain":"Non-Tech","level":"Intermediate","duration_hrs":6,"prereqs":["OP01"]},
    {"id":"OP03","title":"Inventory Control & Demand Planning","skill":"Inventory Management","domain":"Non-Tech","level":"Intermediate","duration_hrs":5,"prereqs":["OP01"]},
    {"id":"OP04","title":"Lean Manufacturing & Six Sigma Green Belt","skill":"Process Improvement","domain":"Non-Tech","level":"Advanced","duration_hrs":8,"prereqs":["OP01"]},
    {"id":"FI01","title":"Financial Accounting Basics","skill":"Accounting","domain":"Non-Tech","level":"Beginner","duration_hrs":6,"prereqs":[]},
    {"id":"FI02","title":"Financial Analysis & Modeling","skill":"Financial Analysis","domain":"Non-Tech","level":"Intermediate","duration_hrs":8,"prereqs":["FI01"]},
    {"id":"FI03","title":"Budgeting & Forecasting","skill":"Budgeting","domain":"Non-Tech","level":"Intermediate","duration_hrs":5,"prereqs":["FI01"]},
    {"id":"LD01","title":"Communication & Presentation Skills","skill":"Communication","domain":"Soft","level":"Beginner","duration_hrs":4,"prereqs":[]},
    {"id":"LD02","title":"Team Leadership & People Management","skill":"Leadership","domain":"Soft","level":"Intermediate","duration_hrs":6,"prereqs":["LD01"]},
    {"id":"LD03","title":"Strategic Thinking & Decision Making","skill":"Strategic Planning","domain":"Soft","level":"Advanced","duration_hrs":6,"prereqs":["LD02"]},
    {"id":"LD04","title":"Cross-Functional Collaboration","skill":"Collaboration","domain":"Soft","level":"Beginner","duration_hrs":3,"prereqs":["LD01"]},
]
CATALOG_BY_ID  = {c["id"]: c for c in CATALOG}
CATALOG_SKILLS = [c["skill"].lower() for c in CATALOG]

MARKET_DEMAND: Dict[str, int] = {
    "python":3,"machine learning":3,"deep learning":3,"aws":3,"docker":3,
    "sql":3,"react":3,"kubernetes":3,"mlops":3,"fastapi":3,"data analysis":3,
    "cloud computing":3,"ci/cd":3,"nlp":3,"rest apis":3,"javascript":3,
    "statistics":2,"data visualization":2,"gcp":2,"linux":2,"html/css":2,
    "agile":2,"databases":2,"cybersecurity":2,"communication":2,"leadership":2,
    "project management":2,"recruitment":2,"performance management":2,
    "financial analysis":2,"application security":2,"process improvement":2,
    "human resources":1,"accounting":1,"budgeting":1,"logistics":1,"scrum":1,
    "collaboration":1,"strategic planning":1,"warehouse management":1,
    "inventory management":1,"l&d strategy":1,
}
OBSOLESCENCE_RISK: Dict[str, str] = {
    "jquery":"Replaced by vanilla JS and React",
    "php":"Declining; Python/Node dominant",
    "hadoop":"Replaced by Spark + cloud-native",
    "excel vba":"Power Query and Python replacing VBA",
    "manual testing":"AI-assisted automation replacing manual QA",
    "waterfall":"Industry fully shifted to Agile/DevOps",
}
TRANSFER_MAP: Dict[str, Dict[str, int]] = {
    "python":{"machine learning":40,"mlops":35,"fastapi":60,"data analysis":50,"deep learning":30,"rest apis":45},
    "machine learning":{"deep learning":50,"mlops":45,"nlp":40,"statistics":30},
    "javascript":{"react":55,"rest apis":40},
    "sql":{"data analysis":35,"databases":60},
    "docker":{"kubernetes":45,"ci/cd":35,"mlops":30},
    "linux":{"docker":40,"ci/cd":30,"aws":20},
    "aws":{"gcp":30,"cloud computing":70,"mlops":25},
    "human resources":{"recruitment":45,"performance management":40,"employee relations":35},
    "communication":{"leadership":35,"project management":25},
    "leadership":{"strategic planning":40},
    "financial analysis":{"budgeting":55,"accounting":40},
}
SENIORITY_MAP = {"Junior":0,"Mid":1,"Senior":2,"Lead":3}

SAMPLES = {
    "junior_swe": {
        "label": "Junior SWE → Full Stack",
        "description": "1 yr experience · Python basics · HTML/CSS · no cloud",
        "resume": """John Smith
Junior Software Developer | 1 year experience
Skills: Python (basic, 4/10), HTML/CSS, some JavaScript
Education: B.Tech Computer Science 2023
Projects: Built a todo app using Flask. Familiar with Git basics.
No professional cloud or DevOps experience.""",
        "jd": """Software Engineer Full Stack - Mid Level
Required: Python, React, FastAPI, Docker, SQL, REST APIs, AWS
Preferred: Kubernetes, CI/CD
Seniority: Mid | Domain: Tech""",
    },
    "senior_ds": {
        "label": "Senior DS → Lead AI",
        "description": "7 yr experience · Python expert · NLP & MLOps unused 2+ yrs",
        "resume": """Priya Patel
Senior Data Scientist | 7 years experience
Skills: Python (expert, 9/10), Machine Learning (expert), Deep Learning (PyTorch, 8/10), SQL (advanced, 8/10), AWS SageMaker (7/10)
Last used NLP: 2022. Last used MLOps: 2021.
Led team of 5. Published 3 ML papers.""",
        "jd": """Lead Data Scientist - AI Products
Required: Python, Machine Learning, Deep Learning, NLP, MLOps, SQL, AWS
Preferred: GCP, Kubernetes, Leadership
Seniority: Lead | Domain: Tech""",
    },
    "hr_manager": {
        "label": "HR Coordinator → Manager",
        "description": "3 yr experience · Recruitment strength · no L&D or leadership",
        "resume": """Amara Johnson
HR Coordinator | 3 years experience
Skills: Human Resources (intermediate, 6/10), Recruitment (good, 7/10), Microsoft Office
Some performance review experience. No formal L&D training.""",
        "jd": """HR Manager - People and Culture
Required: Human Resources, Recruitment, Performance Management, Employee Relations
Preferred: L&D Strategy, Communication, Leadership
Seniority: Senior | Domain: Non-Tech""",
    },
}

# =============================================================================
#  DEPENDENCY GRAPH
# =============================================================================
def _build_graph() -> nx.DiGraph:
    G = nx.DiGraph()
    for c in CATALOG:
        G.add_node(c["id"], **c)
        for p in c["prereqs"]:
            G.add_edge(p, c["id"])
    return G

SKILL_GRAPH = _build_graph()

# =============================================================================
#  FILE PARSER
# =============================================================================
def parse_uploaded_file(f) -> Tuple[str, Optional[str]]:
    if f is None: return "", None
    name = f.name.lower()
    # Seek to start — Streamlit file objects may be at EOF on rerun
    try: f.seek(0)
    except: pass
    raw = f.read()
    if not raw:
        return f"[Empty file: {name}]", None
    if name.endswith(".pdf"):
        try:
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages), None
        except Exception as e: return f"[PDF error: {e}]", None
    if name.endswith(".docx"):
        try:
            doc = Document(io.BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs), None
        except Exception as e: return f"[DOCX error: {e}]", None
    if any(name.endswith(x) for x in [".jpg",".jpeg",".png",".webp"]):
        media = ("image/jpeg" if name.endswith((".jpg",".jpeg"))
                 else "image/png" if name.endswith(".png") else "image/webp")
        return "", f"data:{media};base64,{base64.b64encode(raw).decode()}"
    return raw.decode("utf-8", errors="ignore"), None

# =============================================================================
#  GROQ CALLS
# =============================================================================
_audit_log: List[dict] = []

# FIX #2 & #19: Separate vision call from text call — vision models don't support
# response_format=json_object and need a different content format + JSON parsing
def _groq_call_vision(prompt: str, system: str, image_b64: str, max_tokens: int = 3200) -> dict:
    """Call a vision-capable Groq model for image resume analysis."""
    content = [
        {"type": "image_url", "image_url": {"url": image_b64}},
        {"type": "text", "text": prompt}
    ]
    messages = [{"role": "system", "content": system}, {"role": "user", "content": content}]
    t0 = time.time()
    try:
        r = GROQ_CLIENT.chat.completions.create(
            model=MODEL_VISION,
            messages=messages,
            temperature=0.1,
            max_tokens=max_tokens,
            # FIX #19: Vision models do NOT support response_format=json_object
            # We parse JSON from the text response manually
        )
        raw_text = r.choices[0].message.content or "{}"
        usage = r.usage
        in_tok  = usage.prompt_tokens if usage else 0
        out_tok = usage.completion_tokens if usage else 0
        cost = round((in_tok*0.00000011)+(out_tok*0.00000034),6)
        _audit_log.append({"ts":datetime.now().strftime("%H:%M:%S"),
                           "model": MODEL_VISION.split("/")[-1][:22],
                           "in":in_tok,"out":out_tok,
                           "ms":round((time.time()-t0)*1000),"cost":cost,"status":"ok"})
        # FIX: Extract JSON from response — strip markdown fences if present
        json_text = _extract_json(raw_text)
        return json.loads(json_text)
    except json.JSONDecodeError as e:
        _audit_log.append({"ts":datetime.now().strftime("%H:%M:%S"),
                           "model":MODEL_VISION.split("/")[-1][:22],
                           "status":f"json_err:{str(e)[:40]}","in":0,"out":0,"ms":0,"cost":0})
        return {"error": f"vision_json_parse_failed: {e}"}
    except Exception as e:
        err = str(e); wait_s = 0
        if "429" in err or "rate_limit" in err:
            m = re.search(r"try again in (\d+)m([\d.]+)s", err)
            if m: wait_s = int(m.group(1))*60+float(m.group(2))
            return {"error":"rate_limited","wait_seconds":int(wait_s),
                    "message":f"Rate limited. Retry in {int(wait_s//60)}m{int(wait_s%60)}s."}
        _audit_log.append({"ts":datetime.now().strftime("%H:%M:%S"),
                           "model":MODEL_VISION.split("/")[-1][:22],
                           "status":f"err:{err[:40]}","in":0,"out":0,"ms":0,"cost":0})
        return {"error": err}


def _extract_json(text: str) -> str:
    """Extract JSON object from text that may contain markdown fences or prose."""
    # Strip ```json ... ``` or ``` ... ``` fences
    text = re.sub(r"```(?:json)?\s*", "", text)
    text = text.replace("```", "").strip()
    # Find first { ... } block
    start = text.find("{")
    if start == -1:
        return "{}"
    # Find matching closing brace
    depth, end = 0, -1
    for i, ch in enumerate(text[start:], start):
        if ch == "{": depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                end = i + 1
                break
    return text[start:end] if end > start else text[start:]


def _groq_call(prompt: str, system: str, model: str = MODEL_FAST,
               max_tokens: int = 2800) -> dict:
    """Standard text-only Groq call with JSON response format."""
    messages = [{"role":"system","content":system},{"role":"user","content":prompt}]
    t0 = time.time()
    try:
        r = GROQ_CLIENT.chat.completions.create(
            model=model, messages=messages, temperature=0.1,
            max_tokens=max_tokens, response_format={"type":"json_object"},
        )
        usage = r.usage
        in_tok  = usage.prompt_tokens if usage else 0
        out_tok = usage.completion_tokens if usage else 0
        cost = round((in_tok*0.00000011)+(out_tok*0.00000034),6)
        _audit_log.append({"ts":datetime.now().strftime("%H:%M:%S"),
                           "model":model.split("/")[-1][:22],
                           "in":in_tok,"out":out_tok,
                           "ms":round((time.time()-t0)*1000),"cost":cost,"status":"ok"})
        return json.loads(r.choices[0].message.content or "{}")
    except json.JSONDecodeError:
        return {"error":"json_parse_failed"}
    except Exception as e:
        err = str(e); wait_s = 0
        if "429" in err or "rate_limit" in err:
            m = re.search(r"try again in (\d+)m([\d.]+)s", err)
            if m: wait_s = int(m.group(1))*60+float(m.group(2))
            return {"error":"rate_limited","wait_seconds":int(wait_s),
                    "message":f"Rate limited. Retry in {int(wait_s//60)}m{int(wait_s%60)}s."}
        _audit_log.append({"ts":datetime.now().strftime("%H:%M:%S"),
                           "model":model.split("/")[-1][:22],
                           "status":f"err:{err[:40]}","in":0,"out":0,"ms":0,"cost":0})
        return {"error":err}


_MEGA_SYS = """You are a world-class senior tech recruiter, ATS specialist, and L&D expert.
Extract ALL sections in ONE response as valid JSON. Be precise and evidence-based. No hallucinations.
Return ONLY the JSON object — no preamble, no markdown fences, no explanation text."""

# FIX #3 & #10: New dedicated vision system prompt for image resumes
_VISION_SYS = """You are an expert resume parser and tech recruiter. The user has uploaded an IMAGE of their resume.

STEP 1 — OCR: Read every piece of text visible in the resume image carefully, including:
  - Name, contact details
  - Work experience: company names, job titles, dates, responsibilities
  - Education: degree, institution, year
  - Skills section: all listed skills with any proficiency indicators
  - Projects, certifications, awards
  - Any other visible sections

STEP 2 — ANALYZE: Compare the extracted resume against the job description provided.

Return ONLY a valid JSON object (no markdown, no prose before/after the JSON)."""


def mega_call(resume_text: str, jd_text: str,
              modules_hint: Optional[List[dict]] = None,
              resume_image_b64: Optional[str] = None) -> dict:
    reasoning_block = ""
    if modules_hint:
        reasoning_block = '\n  "reasoning": {"<module_id>": "<2-sentence why this candidate needs this module>"},'

    json_schema = f"""{{
  "candidate": {{
    "name": "<full name or Unknown>",
    "current_role": "<latest title>",
    "years_experience": <int>,
    "seniority": "<Junior|Mid|Senior|Lead>",
    "domain": "<Tech|Non-Tech|Hybrid>",
    "education": "<degree + field>",
    "skills": [{{"skill":"<n>","proficiency":<0-10>,"year_last_used":<year or 0>,"context":"<1-line evidence>"}}],
    "strengths": ["<s1>","<s2>","<s3>"],
    "red_flags": ["<f1>","<f2>"]
  }},
  "jd": {{
    "role_title": "<title>",
    "seniority_required": "<Junior|Mid|Senior|Lead>",
    "domain": "<Tech|Non-Tech|Hybrid>",
    "required_skills": ["<skill>"],
    "preferred_skills": ["<skill>"],
    "key_responsibilities": ["<resp>"]
  }},
  "audit": {{
    "ats_score": <0-100>,
    "completeness_score": <0-100>,
    "clarity_score": <0-100>,
    "overall_grade": "<A|B|C|D>",
    "ats_issues": ["<issue>"],
    "improvement_tips": ["<tip1>","<tip2>","<tip3>","<tip4>","<tip5>"],
    "missing_keywords": ["<kw>"],
    "interview_talking_points": ["<pt1>","<pt2>","<pt3>"]
  }}{reasoning_block}
}}"""

    if resume_image_b64:
        # FIX #1, #2, #3: Use vision model with detailed image extraction prompt
        prompt = f"""I have uploaded an IMAGE of my resume. Please:

1. Carefully OCR and read ALL text from the resume image
2. Extract every skill, job title, company, date, project, and education detail visible
3. Analyze the gap against this job description:

JOB DESCRIPTION:
{jd_text[:2000]}

Return EXACTLY this JSON schema (fill all fields based on what you can read from the image):
{json_schema}

Important: Extract skills from the resume image with realistic proficiency scores (1-10) based on how they are described (e.g., 'expert' = 9, 'proficient' = 7, 'familiar' = 4, 'basic' = 3). For year_last_used, use the most recent job/project year where that skill appeared."""

        return _groq_call_vision(
            prompt=prompt,
            system=_VISION_SYS,
            image_b64=resume_image_b64,
            max_tokens=3200
        )
    else:
        # Text resume — use standard fast model
        prompt = f"""Analyze this resume and job description.

RESUME (TEXT):
{resume_text[:4000]}

JOB DESCRIPTION:
{jd_text[:2000]}

Return EXACTLY this JSON:
{json_schema}"""
        return _groq_call(prompt=prompt, system=_MEGA_SYS, model=MODEL_FAST, max_tokens=2800)


def rewrite_resume(resume_text: str, jd: dict, missing_kw: List[str]) -> str:
    r = _groq_call(
        f'Rewrite this resume for the target role. Naturally add missing keywords: {missing_kw[:8]}. '
        f'Keep all facts true. Return JSON: {{"rewritten_resume":"<text>"}}\n\n'
        f'Resume:\n{resume_text[:1500]}\n\nTarget: {jd.get("role_title","--")}  '
        f'Required: {jd.get("required_skills",[])}',
        system="Expert resume writer. Return JSON only.", model=MODEL_FAST, max_tokens=1500,
    )
    return r.get("rewritten_resume","Could not rewrite resume.")

# =============================================================================
#  WEB SEARCH FEATURES
# =============================================================================
def search_real_salary(role: str, location: str) -> dict:
    results = ddg_search(f"{role} salary {location} 2025 average annual", max_results=6)
    if not results:
        return {}
    snippets = "\n".join([f"- {r.get('title','')}: {r.get('body','')[:200]}" for r in results[:5]])
    r = _groq_call(
        f'Extract salary data for "{role}" in {location} from these search snippets.\n\n{snippets}\n\n'
        f'Return JSON: {{"min_lpa":<number>,"max_lpa":<number>,"median_lpa":<number>,'
        f'"currency":"INR or USD","source":"<website name>","note":"<key caveat>"}}',
        system="Extract structured salary info from web snippets. Return JSON only.",
        model=MODEL_FAST, max_tokens=400,
    )
    median = r.get("median_lpa", 0)
    return r if ("error" not in r and isinstance(median, (int, float)) and median > 0) else {}

def _is_english(text: str) -> bool:
    """Return True if text is mostly ASCII/English."""
    if not text: return True
    ascii_chars = sum(1 for c in text if ord(c) < 128)
    return ascii_chars / max(len(text), 1) > 0.75

def search_course_links(skill: str) -> List[dict]:
    # Force English results by adding language hints to query
    results = ddg_search(
        f'{skill} online course english 2024 2025 site:coursera.org OR site:udemy.com OR site:youtube.com OR site:edx.org',
        max_results=8
    )
    courses = []
    for r in results:
        url   = r.get("href","")
        title = r.get("title","")
        body  = r.get("body","")
        if not url or not title: continue
        # Skip non-English titles
        if not _is_english(title): continue
        if   "coursera.org" in url: plat, icon = "Coursera", "🎓"
        elif "udemy.com"    in url: plat, icon = "Udemy",    "🎯"
        elif "youtube.com"  in url: plat, icon = "YouTube",  "▶"
        elif "edx.org"      in url: plat, icon = "edX",      "📘"
        elif "linkedin.com" in url: plat, icon = "LinkedIn", "💼"
        else: continue
        courses.append({"title":title[:65],"url":url,
                        "platform":plat,"icon":icon,"snippet":body[:120]})
    return courses[:4]

def search_skill_trends(skills: List[str]) -> Dict[str, str]:
    if not skills: return {}
    query = " ".join(skills[:6])
    results = ddg_search(f"most in-demand tech skills 2025 hiring india {query}", max_results=5)
    # Use only English results
    text = " ".join([r.get("body","") for r in results if _is_english(r.get("body",""))]).lower()
    out = {}
    for skill in skills:
        sl = skill.lower()
        count = text.count(sl)
        out[skill] = ("🔥 Hot" if count >= 3 else "📈 Growing" if count >= 1 else "✓ Stable")
    return out

def search_job_market(role: str) -> List[str]:
    results = ddg_search(f'{role} job market hiring demand 2025 india', max_results=5)
    if not results: return []
    # Filter to English snippets only
    eng_snippets = [r.get("body","")[:300] for r in results[:5] if _is_english(r.get("body",""))]
    if not eng_snippets: return []
    snippets = "\n".join(eng_snippets[:4])
    r = _groq_call(
        f'Based on these search results about "{role}" job market, give 3 short specific insights in English.\n\n'
        f'{snippets}\n\nReturn JSON: {{"insights":["<insight1>","<insight2>","<insight3>"]}}',
        system="Job market analyst. Give insights in English only. Return JSON only.", model=MODEL_FAST, max_tokens=300,
    )
    return r.get("insights",[]) if "error" not in r else []

# =============================================================================
#  CACHE
# =============================================================================
_CACHE_PATH = "/tmp/skillforge_v10"  # bumped version — forces fresh cache

def _ckey(r: str, j: str) -> str: return hashlib.md5((r+"||"+j).encode()).hexdigest()

def cache_get(r, j):
    try:
        with shelve.open(_CACHE_PATH) as db: return db.get(_ckey(r,j))
    except: return None

def cache_set(r, j, v):
    try:
        with shelve.open(_CACHE_PATH) as db: db[_ckey(r,j)] = v
    except: pass

# =============================================================================
#  ANALYSIS ENGINE
# =============================================================================
def _match_skill(skill: str) -> int:
    sl = (skill.lower().replace(".js","").replace(".ts","")
                       .replace("(","").replace(")","").strip())
    for i, cs in enumerate(CATALOG_SKILLS):
        if sl == cs or sl in cs or cs in sl: return i
    if SEMANTIC and _ST and _CEMBS is not None:
        try:
            from sklearn.metrics.pairwise import cosine_similarity
            import numpy as np
            sims = cosine_similarity(_ST.encode([sl]), _CEMBS)[0]
            best = int(np.argmax(sims))
            if sims[best] >= 0.52: return best
        except: pass
    tokens = set(sl.split()); best_s, best_i = 0.0, -1
    for i, cs in enumerate(CATALOG_SKILLS):
        ov = len(tokens & set(cs.split())) / max(len(tokens), 1)
        if ov > best_s: best_s, best_i = ov, i
    return best_i if best_s >= 0.4 else -1

def skill_decay(p, yr):
    p  = p  or 0   # guard None from LLM
    yr = yr or 0   # guard None from LLM
    if yr <= 0 or yr >= CURRENT_YEAR - 1: return p, False
    yrs = CURRENT_YEAR - yr
    if yrs <= 2: return p, False
    a = round(p * max(0.5, 1 - yrs / 5)); return a, a < p

def analyze_gap(candidate, jd):
    rs = {s["skill"].lower(): s for s in candidate.get("skills", [])}
    all_s = [(s, True) for s in jd.get("required_skills", [])] + \
            [(s, False) for s in jd.get("preferred_skills", [])]
    out = []
    for skill, req in all_s:
        sl = skill.lower().replace(".js","").replace(".ts","").strip()
        status, prof, ctx, dec, orig = "Missing", 0, "", False, 0
        src = rs.get(sl) or next((v for k,v in rs.items() if sl in k or k in sl), None)
        if src:
            raw_p = src.get("proficiency", 0) or 0; prof, dec = skill_decay(raw_p, src.get("year_last_used", 0) or 0)
            orig, ctx = raw_p, src.get("context",""); status = "Known" if prof >= 7 else "Partial"
        idx = _match_skill(skill)
        demand = MARKET_DEMAND.get(sl, MARKET_DEMAND.get(skill.lower(), 1))
        obs = OBSOLESCENCE_RISK.get(sl)
        out.append({"skill":skill,"status":status,"proficiency":prof,"original_prof":orig,
                    "decayed":dec,"is_required":req,"context":ctx,
                    "catalog_course":CATALOG[idx] if idx >= 0 else None,
                    "demand":demand,"obsolescence_risk":obs})
    return out

def seniority_check(c, jd):
    cs, rs = c.get("seniority","Mid"), jd.get("seniority_required","Mid")
    gap = SENIORITY_MAP.get(rs,1) - SENIORITY_MAP.get(cs,1)
    return {"has_mismatch":gap>0,"gap_levels":gap,"candidate":cs,"required":rs,
            "add_leadership":gap>=1,"add_strategic":gap>=2}

def build_path(gp, c, jd=None):
    needed = set(); id2gap = {}
    for g in gp:
        if g["status"] == "Known": continue
        co = g.get("catalog_course")
        if not co: continue
        needed.add(co["id"]); id2gap[co["id"]] = g
        try:
            for anc in nx.ancestors(SKILL_GRAPH, co["id"]):
                ad = CATALOG_BY_ID.get(anc)
                if ad and not any(x["status"]=="Known" and x["skill"].lower() in ad["skill"].lower()
                                  for x in gp):
                    needed.add(anc)
        except: pass
    if jd:
        sm = seniority_check(c, jd)
        if sm["add_leadership"]: needed.update(["LD01","LD02"])
        if sm["add_strategic"]:  needed.add("LD03")
    sub = SKILL_GRAPH.subgraph(needed)
    try: ordered = list(nx.topological_sort(sub))
    except: ordered = list(needed)
    crit = set()
    try:
        if sub.nodes: crit = set(nx.dag_longest_path(sub))
    except: pass
    path, seen = [], set()
    for cid in ordered:
        if cid in seen: continue
        seen.add(cid); co = CATALOG_BY_ID.get(cid)
        if not co: continue
        g = id2gap.get(cid, {})
        path.append({**co,
                     "gap_skill":  g.get("skill", co["skill"]),
                     "gap_status": g.get("status","Prereq"),
                     "priority":   (0 if g.get("is_required") else 1, g.get("proficiency",0) or 0),
                     "reasoning":  "",
                     "is_critical":cid in crit,
                     "demand":     g.get("demand", 1),
                     "is_required":g.get("is_required", False)})
    path.sort(key=lambda x: x["priority"]); return path

def calc_impact(gp, path):
    tot = len(gp); known = sum(1 for g in gp if g["status"]=="Known")
    covered = len({m["gap_skill"] for m in path}); rhrs = sum(m.get("duration_hrs",0) or 0 for m in path)
    cur  = min(100, round(known/max(tot,1)*100))
    proj = min(100, round((known+covered)/max(tot,1)*100))
    return {"total_skills":tot,"known_skills":known,"gaps_addressed":covered,
            "roadmap_hours":rhrs,"hours_saved":max(0,60-rhrs),
            "current_fit":cur,"projected_fit":proj,"fit_delta":proj-cur,
            "modules_count":len(path),"critical_count":sum(1 for m in path if m.get("is_critical")),
            "decayed_skills":sum(1 for g in gp if g.get("decayed"))}

def interview_readiness(gp, c):
    rk = [g for g in gp if g["status"]=="Known"   and g["is_required"]]
    rp = [g for g in gp if g["status"]=="Partial"  and g["is_required"]]
    rm = [g for g in gp if g["status"]=="Missing"  and g["is_required"]]
    tot = max(len(rk)+len(rp)+len(rm), 1)
    sc = max(0, min(100, round((len(rk)+len(rp)*0.4)/tot*100)
                    + {"Junior":5,"Mid":0,"Senior":-5,"Lead":-10}.get(c.get("seniority","Mid"),0)))
    if   sc >= 75: v = ("Strong",    "#4ade80", "Ready for most rounds")
    elif sc >= 50: v = ("Moderate",  "#fbbf24", "Pass screening; prep gaps")
    elif sc >= 30: v = ("Weak",      "#fb923c", "Gap work before applying")
    else:          v = ("Not Ready", "#f87171", "Significant prep needed")
    return {"score":sc,"label":v[0],"color":v[1],"advice":v[2],
            "req_known":len(rk),"req_partial":len(rp),"req_missing":len(rm)}

def weekly_plan(path, hpd=2.0):
    cap = max(hpd, 0.5) * 5  # FIX: guard against 0
    weeks, cur, hrs, wn = [], [], 0.0, 1
    for m in path:
        rem = float(m.get("duration_hrs",0) or 0)
        while rem > 0:
            avail = cap - hrs
            if avail <= 0:
                weeks.append({"week":wn,"modules":cur,"total_hrs":hrs})
                cur, hrs, wn = [], 0.0, wn+1; avail = cap
            chunk = min(rem, avail)
            ex = next((x for x in cur if x["id"]==m["id"]), None)
            if ex: ex["hrs_this_week"] += chunk
            else:  cur.append({"id":m["id"],"title":m["title"],"level":m["level"],
                                "domain":m["domain"],"is_critical":m.get("is_critical",False),
                                "hrs_this_week":chunk,"total_hrs":m.get("duration_hrs",0) or 0})
            hrs += chunk; rem -= chunk
    if cur: weeks.append({"week":wn,"modules":cur,"total_hrs":hrs})
    return weeks

def transfer_map_calc(c, gp):
    known = {g["skill"].lower() for g in c.get("skills",[]) if (g.get("proficiency",0) or 0) >= 6}
    out = []
    for g in gp:
        if g["status"] == "Known": continue
        sl = g["skill"].lower()
        for k in known:
            pct = TRANSFER_MAP.get(k,{}).get(sl, 0)
            if pct: out.append({"gap_skill":g["skill"],"known_skill":k.title(),
                                  "transfer_pct":pct,
                                  "label":f"Your {k.title()} → {pct}% head start on {g['skill']}"})
    return sorted(out, key=lambda x: x["transfer_pct"], reverse=True)

def roi_rank(gp, path):
    out = []
    for m in path:
        g = next((x for x in gp if x["skill"]==m.get("gap_skill")), {})
        roi = round((g.get("demand",1)*(1.5 if g.get("is_required") else 1)*10)/max(m.get("duration_hrs",1) or 1,1), 2)
        out.append({"id":m["id"],"title":m["title"],"skill":m["skill"],"roi":roi,
                    "hrs":m.get("duration_hrs",0) or 0,"is_required":g.get("is_required",False)})
    return sorted(out, key=lambda x: x["roi"], reverse=True)

def weeks_ready(hrs, hpd):
    hpd = max(hpd, 0.5)  # FIX: guard against 0
    if hrs <= 0: return "0d"
    w = (hrs/hpd)/5
    if w < 1:   return f"{int(hrs/hpd)}d"
    elif w < 4: return f"{w:.1f}w"
    return f"{(w/4):.1f}mo"

# =============================================================================
#  FULL PIPELINE
# =============================================================================
def run_analysis(resume_text, jd_text, resume_image_b64=None):
    # Cache key: always use md5 hash — avoids shelve key-length limits and
    # ensures different PDFs with same JD get separate cache entries
    if resume_text and resume_text.strip():
        cache_k = "txt:" + hashlib.md5(resume_text.encode()).hexdigest()
    elif resume_image_b64:
        cache_k = "img:" + hashlib.md5(resume_image_b64.encode()).hexdigest()
    else:
        cache_k = "empty"
    cached = cache_get(cache_k, jd_text)
    if cached: cached["_cache_hit"] = True; return cached
    kws = [w.strip() for w in jd_text.split() if len(w)>3][:20]
    potential_mods = [c for c in CATALOG
                      if any(kw.lower() in c["skill"].lower() or c["skill"].lower() in kw.lower()
                             for kw in kws)][:10]
    raw = mega_call(resume_text=resume_text, jd_text=jd_text,
                    modules_hint=potential_mods, resume_image_b64=resume_image_b64)
    if "error" in raw: return raw
    candidate = raw.get("candidate",{}); jd_data = raw.get("jd",{})
    quality = raw.get("audit",{}); rsn_map = raw.get("reasoning",{})
    if not candidate or not jd_data:
        return {"error":"parse_failed — empty candidate or JD"}
    gp   = analyze_gap(candidate, jd_data)
    path = build_path(gp, candidate, jd_data)
    for m in path: m["reasoning"] = rsn_map.get(m["id"], f"Addresses gap in {m['gap_skill']}.")
    im  = calc_impact(gp, path);      sm  = seniority_check(candidate, jd_data)
    iv  = interview_readiness(gp, candidate); wp = weekly_plan(path)
    tf  = transfer_map_calc(candidate, gp);   roi = roi_rank(gp, path)
    obs = [{"skill":g["skill"],"status":g["status"],"reason":OBSOLESCENCE_RISK[g["skill"].lower()]}
           for g in gp if OBSOLESCENCE_RISK.get(g["skill"].lower())]
    cgm = max(0, SENIORITY_MAP.get(jd_data.get("seniority_required","Mid"),1)
                 - SENIORITY_MAP.get(candidate.get("seniority","Mid"),1)) * 18
    result = {"candidate":candidate,"jd":jd_data,"gap_profile":gp,"path":path,
              "impact":im,"seniority":sm,"quality":quality,"interview":iv,
              "weekly_plan":wp,"transfers":tf,"roi":roi,"obsolescence":obs,
              "career_months":cgm,"_cache_hit":False}
    cache_set(cache_k, jd_text, result); return result

def run_analysis_with_web(resume_text, jd_text, resume_image_b64=None, location="India"):
    result = run_analysis(resume_text, jd_text, resume_image_b64)
    if "error" in result: return result
    role = result["jd"].get("role_title","")
    gap_skills = [g["skill"] for g in result["gap_profile"] if g["status"] != "Known"][:6]
    with ThreadPoolExecutor(max_workers=3) as ex:
        sal_f   = ex.submit(search_real_salary, role, location)
        trend_f = ex.submit(search_skill_trends, gap_skills)
        mkt_f   = ex.submit(search_job_market, role)
    result["salary"]          = sal_f.result()
    result["skill_trends"]    = trend_f.result()
    result["market_insights"] = mkt_f.result()
    return result

# =============================================================================
#  PDF EXPORT
# =============================================================================
def build_pdf(c, jd, gp, path, im, ql=None, iv=None):
    buf = io.BytesIO()
    if not REPORTLAB: return buf
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            topMargin=48, bottomMargin=48, leftMargin=48, rightMargin=48)
    styles = getSampleStyleSheet()
    TEAL = rl_colors.HexColor("#2dd4bf")
    BD = ParagraphStyle("BD", parent=styles["Normal"], fontSize=10, spaceAfter=5)
    H1 = ParagraphStyle("H1", parent=styles["Title"],   fontSize=20, spaceAfter=4, textColor=TEAL)
    H2 = ParagraphStyle("H2", parent=styles["Heading2"],fontSize=13, spaceAfter=6, spaceBefore=14)
    IT = ParagraphStyle("IT", parent=styles["Normal"],  fontSize=9,  spaceAfter=4, leftIndent=18,
                        textColor=rl_colors.HexColor("#555"))
    story = [
        Paragraph("SkillForge — AI Adaptive Onboarding Report", H1),
        Paragraph(f"Candidate: <b>{c.get('name','--')}</b>  |  Role: <b>{jd.get('role_title','--')}</b>  "
                  f"|  Generated: {datetime.now().strftime('%d %b %Y %H:%M')}", BD),
        Spacer(1, 14),
    ]
    if ql or iv:
        story.append(Paragraph("Scores", H2))
        rows = []
        if ql: rows += [["ATS Score", f"{ql.get('ats_score','--')}%"],
                        ["Grade", ql.get("overall_grade","--")],
                        ["Completeness", f"{ql.get('completeness_score','--')}%"],
                        ["Clarity", f"{ql.get('clarity_score','--')}%"]]
        if iv: rows += [["Interview Ready", f"{iv['score']}% — {iv['label']}"],
                        ["Known", str(iv["req_known"])],
                        ["Missing", str(iv["req_missing"])]]
        t = Table([["Metric","Value"]]+rows, colWidths=[200,260])
        t.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0),TEAL),
            ("TEXTCOLOR",(0,0),(-1,0),rl_colors.white),
            ("FONTSIZE",(0,0),(-1,-1),10),
            ("GRID",(0,0),(-1,-1),0.4,rl_colors.grey),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[rl_colors.whitesmoke,rl_colors.white]),
            ("LEFTPADDING",(0,0),(-1,-1),8),
        ]))
        story += [t, Spacer(1,14)]
    story.append(Paragraph("Learning Roadmap", H2))
    for i, m in enumerate(path):
        story.append(Paragraph(
            f"<b>{i+1}. {'[CRITICAL] ' if m.get('is_critical') else ''}{m['title']}</b>"
            f" — {m['level']} / {m['duration_hrs']}h", BD))
        if m.get("reasoning"): story.append(Paragraph(f"→ {m['reasoning']}", IT))
    doc.build(story); buf.seek(0); return buf


# =============================================================================
#  CHARTS
# =============================================================================
_BG     = "rgba(0,0,0,0)"
_GRID   = "rgba(255,255,255,0.06)"
_TEAL   = "#2dd4bf"
_AMBER  = "#f59e0b"
_RED    = "#ef4444"
_GREEN  = "#4ade80"
_FONT   = dict(color="#94a3b8", family="'DM Mono', 'IBM Plex Mono', monospace")

def _bl(**kw):
    return dict(paper_bgcolor=_BG, plot_bgcolor=_BG, font=_FONT,
                margin=dict(l=8,r=8,t=8,b=36), **kw)

def radar_chart(gp):
    items = [g for g in gp if g["is_required"]][:10] or gp[:10]
    if not items: return go.Figure()
    theta = [g["skill"][:14] for g in items]
    fig = go.Figure(data=[
        go.Scatterpolar(r=[10]*len(items), theta=theta, fill="toself",
                        name="Required", line=dict(color=_RED,width=1), opacity=0.08),
        go.Scatterpolar(r=[g.get("proficiency") or 0 for g in items], theta=theta, fill="toself",
                        name="Current", line=dict(color=_TEAL,width=2.5), opacity=0.7),
    ])
    fig.update_layout(**_bl(height=320),
        polar=dict(bgcolor=_BG,
                   radialaxis=dict(visible=True, range=[0,10], gridcolor=_GRID,
                                   tickfont=dict(size=9,color="#475569")),
                   angularaxis=dict(gridcolor=_GRID, tickfont=dict(size=11))),
        showlegend=True,
        legend=dict(bgcolor=_BG,x=0.72,y=1.22,font=dict(size=10)),
    )
    return fig

def timeline_chart(path):
    if not path: return go.Figure()
    lc = {"Critical":_RED,"Beginner":_TEAL,"Intermediate":_AMBER,"Advanced":"#f97316"}
    shown, fig = set(), go.Figure()
    for m in path:
        k = "Critical" if m.get("is_critical") else m["level"]
        show = k not in shown; shown.add(k)
        fig.add_trace(go.Bar(
            x=[m.get("duration_hrs",0) or 0], y=[m["title"][:30]],
            orientation="h",
            marker=dict(color=lc.get(k,"#64748b"), opacity=0.85, line=dict(width=0)),
            name=k, legendgroup=k, showlegend=show,
            hovertemplate=f"<b>{m['title']}</b><br>{m['level']} · {m['duration_hrs']}h<extra></extra>",
        ))
    fig.update_layout(**_bl(height=max(260, len(path)*38)),
        xaxis=dict(title="Hours", gridcolor=_GRID, zeroline=False, tickfont=dict(size=11)),
        yaxis=dict(gridcolor="rgba(0,0,0,0)", tickfont=dict(size=11), autorange="reversed"),
        legend=dict(bgcolor=_BG, orientation="h", y=1.05, font=dict(size=11)),
        barmode="overlay",
    )
    return fig

def salary_chart(s):
    if not s: return go.Figure()
    try: _med = float(s.get("median_lpa") or 0)
    except: _med = 0
    if _med <= 0: return go.Figure()
    def _to_num(v): 
        try: return float(v or 0)
        except: return 0.0
    vals  = [_to_num(s.get("min_lpa",0)), _to_num(s.get("median_lpa",0)), _to_num(s.get("max_lpa",0))]
    curr  = s.get("currency","INR")
    sym   = "₹" if curr=="INR" else "$"
    unit  = "L/yr" if curr=="INR" else "k/yr"
    lbls  = [f"{sym}{v}{unit}" for v in vals]
    fig = go.Figure(go.Bar(x=["Min","Median","Max"], y=vals,
        marker_color=[_TEAL,_AMBER,_RED], opacity=0.82,
        text=lbls, textposition="outside",
        textfont=dict(size=13,family="'DM Mono','IBM Plex Mono',monospace"),
    ))
    fig.update_layout(**_bl(height=230),
        yaxis=dict(title=unit, gridcolor=_GRID, zeroline=False, tickfont=dict(size=11)),
        xaxis=dict(gridcolor="rgba(0,0,0,0)", tickfont=dict(size=12)),
    )
    return fig

def roi_bar(roi_list):
    if not roi_list: return go.Figure()
    top = roi_list[:10]
    fig = go.Figure(go.Bar(
        x=[m["roi"] for m in top],
        y=[m["title"][:28] for m in top],
        orientation="h",
        marker=dict(color=[_RED if m["is_required"] else _TEAL for m in top], opacity=0.85, line=dict(width=0)),
        hovertemplate="<b>%{y}</b><br>ROI Index: %{x}<extra></extra>",
    ))
    fig.update_layout(**_bl(height=max(200, len(top)*36)),
        xaxis=dict(title="ROI Index", gridcolor=_GRID, zeroline=False, tickfont=dict(size=11)),
        yaxis=dict(gridcolor="rgba(0,0,0,0)", autorange="reversed", tickfont=dict(size=11)),
    )
    return fig


# =============================================================================
#  CSS
# =============================================================================
CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&family=DM+Mono:wght@400;500&display=swap');
:root{
  --bg:      #0b0d14;
  --s1:      #131720;
  --s2:      #1a1f2e;
  --s3:      #222840;
  --border:  rgba(255,255,255,0.07);
  --bhi:     rgba(45,212,191,0.20);
  --teal:    #2dd4bf;
  --teal-bg: rgba(45,212,191,0.08);
  --amber:   #f59e0b;
  --red:     #ef4444;
  --green:   #4ade80;
  --purple:  #a78bfa;
  --t1:      #f1f5f9;
  --t2:      #94a3b8;
  --t3:      #475569;
  --t4:      #2d3a52;
  --sans:    'DM Sans',sans-serif;
  --mono:    'DM Mono','IBM Plex Mono',monospace;
}
*,*::before,*::after{box-sizing:border-box}
html,body,[class*="css"]{
  font-family:var(--sans)!important;
  background:var(--bg)!important;
  color:var(--t2)!important;
  font-size:15px!important;
}
.stApp{background:var(--bg)!important}
.main .block-container{padding:0!important;max-width:100%!important}
footer,
#MainMenu,
header[data-testid="stHeader"],
[data-testid="stToolbar"],
[data-testid="stDecoration"],
[data-testid="stStatusWidget"],
[data-testid="stAppDeployButton"],
[data-testid="stBaseButton-header"],
[data-testid="stBaseButton-minimal"],
[data-testid="stBaseButton-borderless"],
.stDeployButton,
button[kind="header"],
button[kind="minimal"],
[data-testid="stHeader"] button,
[data-testid="stHeader"] a,
[data-testid="stHeader"] > div,
#stDecoration
{display:none!important;visibility:hidden!important;height:0!important;
 width:0!important;overflow:hidden!important;pointer-events:none!important;
 position:absolute!important;opacity:0!important}
[data-testid="stHeader"]{height:0!important;min-height:0!important;padding:0!important}
section[data-testid="stSidebar"]>div:first-child{
  background:var(--s1)!important;border-right:1px solid var(--border)!important;
}
::-webkit-scrollbar{width:3px;height:3px}
::-webkit-scrollbar-thumb{background:rgba(255,255,255,0.1);border-radius:99px}

/* TOPBAR */
.sf-top{
  height:52px;display:flex;align-items:center;justify-content:space-between;
  padding:0 32px;border-bottom:1px solid var(--border);
  position:sticky;top:0;z-index:200;
  background:rgba(11,13,20,0.96);backdrop-filter:blur(20px);
}
.sf-logo{font-family:var(--sans);font-size:1.1rem;font-weight:700;color:var(--t1);letter-spacing:-0.02em}
.sf-logo em{color:var(--teal);font-style:normal}
.sf-top-right{display:flex;align-items:center;gap:8px;font-family:var(--mono);font-size:0.65rem;color:var(--t3)}
.sf-chip{padding:3px 10px;border-radius:4px;border:1px solid var(--border);color:var(--t3);font-size:0.63rem}
.sf-chip.on{border-color:var(--bhi);color:var(--teal)}

/* PAGE WRAP */
.sf-page{padding:0 32px 80px;max-width:1200px;margin:0 auto}
@media(max-width:640px){.sf-page{padding:0 16px 80px}}

/* ── INPUT PAGE ─────────────────────────────────── */
.sf-hero{padding:44px 0 32px}
.sf-eyebrow{
  font-family:var(--mono);font-size:0.68rem;font-weight:500;
  letter-spacing:0.12em;text-transform:uppercase;color:var(--teal);
  display:flex;align-items:center;gap:10px;margin-bottom:12px;
}
.sf-eyebrow::before{content:'';width:28px;height:1px;background:var(--teal)}
.sf-h1{
  font-family:var(--sans);font-size:clamp(2rem,4vw,3.2rem);
  font-weight:700;color:var(--t1);line-height:1.1;
  letter-spacing:-0.03em;margin-bottom:14px;
}
.sf-h1 span{color:var(--teal)}
.sf-sub{font-size:1rem;color:var(--t2);line-height:1.6;max-width:480px;margin-bottom:0}

/* sample pills */
.sf-samples{display:flex;align-items:center;gap:8px;flex-wrap:wrap;margin-bottom:28px}
.sf-sample-lbl{font-family:var(--mono);font-size:0.65rem;color:var(--t3)}

/* input grid */
.sf-input-grid{display:grid;grid-template-columns:1fr 1fr;gap:16px;margin-bottom:20px}
@media(max-width:700px){.sf-input-grid{grid-template-columns:1fr}}
.sf-panel{
  background:var(--s1);border:1px solid var(--border);border-radius:12px;padding:20px 22px;
}
.sf-panel.ready{border-color:var(--bhi)}
.sf-panel-hd{
  font-size:0.82rem;font-weight:600;color:var(--t1);
  margin-bottom:14px;display:flex;align-items:center;gap:8px;
}
.sf-panel-icon{font-size:1rem}
.sf-ready-badge{
  font-family:var(--mono);font-size:0.6rem;padding:2px 8px;border-radius:3px;
  background:rgba(45,212,191,0.1);color:var(--teal);border:1px solid var(--bhi);
  margin-left:auto;
}
.sf-wc{font-family:var(--mono);font-size:0.65rem;color:var(--t3);margin-top:8px}

/* image upload hint */
.sf-img-hint{
  background:rgba(45,212,191,0.04);border:1px solid var(--bhi);border-radius:7px;
  padding:9px 13px;margin-top:8px;font-family:var(--mono);font-size:0.68rem;color:var(--teal);
}

/* file uploader */
[data-testid="stFileUploadDropzone"]{
  background:rgba(45,212,191,0.02)!important;
  border:1.5px dashed rgba(45,212,191,0.14)!important;border-radius:8px!important;
}
[data-testid="stFileUploadDropzone"]:hover{
  border-color:rgba(45,212,191,0.32)!important;background:rgba(45,212,191,0.04)!important;
}
[data-testid="stFileUploadDropzone"] button{
  background:transparent!important;border:1px solid var(--bhi)!important;
  color:var(--teal)!important;font-family:var(--mono)!important;font-size:0.72rem!important;
  border-radius:5px!important;
}

/* textarea */
textarea{
  background:#0c0f1a!important;border:1px solid var(--border)!important;
  border-radius:8px!important;color:#b8ccd8!important;
  font-family:var(--mono)!important;font-size:0.82rem!important;
  resize:vertical!important;line-height:1.6!important;
}
textarea:focus{border-color:var(--bhi)!important;outline:none!important}
textarea::placeholder{color:var(--t4)!important}

/* options row */
.sf-opts{display:flex;align-items:center;gap:16px;flex-wrap:wrap;margin-bottom:20px}

/* BUTTONS */
.stButton>button{
  background:var(--teal)!important;border:none!important;border-radius:8px!important;
  color:#061412!important;font-family:var(--sans)!important;font-weight:700!important;
  font-size:0.9rem!important;padding:11px 0!important;width:100%!important;
  letter-spacing:0.01em!important;transition:opacity 0.15s!important;
}
.stButton>button:hover{opacity:0.84!important}
.stButton>button:disabled{opacity:0.3!important}
.sf-ghost .stButton>button{
  background:var(--s2)!important;border:1px solid var(--border)!important;
  color:var(--t2)!important;font-weight:500!important;
}
.sf-ghost .stButton>button:hover{border-color:rgba(255,255,255,0.15)!important;color:var(--t1)!important}
.sf-danger .stButton>button{background:rgba(239,68,68,0.12)!important;border:1px solid rgba(239,68,68,0.2)!important;color:var(--red)!important}

/* ── SCORE BANNER ───────────────────────────────── */
.sf-banner{
  background:var(--s1);border:1px solid var(--border);border-radius:14px;
  padding:28px 32px;margin:24px 0 20px;
}
.sf-banner-top{
  display:flex;align-items:flex-start;gap:8px;margin-bottom:6px;
}
.sf-candidate-name{font-size:1.2rem;font-weight:700;color:var(--t1);letter-spacing:-0.02em}
.sf-candidate-sub{font-family:var(--mono);font-size:0.72rem;color:var(--t3);margin-top:2px}
.sf-cache-badge{
  font-family:var(--mono);font-size:0.6rem;padding:2px 8px;border-radius:3px;
  background:rgba(167,139,250,0.1);color:var(--purple);border:1px solid rgba(167,139,250,0.2);
  margin-left:auto;margin-top:4px;
}
.sf-vision-badge{
  font-family:var(--mono);font-size:0.6rem;padding:2px 8px;border-radius:3px;
  background:rgba(45,212,191,0.1);color:var(--teal);border:1px solid var(--bhi);
  margin-left:8px;margin-top:4px;
}

/* score circles row */
.sf-scores{display:grid;grid-template-columns:repeat(3,1fr);gap:14px;margin-top:20px}
@media(max-width:700px){.sf-scores{grid-template-columns:1fr}}
.sf-score-card{
  background:var(--s2);border:1px solid var(--border);border-radius:10px;
  padding:20px 22px;text-align:center;
}
.sf-score-num{
  font-family:var(--mono);font-size:3.2rem;font-weight:500;line-height:1;
  letter-spacing:-0.04em;
}
.sf-score-lbl{
  font-family:var(--mono);font-size:0.62rem;font-weight:500;
  letter-spacing:0.1em;text-transform:uppercase;color:var(--t3);margin-top:6px;
}
.sf-score-sub{font-size:0.78rem;color:var(--t2);margin-top:5px;line-height:1.4}

/* kpi row */
.sf-kpis{display:grid;grid-template-columns:repeat(4,1fr);gap:10px;margin-top:16px}
@media(max-width:600px){.sf-kpis{grid-template-columns:repeat(2,1fr)}}
.sf-kpi{
  background:var(--s2);border:1px solid var(--border);border-radius:8px;padding:14px 16px;
}
.sf-kpi-v{font-family:var(--mono);font-size:1.45rem;font-weight:500;color:var(--t1);line-height:1}
.sf-kpi-l{font-family:var(--mono);font-size:0.6rem;text-transform:uppercase;letter-spacing:0.08em;color:var(--t3);margin-top:4px}
.sf-kpi-d{font-family:var(--mono);font-size:0.7rem;color:var(--teal);margin-top:3px}

/* ── TABS ───────────────────────────────────────── */
[data-testid="stTabs"] button{
  font-family:var(--sans)!important;font-size:0.9rem!important;font-weight:500!important;
  padding:10px 18px!important;
}
[data-testid="stTabs"] button[aria-selected="true"]{color:var(--teal)!important}

/* ── SKILL GRID ─────────────────────────────────── */
.sf-skill-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(210px,1fr));gap:10px;margin-bottom:8px}
.sf-skill-card{
  background:var(--s1);border:1px solid var(--border);border-radius:9px;
  padding:14px 16px;transition:border-color 0.12s;cursor:default;
}
.sf-skill-card:hover{border-color:rgba(255,255,255,0.12)}
.sf-skill-card.known{border-left:3px solid var(--teal)}
.sf-skill-card.partial{border-left:3px solid var(--amber)}
.sf-skill-card.missing{border-left:3px solid var(--red)}
.sf-skill-top{display:flex;align-items:center;justify-content:space-between;gap:8px;margin-bottom:8px}
.sf-skill-name{font-size:0.88rem;font-weight:600;color:var(--t1)}
.sf-st-badge{
  font-family:var(--mono);font-size:0.6rem;font-weight:600;padding:2px 8px;border-radius:3px;
}
.sf-st-known{background:rgba(45,212,191,0.1);color:var(--teal);border:1px solid rgba(45,212,191,0.2)}
.sf-st-partial{background:rgba(245,158,11,0.1);color:var(--amber);border:1px solid rgba(245,158,11,0.2)}
.sf-st-missing{background:rgba(239,68,68,0.1);color:var(--red);border:1px solid rgba(239,68,68,0.2)}
.sf-skill-bar{height:4px;background:rgba(255,255,255,0.05);border-radius:99px;margin-bottom:6px}
.sf-skill-bar-fill{height:100%;border-radius:99px}
.sf-skill-bottom{display:flex;align-items:center;justify-content:space-between}
.sf-skill-score{font-family:var(--mono);font-size:0.75rem;color:var(--t2)}
.sf-skill-demand{font-family:var(--mono);font-size:0.65rem}
.sf-decay-tag{
  font-family:var(--mono);font-size:0.6rem;color:var(--amber);
  background:rgba(245,158,11,0.08);border:1px solid rgba(245,158,11,0.2);
  border-radius:3px;padding:1px 5px;margin-top:5px;display:inline-block;
}
.sf-skill-ctx{font-size:0.72rem;color:var(--t3);margin-top:6px;font-style:italic;line-height:1.4}

/* ── MODULE CARDS ───────────────────────────────── */
.sf-phase-hd{
  font-size:0.72rem;font-weight:600;letter-spacing:0.1em;text-transform:uppercase;
  color:var(--t3);margin:22px 0 10px;display:flex;align-items:center;gap:10px;
}
.sf-phase-hd::after{content:'';flex:1;height:1px;background:var(--border)}
.sf-mod{
  background:var(--s1);border:1px solid var(--border);
  border-left:3px solid transparent;border-radius:0 9px 9px 0;
  padding:14px 16px 14px 14px;margin-bottom:8px;
  transition:border-color 0.12s;
}
.sf-mod:hover{border-right-color:rgba(255,255,255,0.1)!important}
.sf-mod.crit{border-left-color:var(--red)!important}
.sf-mod.adv{border-left-color:#f97316}
.sf-mod.inter{border-left-color:var(--amber)}
.sf-mod.beg{border-left-color:var(--teal)}
.sf-mod.done{opacity:0.45}
.sf-mod-row{display:flex;align-items:flex-start;gap:12px}
.sf-mod-num{
  font-family:var(--mono);font-size:0.68rem;color:var(--t4);
  min-width:28px;padding-top:1px;flex-shrink:0;
}
.sf-mod-body{flex:1;min-width:0}
.sf-mod-title{font-size:0.9rem;font-weight:600;color:var(--t1);margin-bottom:3px;line-height:1.3}
.sf-mod-meta{font-family:var(--mono);font-size:0.68rem;color:var(--t3);margin-bottom:6px}
.sf-mod-tags{display:flex;gap:5px;flex-wrap:wrap}
.sf-tag{
  font-family:var(--mono);font-size:0.6rem;padding:2px 8px;border-radius:3px;
  background:var(--s2);color:var(--t2);border:1px solid var(--border);
}
.sf-tag-crit{color:var(--red);border-color:rgba(239,68,68,0.25);background:rgba(239,68,68,0.06)}
.sf-tag-req{color:var(--teal);border-color:var(--bhi);background:var(--teal-bg)}
.sf-mod-reason{
  font-size:0.76rem;color:var(--t3);font-style:italic;line-height:1.55;
  margin-top:9px;padding-top:9px;border-top:1px solid var(--border);
}
.sf-mod-hrs{
  font-family:var(--mono);font-size:0.8rem;color:var(--t2);
  white-space:nowrap;flex-shrink:0;padding-top:1px;
}
.sf-course-link{
  background:rgba(45,212,191,0.04);border:1px solid var(--bhi);border-radius:6px;
  padding:7px 12px;margin-top:9px;
}
.sf-course-link a{font-size:0.78rem;color:var(--teal);text-decoration:none;font-weight:500}
.sf-course-link a:hover{text-decoration:underline}
.sf-course-plat{font-family:var(--mono);font-size:0.64rem;color:var(--t3);margin-top:2px}

/* ── ATS SCORES ─────────────────────────────────── */
.sf-ats-row{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin-bottom:18px}
@media(max-width:600px){.sf-ats-row{grid-template-columns:repeat(2,1fr)}}
.sf-ats-card{
  background:var(--s1);border:1px solid var(--border);border-radius:9px;
  padding:18px 16px;text-align:center;
}
.sf-ats-n{font-family:var(--mono);font-size:1.8rem;font-weight:500;color:var(--t1);line-height:1}
.sf-ats-l{font-family:var(--mono);font-size:0.62rem;text-transform:uppercase;letter-spacing:0.08em;color:var(--t3);margin-top:5px}
.sf-prog{height:3px;background:rgba(255,255,255,0.05);border-radius:99px;overflow:hidden;margin-bottom:20px}
.sf-prog-fill{height:100%;border-radius:99px;background:var(--teal)}

/* tips / talking pts */
.sf-tip{display:flex;gap:12px;margin-bottom:10px;font-size:0.82rem;color:var(--t2);line-height:1.6}
.sf-tip-n{
  font-family:var(--mono);font-size:0.62rem;color:var(--teal);
  background:var(--teal-bg);border:1px solid var(--bhi);border-radius:3px;
  padding:2px 7px;font-weight:500;min-width:26px;text-align:center;flex-shrink:0;height:fit-content;
}
.sf-talk{
  font-size:0.82rem;color:var(--t2);padding:8px 0 8px 14px;
  border-left:2px solid var(--teal);margin-bottom:7px;line-height:1.55;
}
.sf-kw{
  display:inline-block;font-family:var(--mono);font-size:0.65rem;
  padding:3px 9px;border-radius:3px;margin:3px;
  background:rgba(239,68,68,0.07);color:var(--red);border:1px solid rgba(239,68,68,0.18);
}

/* ── RESEARCH TAB ───────────────────────────────── */
.sf-search-result{
  background:var(--s1);border:1px solid var(--border);border-radius:9px;
  padding:14px 16px;margin-bottom:8px;
}
.sf-search-title{font-size:0.9rem;font-weight:600;color:var(--teal);text-decoration:none}
.sf-search-title:hover{text-decoration:underline}
.sf-search-url{font-family:var(--mono);font-size:0.65rem;color:var(--t4);margin:3px 0 5px}
.sf-search-body{font-size:0.78rem;color:var(--t2);line-height:1.55}
.sf-insight{
  background:rgba(45,212,191,0.03);border-left:2px solid var(--teal);
  border-radius:0 5px 5px 0;padding:9px 13px;margin-bottom:6px;
  font-size:0.82rem;color:var(--t2);line-height:1.55;
}
.sf-trend-pill{
  display:inline-flex;align-items:center;gap:6px;
  background:var(--s2);border:1px solid var(--border);border-radius:6px;
  padding:6px 12px;margin:4px;font-size:0.78rem;
}

/* ── TRANSFER ────────────────────────────────────── */
.sf-xfer{
  background:var(--s2);border:1px solid var(--border);border-radius:7px;
  padding:10px 14px;margin-bottom:6px;
  display:flex;align-items:center;gap:10px;font-size:0.78rem;color:var(--t2);
}
.sf-xfer-pct{color:var(--purple);font-family:var(--mono);font-weight:500;font-size:0.88rem}

/* ── EXPORT CARDS ───────────────────────────────── */
.sf-export-card{
  background:var(--s1);border:1px solid var(--border);border-radius:10px;padding:20px 22px;
}
.sf-export-hd{font-size:0.9rem;font-weight:700;color:var(--t1);margin-bottom:4px}
.sf-export-sub{font-size:0.78rem;color:var(--t3);margin-bottom:14px;line-height:1.5}
.sf-export-row{display:flex;justify-content:space-between;font-family:var(--mono);font-size:0.72rem;padding:5px 0;border-bottom:1px solid var(--border)}
.sf-ek{color:var(--t3)}.sf-ev{color:var(--t1);font-weight:500}

/* download button */
[data-testid="stDownloadButton"]>button{
  background:var(--s2)!important;border:1px solid var(--border)!important;
  color:var(--t2)!important;font-family:var(--sans)!important;font-weight:500!important;
  font-size:0.82rem!important;
}
[data-testid="stDownloadButton"]>button:hover{border-color:var(--bhi)!important;color:var(--teal)!important}

/* Streamlit metric override */
[data-testid="stMetric"]{background:var(--s2)!important;border:1px solid var(--border)!important;border-radius:8px!important;padding:13px 15px!important}
[data-testid="stMetricValue"]{font-family:var(--mono)!important;font-size:1.5rem!important;color:var(--t1)!important}
[data-testid="stMetricLabel"]{font-family:var(--mono)!important;color:var(--t3)!important;font-size:0.6rem!important;text-transform:uppercase!important;letter-spacing:0.08em!important}

/* Selectbox */
[data-testid="stSelectbox"]>div>div{background:var(--s1)!important;border:1px solid var(--border)!important;color:var(--t1)!important;font-family:var(--sans)!important;font-size:0.85rem!important}

/* Expander */
[data-testid="stExpander"]{background:var(--s1)!important;border:1px solid var(--border)!important;border-radius:8px!important;margin-bottom:5px!important}
[data-testid="stExpander"] summary{font-family:var(--sans)!important;color:var(--t2)!important;font-size:0.85rem!important}

/* Checkbox */
[data-testid="stCheckbox"] label{font-family:var(--sans)!important;font-size:0.85rem!important;color:var(--t2)!important}

/* Select slider */
[data-testid="stSlider"] .st-ae{background:var(--teal)!important}

/* Progress */
[data-testid="stProgressBar"]>div>div{background:var(--teal)!important}
[data-testid="stProgressBar"]>div{background:rgba(255,255,255,0.05)!important;border-radius:99px!important}

/* Diff pane */
.sf-diff{
  background:#090c16;border:1px solid var(--border);border-radius:8px;
  padding:14px 16px;font-family:var(--mono);font-size:0.78rem;color:var(--t2);
  white-space:pre-wrap;line-height:1.6;max-height:320px;overflow-y:auto;
}

/* API log */
.sf-log{font-family:var(--mono);font-size:0.68rem;color:var(--t3);padding:5px 10px;background:var(--s1);border:1px solid var(--border);border-radius:4px;margin-bottom:3px;display:flex;gap:12px}

/* Sidebar nav */
.sf-nav-item{display:flex;align-items:center;gap:9px;padding:9px 12px;border-radius:6px;font-size:0.82rem;color:var(--t2);text-decoration:none;transition:all 0.12s;margin-bottom:2px}
.sf-nav-item:hover{background:var(--s2);color:var(--t1)}
.sf-nav-dot{width:6px;height:6px;border-radius:50%;flex-shrink:0}

/* Seniority warning */
.sf-warn{background:rgba(245,158,11,0.07);border:1px solid rgba(245,158,11,0.2);border-radius:8px;padding:12px 16px;font-size:0.82rem;color:var(--amber);margin-bottom:12px}

/* Footer */
.sf-foot{
  position:fixed;bottom:0;left:0;right:0;background:rgba(11,13,20,0.97);
  border-top:1px solid var(--border);padding:6px 32px;
  font-family:var(--mono);font-size:0.62rem;color:var(--t3);
  display:flex;align-items:center;gap:16px;z-index:99;
}
.sf-fdot{width:5px;height:5px;border-radius:50%;display:inline-block;margin-right:4px}
.sf-fr{margin-left:auto}

/* Section headings */
.sf-sh{font-size:1.15rem;font-weight:700;color:var(--t1);letter-spacing:-0.02em;margin-bottom:4px}
.sf-ss{font-family:var(--mono);font-size:0.68rem;color:var(--t3);margin-bottom:18px}
.sf-divider{height:1px;background:var(--border);margin:28px 0}
/* ── HOW IT WORKS STRIP ─────────────────────────────── */
.sf-how{display:flex;align-items:center;gap:0;margin:0 0 24px;background:var(--s1);border:1px solid var(--border);border-radius:12px;padding:24px 28px}
.sf-how-step{flex:1;text-align:center}
.sf-how-num{font-family:var(--mono);font-size:1.6rem;font-weight:500;color:var(--teal);line-height:1;margin-bottom:6px}
.sf-how-title{font-size:0.9rem;font-weight:700;color:var(--t1);margin-bottom:4px}
.sf-how-sub{font-family:var(--mono);font-size:0.65rem;color:var(--t3);line-height:1.5}
.sf-how-arrow{font-size:1.4rem;color:var(--border);padding:0 16px;flex-shrink:0}

/* ── STATS STRIP ─────────────────────────────────────── */
.sf-stats-strip{display:flex;align-items:center;justify-content:center;gap:0;margin:0 0 24px;background:rgba(45,212,191,0.04);border:1px solid var(--bhi);border-radius:10px;padding:18px 28px}
.sf-stat{text-align:center;flex:1}
.sf-stat-n{display:block;font-family:var(--mono);font-size:1.8rem;font-weight:500;color:var(--teal);line-height:1;margin-bottom:4px}
.sf-stat-l{font-family:var(--mono);font-size:0.62rem;text-transform:uppercase;letter-spacing:0.08em;color:var(--t3)}
.sf-stat-div{width:1px;height:40px;background:var(--border);margin:0 16px;flex-shrink:0}

/* ── HERO DELTA ──────────────────────────────────────── */
.sf-hero-delta{text-align:center;padding:28px 0 20px;border-bottom:1px solid var(--border);margin-bottom:20px}
.sf-hero-delta-num{font-family:var(--mono);font-size:4rem;font-weight:500;color:var(--green);line-height:1;letter-spacing:-0.04em}
.sf-hero-delta-label{font-family:var(--mono);font-size:0.7rem;text-transform:uppercase;letter-spacing:0.12em;color:var(--t3);margin:8px 0 6px}
.sf-hero-delta-sub{font-size:0.9rem;color:var(--t2)}

/* ── GROUNDING BADGE ─────────────────────────────────── */
.sf-ground-badge{display:flex;align-items:center;gap:10px;margin-top:16px;padding:10px 16px;background:rgba(45,212,191,0.04);border:1px solid var(--bhi);border-radius:7px;font-family:var(--mono);font-size:0.68rem;color:var(--teal)}
.sf-ground-dot{width:7px;height:7px;border-radius:50%;background:var(--teal);flex-shrink:0;animation:pulse 2s infinite}
@keyframes pulse{0%,100%{opacity:1}50%{opacity:0.4}}

/* ── LOADING STEPS ───────────────────────────────────── */
.sf-lstep{display:flex;align-items:flex-start;gap:14px;padding:12px 16px;border-radius:8px;margin-bottom:6px;font-size:0.85rem}
.sf-lstep-done{background:rgba(74,222,128,0.05);border:1px solid rgba(74,222,128,0.15)}
.sf-lstep-active{background:rgba(45,212,191,0.07);border:1px solid var(--bhi)}
.sf-lstep-wait{background:var(--s1);border:1px solid var(--border);opacity:0.45}
.sf-lstep-icon{font-family:var(--mono);font-size:1rem;min-width:24px;text-align:center;margin-top:1px}
.sf-lstep-spin{display:inline-block;animation:spin 1.2s linear infinite}
@keyframes spin{from{transform:rotate(0deg)}to{transform:rotate(360deg)}}
.sf-lstep-title{font-weight:600;color:var(--t1);margin-bottom:2px}
.sf-lstep-sub{font-family:var(--mono);font-size:0.65rem;color:var(--t3)}
.sf-lprog{height:4px;background:rgba(255,255,255,0.05);border-radius:99px;overflow:hidden;margin:16px 0 24px;max-width:560px}
.sf-lprog-fill{height:100%;border-radius:99px;background:linear-gradient(90deg,var(--teal),var(--green));transition:width 0.6s ease}

/* ── COURSE LINK BUTTONS ────────────────────────────── */
[data-testid="stLinkButton"]>a{
  background:rgba(45,212,191,0.06)!important;
  border:1px solid var(--bhi)!important;
  border-radius:6px!important;
  color:var(--teal)!important;
  font-family:var(--mono)!important;
  font-size:0.75rem!important;
  padding:6px 12px!important;
  text-decoration:none!important;
  display:flex!important;
  align-items:center!important;
  gap:6px!important;
  margin-top:8px!important;
}
[data-testid="stLinkButton"]>a:hover{
  background:rgba(45,212,191,0.12)!important;
  border-color:var(--teal)!important;
}

/* ── REASONING TRACE ─────────────────────────────────── */
.sf-mod-trace{margin-top:10px;padding:10px 14px;background:rgba(45,212,191,0.04);border:1px solid var(--bhi);border-radius:7px}
.sf-mod-trace-lbl{font-family:var(--mono);font-size:0.6rem;font-weight:600;letter-spacing:0.1em;text-transform:uppercase;color:var(--teal);margin-bottom:5px}
.sf-mod-trace-body{font-size:0.78rem;color:var(--t2);line-height:1.6}

/* ── IMPACT BOX ──────────────────────────────────────── */
.sf-impact-box{background:var(--s1);border:1px solid var(--border);border-radius:10px;padding:22px 28px;margin:24px 0}
.sf-impact-row{display:flex;align-items:center;justify-content:center;gap:0;margin-bottom:12px}
.sf-impact-item{text-align:center;flex:1}
.sf-impact-lbl{font-family:var(--mono);font-size:0.62rem;text-transform:uppercase;letter-spacing:0.08em;color:var(--t3);margin-bottom:6px}
.sf-impact-val{font-family:var(--mono);font-size:2rem;font-weight:500;line-height:1}
.sf-impact-arrow{font-size:1.2rem;color:var(--t4);padding:0 20px;flex-shrink:0}
.sf-impact-sub{font-family:var(--mono);font-size:0.65rem;color:var(--t3);text-align:center;line-height:1.6}


</style>
"""

# =============================================================================
#  SESSION STATE
# =============================================================================
def _init_state():
    defaults = {
        "step":           "input",
        "resume_text":    "",
        "resume_image":   None,
        "jd_text":        "",
        "result":         None,
        "completed":      [],
        "hpd":            2,
        "rw_result":      None,
        "course_cache":   {},
        "sal_location":   "India",
        "force_fresh":    False,
        "search_query":   "",
        "search_results": [],
        "_resume_hash":   "",
        "_resume_source": "",
        "_resume_fname":  "",
        "_jd_source":     "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

_RESET_KEYS = [
    "step","resume_text","resume_image","jd_text","result","completed",
    "rw_result","course_cache","force_fresh","search_query","search_results",
    "res_paste","jd_paste","_resume_source","_resume_fname","_jd_source",
    "search_input","hpd_s","_resume_hash",
]

def _full_reset():
    for k in _RESET_KEYS:
        if k in st.session_state: del st.session_state[k]
    st.rerun()

def _parse_bytes(raw_bytes: bytes, filename: str):
    """Parse resume from raw bytes. Returns (text, image_b64)."""
    name = filename.lower()
    if name.endswith(".pdf"):
        try:
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages), None
        except Exception as e:
            return f"[PDF error: {e}]", None
    if name.endswith(".docx"):
        try:
            doc = Document(io.BytesIO(raw_bytes))
            return "\n".join(p.text for p in doc.paragraphs), None
        except Exception as e:
            return f"[DOCX error: {e}]", None
    if any(name.endswith(x) for x in [".jpg",".jpeg",".png",".webp"]):
        media = ("image/jpeg" if name.endswith((".jpg",".jpeg"))
                 else "image/png" if name.endswith(".png") else "image/webp")
        return "", f"data:{media};base64,{base64.b64encode(raw_bytes).decode()}"
    return raw_bytes.decode("utf-8", errors="ignore"), None

# =============================================================================
#  TOPBAR
# =============================================================================
def render_topbar():
    st.markdown("""
    <div class="sf-top">
      <div class="sf-logo">Skill<em>Forge</em></div>
      <div class="sf-top-right">
        <span class="sf-chip on">⚡ Groq-powered</span>
        <span class="sf-chip on">🖼 Vision OCR</span>
        <span class="sf-chip">NetworkX DAG</span>
        <span class="sf-chip">ARTPARK CodeForge 2025</span>
      </div>
    </div>""", unsafe_allow_html=True)

# =============================================================================
#  INPUT PAGE
# =============================================================================
def render_input():
    st.markdown('<div class="sf-page">', unsafe_allow_html=True)

    st.markdown("""
    <div class="sf-hero">
      <div class="sf-eyebrow">ARTPARK CodeForge Hackathon · AI Adaptive Onboarding Engine</div>
      <div class="sf-h1">Skip what you know.<br><span>Learn what you need.</span></div>
      <div class="sf-sub">Upload your resume and target job description. SkillForge maps your exact skill gap and generates a dependency-ordered, personalized learning roadmap.</div>
    </div>""", unsafe_allow_html=True)

    st.markdown("""
    <div class="sf-how">
      <div class="sf-how-step"><div class="sf-how-num">01</div><div class="sf-how-title">Upload Resume + JD</div><div class="sf-how-sub">PDF, DOCX, or image — Vision AI reads it all</div></div>
      <div class="sf-how-arrow">→</div>
      <div class="sf-how-step"><div class="sf-how-num">02</div><div class="sf-how-title">AI Maps Your Gap</div><div class="sf-how-sub">Groq LLaMA extracts skills · detects decay · scores proficiency</div></div>
      <div class="sf-how-arrow">→</div>
      <div class="sf-how-step"><div class="sf-how-num">03</div><div class="sf-how-title">Get Your Roadmap</div><div class="sf-how-sub">NetworkX DAG orders modules by dependency — zero redundancy</div></div>
    </div>""", unsafe_allow_html=True)

    st.markdown("""
    <div class="sf-stats-strip">
      <div class="sf-stat"><span class="sf-stat-n">47</span><span class="sf-stat-l">Courses in catalog</span></div>
      <div class="sf-stat-div"></div>
      <div class="sf-stat"><span class="sf-stat-n">6</span><span class="sf-stat-l">Skill domains</span></div>
      <div class="sf-stat-div"></div>
      <div class="sf-stat"><span class="sf-stat-n">~38h</span><span class="sf-stat-l">Avg hours saved</span></div>
      <div class="sf-stat-div"></div>
      <div class="sf-stat"><span class="sf-stat-n">0</span><span class="sf-stat-l">Hallucinations</span></div>
    </div>""", unsafe_allow_html=True)

    # Sample buttons
    st.markdown('<div style="margin:20px 0 8px"><span class="sf-sample-lbl">Try a sample →</span></div>', unsafe_allow_html=True)
    pc1, pc2, pc3, _ = st.columns([1,1,1,2])
    for col, key, dlbl in zip([pc1,pc2,pc3], SAMPLES, ["💻 Tech Role","🧠 Data / AI Role","👔 Non-Tech Role"]):
        with col:
            if st.button(dlbl, key=f"pre_{key}", use_container_width=True):
                for wk in _RESET_KEYS:
                    if wk in st.session_state: del st.session_state[wk]
                st.session_state["resume_text"]    = SAMPLES[key]["resume"]
                st.session_state["jd_text"]        = SAMPLES[key]["jd"]
                st.session_state["_resume_source"] = "paste"
                st.session_state["step"]           = "analyzing"
                st.rerun()

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    left, right = st.columns(2, gap="large")

    # ── LEFT: Resume ──────────────────────────────────────
    with left:
        src_flag = st.session_state.get("_resume_source","")
        has_resume = bool(st.session_state.get("resume_text","").strip() or st.session_state.get("resume_image"))
        badge = '<span class="sf-ready-badge">✓ Ready</span>' if has_resume else ''
        st.markdown(f'<div class="sf-panel-hd"><span class="sf-panel-icon">📄</span> Your resume {badge}</div>', unsafe_allow_html=True)

        up_tab, paste_tab = st.tabs(["📄 Upload Resume", "✏️ Paste Resume"])

        with up_tab:
            rf = st.file_uploader("Resume", type=["pdf","docx","jpg","jpeg","png","webp"],
                                  key="res_file", label_visibility="collapsed")
            if rf is not None:
                # Read bytes RIGHT NOW — this is the only reliable way
                try:
                    rf.seek(0)
                    raw_bytes = rf.read()
                except Exception:
                    raw_bytes = b""

                if raw_bytes:
                    new_hash = hashlib.md5(raw_bytes).hexdigest()
                    # Only re-parse if genuinely new file
                    if new_hash != st.session_state.get("_resume_hash", ""):
                        txt, img = _parse_bytes(raw_bytes, rf.name)
                        # Wipe everything old
                        for k in list(st.session_state.keys()):
                            if k not in ("sal_location","force_fresh","hpd"):
                                del st.session_state[k]
                        _init_state()
                        st.session_state["resume_text"]    = txt
                        st.session_state["resume_image"]   = img
                        st.session_state["_resume_source"] = "file"
                        st.session_state["_resume_fname"]  = rf.name
                        st.session_state["_resume_hash"]   = new_hash
                        st.session_state["step"]           = "input"

                    # Show status
                    ss_txt = st.session_state.get("resume_text","")
                    ss_img = st.session_state.get("resume_image")
                    if ss_img:
                        st.success(f"✓ {rf.name} — image loaded (Vision AI will analyze)")
                        st.image(rf, use_container_width=True)
                    elif ss_txt and not ss_txt.startswith("["):
                        wc = len(ss_txt.split())
                        st.success(f"✓ {rf.name} — {wc} words extracted")
                        st.caption(f"Preview: {ss_txt[:200]}…")
                    else:
                        st.error(f"Could not read {rf.name}. Try paste or JPG/PNG.")
                else:
                    st.error("File appears empty. Please try again.")

        with paste_tab:
            if st.session_state.get("_resume_source") == "file":
                st.info("📄 File loaded above. Type here to switch to pasted text.")
                if st.button("Switch to paste mode", key="sw_paste"):
                    for k in ["_resume_source","_resume_hash","_resume_fname","resume_text","resume_image","result","rw_result"]:
                        st.session_state.pop(k, None)
                    st.rerun()
            else:
                rp = st.text_area("Resume text", height=220,
                                  placeholder="Paste your resume here...",
                                  key="res_paste", label_visibility="collapsed")
                if rp and rp.strip():
                    st.session_state["resume_text"]    = rp.strip()
                    st.session_state["_resume_source"] = "paste"
                    st.session_state["_resume_hash"]   = ""
                    wc = len(rp.split())
                    if wc > 5:
                        st.markdown(f'<div class="sf-wc">{wc} words</div>', unsafe_allow_html=True)

        if has_resume and src_flag == "file":
            st.markdown('<div class="sf-ghost" style="margin-top:8px">', unsafe_allow_html=True)
            if st.button("✕ Clear resume", key="clr_res"):
                for k in ["resume_text","resume_image","res_paste","_resume_source","_resume_fname","_resume_hash","result","rw_result"]:
                    st.session_state.pop(k, None)
                _init_state()
                st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)

    # ── RIGHT: JD ─────────────────────────────────────────
    with right:
        has_jd = bool(st.session_state.get("jd_text","").strip())
        badge_jd = '<span class="sf-ready-badge">✓ Ready</span>' if has_jd else ''
        st.markdown(f'<div class="sf-panel-hd"><span class="sf-panel-icon">💼</span> Job description {badge_jd}</div>', unsafe_allow_html=True)

        jup_tab, jpaste_tab = st.tabs(["📤 Upload JD", "📝 Paste JD"])

        with jup_tab:
            jf = st.file_uploader("JD", type=["pdf","docx"], key="jd_file", label_visibility="collapsed")
            if jf is not None:
                try:
                    jf.seek(0)
                    jraw = jf.read()
                except Exception:
                    jraw = b""
                if jraw:
                    jtxt, _ = _parse_bytes(jraw, jf.name)
                    if jtxt and not jtxt.startswith("["):
                        for k in ["result","rw_result","course_cache"]:
                            st.session_state.pop(k, None)
                        st.session_state["jd_text"]    = jtxt
                        st.session_state["_jd_source"] = "file"
                        st.session_state["step"]       = "input"
                        st.session_state.pop("jd_paste", None)
                        st.success(f"✓ {jf.name} — {len(jtxt.split())} words")
                    else:
                        st.error(f"Could not read {jf.name}")

        with jpaste_tab:
            jp = st.text_area("Job description", height=220,
                              placeholder="Paste the job description here...",
                              key="jd_paste", label_visibility="collapsed")
            if jp and jp.strip():
                st.session_state["jd_text"]    = jp.strip()
                st.session_state["_jd_source"] = "paste"
                wc_jd = len(jp.split())
                if wc_jd > 5:
                    st.markdown(f'<div class="sf-wc">{wc_jd} words</div>', unsafe_allow_html=True)

    # ── Analyze button ─────────────────────────────────────
    st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)
    opt1, opt2, _, btn_col = st.columns([1, 1.2, 0.6, 1.6])
    with opt1:
        st.selectbox("Salary location", ["India","USA","UK","Germany","Canada","Singapore"],
                     index=["India","USA","UK","Germany","Canada","Singapore"].index(st.session_state.get("sal_location","India")),
                     key="sal_location", label_visibility="visible")
    with opt2:
        st.checkbox("Force fresh (skip cache)", key="force_fresh")

    with btn_col:
        st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
        resume_ready = bool(st.session_state.get("resume_text","").strip() or st.session_state.get("resume_image"))
        jd_ready     = bool(st.session_state.get("jd_text","").strip())
        if resume_ready and jd_ready:
            is_img = bool(st.session_state.get("resume_image"))
            lbl = "Analyze resume image ⚡" if is_img else "Analyze skill gap ⚡"
            if st.button(lbl, key="go_btn", use_container_width=True):
                st.session_state["step"] = "analyzing"
                st.rerun()
        else:
            missing = []
            if not resume_ready: missing.append("resume")
            if not jd_ready:     missing.append("job description")
            st.markdown(f'<p style="font-family:var(--mono);font-size:0.75rem;color:var(--t3);text-align:center;padding:12px 0">Add {" and ".join(missing)} to continue</p>', unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)



# =============================================================================
#  LOADING
# =============================================================================
def render_loading():
    st.markdown('<div class="sf-page">', unsafe_allow_html=True)
    st.markdown("<div style='height:48px'></div>", unsafe_allow_html=True)

    resume_text = st.session_state.get("resume_text","")
    resume_img  = st.session_state.get("resume_image")
    jd_text     = st.session_state.get("jd_text","") or st.session_state.get("jd_paste","")
    source      = st.session_state.get("_resume_source","paste")

    if not resume_text.strip() and not resume_img:
        st.error("No resume data. Please go back and upload your resume.")
        if st.button("Go back", key="back_no_resume"):
            st.session_state["step"] = "input"; st.rerun()
        return

    if not jd_text.strip():
        st.error("No job description. Please go back and add one.")
        if st.button("Go back", key="back_no_jd"):
            st.session_state["step"] = "input"; st.rerun()
        return

    if source == "file" or st.session_state.get("force_fresh"):
        try:
            ck = ("txt:" + hashlib.md5(resume_text.encode()).hexdigest()) if resume_text.strip() else                  ("img:" + hashlib.md5((resume_img or "").encode()).hexdigest())
            with shelve.open(_CACHE_PATH) as db:
                kk = _ckey(ck, jd_text)
                if kk in db: del db[kk]
        except: pass
        st.session_state.pop("force_fresh", None)

    steps = [
        ("📄","Parsing resume & job description","Extracting text, structure, metadata"),
        ("🔍","Extracting skills with proficiency","Scoring 0-10 per skill, detecting decay"),
        ("🧩","Computing skill gap","Known, Partial, Missing classification"),
        ("🗺","Building dependency roadmap","NetworkX DAG, topological sort"),
        ("🌐","Fetching live market data","Salary, trends, job market via DuckDuckGo"),
    ]
    st.markdown("""<div style="max-width:560px;margin:40px auto 32px">
      <div style="font-family:var(--mono);font-size:0.65rem;letter-spacing:0.12em;text-transform:uppercase;color:var(--teal);margin-bottom:20px">Analyzing your profile</div>
    </div>""", unsafe_allow_html=True)
    slots = [st.empty() for _ in steps]
    prog  = st.empty()

    def show_steps(done):
        for i,(icon,title,sub) in enumerate(steps):
            if   i < done:  s = f'<div class="sf-lstep sf-lstep-done"><span class="sf-lstep-icon">✓</span><div><div class="sf-lstep-title">{title}</div><div class="sf-lstep-sub">{sub}</div></div></div>'
            elif i == done: s = f'<div class="sf-lstep sf-lstep-active"><span class="sf-lstep-icon sf-lstep-spin">{icon}</span><div><div class="sf-lstep-title">{title}</div><div class="sf-lstep-sub">{sub}</div></div></div>'
            else:           s = f'<div class="sf-lstep sf-lstep-wait"><span class="sf-lstep-icon">○</span><div><div class="sf-lstep-title">{title}</div><div class="sf-lstep-sub">{sub}</div></div></div>'
            slots[i].markdown(s, unsafe_allow_html=True)
        prog.markdown(f'<div class="sf-lprog"><div class="sf-lprog-fill" style="width:{int(done/len(steps)*100)}%"></div></div>', unsafe_allow_html=True)

    show_steps(0)
    result = run_analysis_with_web(
        resume_text, jd_text,
        resume_image_b64=resume_img,
        location=st.session_state.get("sal_location","India"),
    )
    if "error" not in result:
        show_steps(3); show_steps(4); show_steps(5)

    if "error" in result:
        err = result.get("error","unknown")
        if err == "rate_limited":
            st.error(f"Rate limited — {result.get('message','')}")
        elif "vision" in str(err).lower():
            st.error(f"Vision model error: {err}")
            st.info("Try pasting resume text instead.")
        else:
            st.error(f"Analysis failed: {err}")
        st.markdown('<div class="sf-ghost">', unsafe_allow_html=True)
        if st.button("Back", key="retry_btn"):
            st.session_state["step"] = "input"; st.rerun()
        st.markdown("</div></div>", unsafe_allow_html=True)
        return

    st.session_state["result"] = result
    st.session_state["step"]   = "results"
    st.rerun()

#  RESULTS: SCORE BANNER
# =============================================================================
def render_banner(res):
    c   = res["candidate"]; jd = res["jd"]
    im  = res["impact"];    iv = res["interview"]
    ql  = res.get("quality",{})

    name   = c.get("name","Unknown")
    crole  = c.get("current_role","")
    yrs    = c.get("years_experience",0)
    sen    = c.get("seniority","")
    trole  = jd.get("role_title","")

    cur   = im.get("current_fit", 0)  or 0
    proj  = im.get("projected_fit", 0) or 0
    delta = im.get("fit_delta", 0)       or 0
    iv_c  = iv.get("color","#4ade80")
    ats   = ql.get("ats_score", 0) or 0
    grade = ql.get("overall_grade","–")

    fit_c = _RED if cur < 40 else _AMBER if cur < 65 else _GREEN

    cache_badge   = '<span class="sf-cache-badge">⚡ Cached</span>' if res.get("_cache_hit") else ""
    vision_badge  = '<span class="sf-vision-badge">🖼 Vision OCR</span>' if res.get("_is_image") else ""

    hpd_val = st.session_state.get("hpd", 2)
    st.markdown(f"""
    <div class="sf-banner">
      <div class="sf-banner-top">
        <div>
          <div class="sf-candidate-name">{name}</div>
          <div class="sf-candidate-sub">{crole} · {yrs}yr experience · {sen} level → <strong style="color:var(--t1)">{trole}</strong></div>
        </div>
        <div style="display:flex;gap:6px;align-items:center;margin-left:auto">
          {cache_badge}{vision_badge}
        </div>
      </div>

      <!-- HERO NUMBER -->
      <div class="sf-hero-delta">
        <div class="sf-hero-delta-num">+{delta}%</div>
        <div class="sf-hero-delta-label">role fit after completing this roadmap</div>
        <div class="sf-hero-delta-sub">
          <span style="color:{fit_c}">{cur}%</span>
          <span style="color:var(--t3);margin:0 10px">→</span>
          <span style="color:var(--green)">{proj}%</span>
          <span style="color:var(--t3);margin-left:10px">· Complete in {weeks_ready(im.get("roadmap_hours",0) or 0, hpd_val)} at {hpd_val}h/day</span>
        </div>
      </div>

      <!-- SUPPORTING SCORES -->
      <div class="sf-scores">
        <div class="sf-score-card">
          <div class="sf-score-lbl">Training hours</div>
          <div class="sf-score-num" style="color:var(--teal)">{im['roadmap_hours']}h</div>
          <div class="sf-score-sub">vs <s>60h</s> generic onboarding · saves <strong style="color:var(--green)">~{im['hours_saved']}h</strong></div>
        </div>
        <div class="sf-score-card">
          <div class="sf-score-lbl">Interview readiness</div>
          <div class="sf-score-num" style="color:{iv_c}">{iv['score']}%</div>
          <div class="sf-score-sub">{iv['label']} · {iv.get('advice','')}</div>
        </div>
        <div class="sf-score-card">
          <div class="sf-score-lbl">ATS score</div>
          <div class="sf-score-num" style="color:var(--t1)">{ats}%</div>
          <div class="sf-score-sub">Grade <strong style="color:var(--teal)">{grade}</strong> · {ql.get('completeness_score',0)}% complete</div>
        </div>
      </div>

      <!-- GROUNDING BADGE -->
      <div class="sf-ground-badge">
        <span class="sf-ground-dot"></span>
        Zero hallucinations &nbsp;·&nbsp; All {im['modules_count']} modules sourced strictly from 47-course catalog &nbsp;·&nbsp; {im['critical_count']} on critical dependency path
      </div>
    </div>""", unsafe_allow_html=True)

    k1,k2,k3,k4 = st.columns(4)
    k1.metric("Modules",          im.get("modules_count",0),    f"{im.get('critical_count',0)} critical path")
    k2.metric("Hours saved",      f"~{im['hours_saved']}h",     "vs 60h generic onboarding")
    k3.metric("Done in",          weeks_ready(im.get("roadmap_hours",0) or 0, hpd_val), f"at {hpd_val}h/day")
    k4.metric("Skills covered",   f"{im['gaps_addressed']}/{im['total_skills']}", f"{im['known_skills']} already known — skipped")

    sm = res.get("seniority",{})
    if sm.get("has_mismatch"):
        st.markdown(f'<div class="sf-warn">⚠ Seniority gap: you are <strong>{sm["candidate"]}</strong>, role requires <strong>{sm["required"]}</strong> — leadership modules injected into your roadmap.</div>', unsafe_allow_html=True)
    if im.get("decayed_skills",0):
        st.markdown(f'<div style="background:rgba(245,158,11,0.06);border:1px solid rgba(245,158,11,0.18);border-radius:7px;padding:10px 14px;font-size:0.82rem;color:var(--amber);margin-bottom:6px">⏱ {im["decayed_skills"]} skill(s) have decayed from inactivity — proficiency reduced in gap analysis</div>', unsafe_allow_html=True)

# =============================================================================
#  TAB 1 — OVERVIEW
# =============================================================================
def render_tab_overview(res):
    gp     = res["gap_profile"]
    trends = res.get("skill_trends",{})
    sal    = res.get("salary",{})

    k_c = sum(1 for g in gp if g["status"]=="Known")
    p_c = sum(1 for g in gp if g["status"]=="Partial")
    m_c = sum(1 for g in gp if g["status"]=="Missing")

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    st.markdown('<div class="sf-sh">Skill gap</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="sf-ss">{k_c} known · {p_c} partial · {m_c} missing — required skills shown first</div>', unsafe_allow_html=True)

    col_flt, col_radar = st.columns([1.4,1], gap="large")

    with col_flt:
        filt = st.selectbox("Filter", ["All","Missing","Partial","Known","Required only"],
                            key="gp_filter", label_visibility="collapsed")

        filtered = [g for g in gp if
                    filt=="All" or
                    (filt=="Missing"       and g["status"]=="Missing") or
                    (filt=="Partial"       and g["status"]=="Partial") or
                    (filt=="Known"         and g["status"]=="Known") or
                    (filt=="Required only" and g["is_required"])]

        # Build all card HTML in one pass then render once (prevents grid layout breaks)
        cards_html = '<div class="sf-skill-grid">'
        for g in filtered:
            s    = g["status"]
            col_map  = {"Known":_TEAL,"Partial":_AMBER,"Missing":_RED}
            bc_map   = {"Known":"sf-st-known","Partial":"sf-st-partial","Missing":"sf-st-missing"}
            col_c    = col_map[s]
            bc       = bc_map[s]
            cls      = s.lower()
            pct      = (g.get("proficiency", 0) or 0)/10*100
            req      = "★ " if g["is_required"] else ""
            trend    = trends.get(g["skill"],"")
            tc       = (_RED if "Hot" in trend else _AMBER if "Growing" in trend else "#3d4d66")
            decay    = '<span class="sf-decay-tag">⏱ decayed</span>' if g.get("decayed") else ""
            ctx      = f'<div class="sf-skill-ctx">{g["context"]}</div>' if g.get("context") else ""
            co       = g.get("catalog_course")
            course_txt = (
                f'<div style="font-family:var(--mono);font-size:0.65rem;color:var(--t3);margin-top:5px">'
                f'📚 {co["title"]} · {co["duration_hrs"]}h · {co["level"]}</div>'
            ) if co else ""
            cards_html += f"""
            <div class="sf-skill-card {cls}">
              <div class="sf-skill-top">
                <div class="sf-skill-name">{req}{g['skill']}</div>
                <span class="sf-st-badge {bc}">{s}</span>
              </div>
              <div class="sf-skill-bar">
                <div class="sf-skill-bar-fill" style="width:{pct}%;background:{col_c}"></div>
              </div>
              <div class="sf-skill-bottom">
                <span class="sf-skill-score">{g['proficiency']}/10</span>
                <span class="sf-skill-demand" style="color:{tc}">{trend}</span>
              </div>
              {decay}{ctx}{course_txt}
            </div>"""
        cards_html += '</div>'
        st.markdown(cards_html, unsafe_allow_html=True)

        obs = res.get("obsolescence",[])
        if obs:
            st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)
            st.markdown('<div class="sf-sh" style="font-size:0.95rem">Obsolescence risks</div>', unsafe_allow_html=True)
            for o in obs:
                st.markdown(f'<div style="background:rgba(239,68,68,0.05);border:1px solid rgba(239,68,68,0.15);border-radius:7px;padding:10px 14px;margin-bottom:6px"><div style="font-size:0.88rem;font-weight:600;color:var(--red)">{o["skill"]}</div><div style="font-family:var(--mono);font-size:0.7rem;color:var(--t3);margin-top:3px">{o["reason"]}</div></div>', unsafe_allow_html=True)

    with col_radar:
        st.plotly_chart(radar_chart(gp), use_container_width=True, config={"displayModeBar":False})

        tf = res.get("transfers",[])
        if tf:
            st.markdown('<div style="font-size:0.88rem;font-weight:600;color:var(--t1);margin:14px 0 8px">Transfer advantages</div>', unsafe_allow_html=True)
            for t in tf[:5]:
                st.markdown(f'<div class="sf-xfer"><span class="sf-xfer-pct">↗{t["transfer_pct"]}%</span><span>{t["label"]}</span></div>', unsafe_allow_html=True)

        if sal and sal.get("median_lpa",0):
            st.markdown(f'<div style="font-size:0.88rem;font-weight:600;color:var(--t1);margin:16px 0 4px">Live salary — {res["jd"].get("role_title","")[:24]}</div>', unsafe_allow_html=True)
            st.plotly_chart(salary_chart(sal), use_container_width=True, config={"displayModeBar":False})
            st.caption(f"Source: {sal.get('source','web')} · {sal.get('note','')}")

# =============================================================================
#  TAB 2 — ROADMAP
# =============================================================================
def render_tab_roadmap(res):
    path      = res["path"]
    completed = set(st.session_state.get("completed", []))  # FIX Bug 4: load as set

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    hd_l, hd_r = st.columns([2,1])
    with hd_l:
        st.markdown('<div class="sf-sh">Learning roadmap</div>', unsafe_allow_html=True)
        st.markdown('<div class="sf-ss">Dependency-ordered · critical path highlighted · check modules off as you complete them</div>', unsafe_allow_html=True)
    with hd_r:
        hpd = st.select_slider("Pace (h/day)", options=[1,2,4,8],
                               value=st.session_state.get("hpd",2), key="hpd_s",
                               label_visibility="visible")
        st.session_state["hpd"] = hpd  # safe: different key than widget "hpd_s"
        rem = sum(m.get("duration_hrs",0) or 0 for m in path if m["id"] not in completed)
        st.markdown(f'<p style="font-family:var(--mono);font-size:0.72rem;color:var(--t2);text-align:right">{rem}h left · done in <strong style="color:var(--teal)">{weeks_ready(rem,hpd)}</strong></p>', unsafe_allow_html=True)

    mod_col, chart_col = st.columns([1.1,1], gap="large")

    with mod_col:
        phases = [
            ("Foundation", [m for m in path if m["level"]=="Beginner"]),
            ("Build",       [m for m in path if m["level"]=="Intermediate"]),
            ("Advanced",    [m for m in path if m["level"]=="Advanced"]),
        ]
        idx = 0
        for phase_name, mods in phases:
            if not mods: continue
            phase_hrs = sum(m.get("duration_hrs",0) or 0 for m in mods)
            st.markdown(f'<div class="sf-phase-hd">{phase_name} &nbsp; <span style="font-weight:400;color:var(--t4)">{len(mods)} modules · {phase_hrs}h</span></div>', unsafe_allow_html=True)

            for m in mods:
                idx += 1
                is_done = m["id"] in completed
                is_crit = m.get("is_critical",False)
                level   = m["level"]
                lc = "crit" if is_crit else "adv" if level=="Advanced" else "inter" if level=="Intermediate" else "beg"
                dc = " done" if is_done else ""

                # FIX: checkbox label must be plain text — no markdown backticks or special chars
                chk = st.checkbox(
                    f"{m['title']} · {m['duration_hrs']}h · {m['level']}",
                    value=is_done, key=f"c_{m['id']}"
                )
                if chk: completed.add(m["id"])
                else:    completed.discard(m["id"])
                st.session_state["completed"] = list(completed)

                prereqs_txt = ", ".join(m.get("prereqs",[]) or []) or "none"
                tags = []
                if is_crit: tags.append('<span class="sf-tag sf-tag-crit">★ critical</span>')
                if m.get("is_required"): tags.append('<span class="sf-tag sf-tag-req">required</span>')
                tags.append(f'<span class="sf-tag">{m["domain"]}</span>')
                tags_html = "".join(tags)

                num_label = "✓" if is_done else f"#{idx:02d}"
                reason_html = (
                    '<div class="sf-mod-trace">' +
                    '<span class="sf-mod-trace-lbl">🧠 AI Reasoning</span>' +
                    f'<div class="sf-mod-trace-body">{m["reasoning"]}</div>' +
                    '</div>'
                ) if m.get("reasoning") else ""

                # Render module card WITHOUT course link (sanitizer strips nested <a> tags)
                st.markdown(
                    f'<div class="sf-mod {lc}{dc}">' +
                    f'  <div class="sf-mod-row">' +
                    f'    <div class="sf-mod-num">{num_label}</div>' +
                    f'    <div class="sf-mod-body">' +
                    f'      <div class="sf-mod-title">{m["title"]}</div>' +
                    f'      <div class="sf-mod-meta">Skill: {m["skill"]} · prereqs: {prereqs_txt}</div>' +
                    f'      <div class="sf-mod-tags">{tags_html}</div>' +
                    f'      {reason_html}' +
                    f'    </div>' +
                    f'    <div class="sf-mod-hrs">{m.get("duration_hrs",0) or 0}h</div>' +
                    f'  </div>' +
                    f'</div>',
                    unsafe_allow_html=True
                )

                # Render course links separately using native Streamlit — bypasses sanitizer
                courses = st.session_state.get("course_cache",{}).get(m["skill"],[])
                if courses:
                    crs = courses[0]
                    lnk_col1, lnk_col2 = st.columns([3,1])
                    with lnk_col1:
                        st.link_button(
                            f"{crs['icon']} {crs['title'][:50]} ({crs['platform']})",
                            crs["url"],
                            use_container_width=True
                        )
                else:
                    q = urllib.parse.quote_plus(m["skill"])
                    lnk_col1, lnk_col2 = st.columns(2)
                    with lnk_col1:
                        st.link_button(
                            f"🎓 {m['skill']} on Coursera",
                            f"https://www.coursera.org/search?query={q}",
                            use_container_width=True
                        )
                    with lnk_col2:
                        st.link_button(
                            f"▶ {m['skill']} on YouTube",
                            f"https://www.youtube.com/results?search_query={q}+tutorial",
                            use_container_width=True
                        )

    with chart_col:
        st.markdown('<div style="font-size:0.88rem;font-weight:600;color:var(--t1);margin-bottom:8px">ROI ranking</div>', unsafe_allow_html=True)
        st.plotly_chart(roi_bar(res.get("roi",[])), use_container_width=True, config={"displayModeBar":False})

    st.markdown('<div class="sf-divider"></div>', unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.88rem;font-weight:600;color:var(--t1);margin-bottom:10px">Training timeline</div>', unsafe_allow_html=True)
    st.plotly_chart(timeline_chart(path), use_container_width=True, config={"displayModeBar":False})

    st.markdown('<div class="sf-divider"></div>', unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.88rem;font-weight:600;color:var(--t1);margin-bottom:10px">Weekly study plan</div>', unsafe_allow_html=True)
    rem_path = [m for m in path if m["id"] not in completed]
    wp = weekly_plan(rem_path, hpd)
    for w in wp[:8]:
        with st.expander(f"Week {w['week']} — {w['total_hrs']:.1f}h / {hpd*5}h capacity"):
            for mx in w["modules"]:
                star = "★ " if mx.get("is_critical") else ""
                st.markdown(f"- {star}**{mx['title']}** &nbsp;·&nbsp; `{mx['hrs_this_week']:.1f}h` of `{mx['total_hrs']}h`")

    # ── Impact summary ───────────────────────────────────────
    im = res["impact"]
    saved = im.get("hours_saved", 0)
    st.markdown(f"""
    <div class="sf-impact-box">
      <div class="sf-impact-row">
        <div class="sf-impact-item">
          <div class="sf-impact-lbl">Generic onboarding</div>
          <div class="sf-impact-val" style="color:var(--t3);text-decoration:line-through">60h</div>
        </div>
        <div class="sf-impact-arrow">→</div>
        <div class="sf-impact-item">
          <div class="sf-impact-lbl">Your personalized path</div>
          <div class="sf-impact-val" style="color:var(--teal)">{im['roadmap_hours']}h</div>
        </div>
        <div class="sf-impact-arrow">→</div>
        <div class="sf-impact-item">
          <div class="sf-impact-lbl">Time saved</div>
          <div class="sf-impact-val" style="color:var(--green)">~{saved}h</div>
        </div>
      </div>
      <div class="sf-impact-sub">
        Powered by NetworkX DAG · {im['modules_count']} modules · {im['critical_count']} critical path nodes · {im['known_skills']} skills skipped (already known)
      </div>
    </div>""", unsafe_allow_html=True)

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    gap_skills_u = list({m["skill"] for m in path})
    loaded = len(st.session_state.get("course_cache",{}))
    if loaded < len(gap_skills_u):
        if st.button(f"Load course links for all {len(gap_skills_u)} skills →", key="load_all_crs"):
            with st.spinner("Searching Coursera · Udemy · YouTube…"):
                cc = {}
                with ThreadPoolExecutor(max_workers=4) as ex:
                    futs = {ex.submit(search_course_links,s): s for s in gap_skills_u[:10]}
                    for f in futs: cc[futs[f]] = f.result()
            st.session_state["course_cache"] = cc
            st.success(f"✓ {len(cc)} skills with course links")
            st.rerun()

# =============================================================================
#  TAB 3 — RESEARCH
# =============================================================================
def render_tab_research(res):
    gp  = res["gap_profile"]
    sal = res.get("salary",{})
    mkt = res.get("market_insights",[])
    trends = res.get("skill_trends",{})
    jd  = res["jd"]

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    st.markdown('<div class="sf-sh">Web research</div>', unsafe_allow_html=True)
    st.markdown('<div class="sf-ss">Live search via DuckDuckGo — courses, salaries, market trends, anything</div>', unsafe_allow_html=True)

    q_col, btn_col = st.columns([5,1])
    with q_col:
        if "search_input" not in st.session_state:
            st.session_state["search_input"] = st.session_state.get("search_query","")
        st.text_input("Search", placeholder='e.g. "React vs Vue 2025" · "FastAPI senior salary Bangalore" · "Docker Kubernetes course"',
                      key="search_input", label_visibility="collapsed")
        st.session_state["search_query"] = st.session_state.get("search_input","")
    with btn_col:
        do_search = st.button("Search →", key="go_search", use_container_width=True)

    gap_skills_s = [g["skill"] for g in gp if g["status"]!="Known"][:4]
    role_name = jd.get("role_title","")
    shortcuts = [(s, f"{s} online course tutorial 2025") for s in gap_skills_s]
    shortcuts.append((f"{role_name[:18]} salary", f"{role_name} salary {st.session_state.get('sal_location','India')} 2025"))

    # FIX #17: Guard columns to match actual shortcuts count
    n_shortcuts = len(shortcuts)
    if n_shortcuts > 0:
        sc_cols = st.columns(n_shortcuts)
        for i, (lbl, q) in enumerate(shortcuts):
            with sc_cols[i]:
                if st.button(lbl[:22], key=f"sc_{i}", use_container_width=True):
                    if "search_input" in st.session_state: del st.session_state["search_input"]
                    st.session_state["search_query"] = q
                    st.session_state["search_results"] = ddg_search(q, max_results=8)
                    st.rerun()

    if do_search and st.session_state.get("search_query","").strip():
        with st.spinner("Searching…"):
            st.session_state["search_results"] = ddg_search(st.session_state["search_query"], max_results=8)

    results = st.session_state.get("search_results",[])
    if _DDG_ERROR:
        st.warning(f"⚠ {_DDG_ERROR}")
    if results:
        # Filter non-English results
        eng_results = [r for r in results if _is_english(r.get("title","")) and _is_english(r.get("body",""))]
        shown = eng_results if eng_results else results  # fallback to all if none pass
        st.markdown(f'<div style="font-family:var(--mono);font-size:0.68rem;color:var(--t3);margin:10px 0 8px">{len(shown)} results for "{st.session_state.get("search_query","")}"</div>', unsafe_allow_html=True)
        for r in shown:
            href  = r.get("href","")
            title = r.get("title","No title")
            body  = r.get("body","")[:200]
            domain = href.split("/")[2] if href.count("/") >= 2 else href
            st.markdown(f"""
            <div class="sf-search-result">
              <a class="sf-search-title" href="{href}" target="_blank">{title}</a>
              <div class="sf-search-url">{domain}</div>
              <div class="sf-search-body">{body}</div>
            </div>""", unsafe_allow_html=True)
    elif do_search:
        st.warning("No results — try rephrasing the query.")

    st.markdown('<div class="sf-divider"></div>', unsafe_allow_html=True)

    sal_col, mkt_col = st.columns(2, gap="large")
    with sal_col:
        st.markdown('<div style="font-size:0.9rem;font-weight:600;color:var(--t1);margin-bottom:10px">Salary benchmark</div>', unsafe_allow_html=True)
        _sal_val = sal.get("median_lpa",0) if sal else 0
        try: _sal_val = float(_sal_val or 0)
        except: _sal_val = 0
        if sal and _sal_val > 0:
            st.plotly_chart(salary_chart(sal), use_container_width=True, config={"displayModeBar":False})
            st.caption(f"Source: {sal.get('source','web')} · {sal.get('note','')}")
        else:
            st.markdown('<p style="font-family:var(--mono);font-size:0.75rem;color:var(--t3)">Salary data not available — fetch it:</p>', unsafe_allow_html=True)
            loc2 = st.selectbox("Location", ["India","USA","UK","Germany","Canada","Singapore"], key="sal_loc2")
            if st.button("🔍 Fetch salary data", key="sal_fetch", use_container_width=True):
                with st.spinner("Searching salary data…"):
                    _sal_new = search_real_salary(role_name, loc2)
                    st.session_state["result"]["salary"] = _sal_new
                if not _sal_new:
                    st.warning("Could not find salary data. Try a different location or search manually above.")
                st.rerun()

    with mkt_col:
        st.markdown('<div style="font-size:0.9rem;font-weight:600;color:var(--t1);margin-bottom:10px">Job market insights</div>', unsafe_allow_html=True)
        if mkt:
            for ins in mkt:
                st.markdown(f'<div class="sf-insight">📌 {ins}</div>', unsafe_allow_html=True)
        else:
            st.markdown('<p style="font-family:var(--mono);font-size:0.75rem;color:var(--t3)">Not fetched yet</p>', unsafe_allow_html=True)
            if st.button("Fetch insights", key="mkt_fetch"):
                with st.spinner("Searching…"):
                    st.session_state["result"]["market_insights"] = search_job_market(role_name)
                st.rerun()

    if trends:
        st.markdown('<div class="sf-divider"></div>', unsafe_allow_html=True)
        st.markdown('<div style="font-size:0.9rem;font-weight:600;color:var(--t1);margin-bottom:10px">Skill demand signals</div>', unsafe_allow_html=True)
        pills = ""
        for skill, sig in trends.items():
            sc = _RED if "Hot" in sig else _AMBER if "Growing" in sig else "#3d4d66"
            pills += f'<span class="sf-trend-pill"><span style="color:var(--t1);font-weight:500">{skill[:14]}</span><span style="color:{sc}">&nbsp;{sig}</span></span>'
        st.markdown(f'<div style="line-height:2.6">{pills}</div>', unsafe_allow_html=True)
        if st.button("Re-fetch trends", key="refetch_trends"):
            gs = [g["skill"] for g in gp if g["status"]!="Known"][:6]
            with st.spinner("Checking latest demand data…"):
                st.session_state["result"]["skill_trends"] = search_skill_trends(gs)
            st.rerun()

    st.markdown('<div class="sf-divider"></div>', unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.9rem;font-weight:600;color:var(--t1);margin-bottom:8px">Course finder</div>', unsafe_allow_html=True)
    gap_skills_f = [g["skill"] for g in gp if g["status"]!="Known"]
    if gap_skills_f:
        crs_c1, crs_c2 = st.columns([3,1])
        with crs_c1:
            sel_s = st.selectbox("Skill to search courses for:", gap_skills_f, key="crs_sel",
                                 label_visibility="collapsed")
        with crs_c2:
            if st.button("Find →", key="crs_go", use_container_width=True):
                with st.spinner(f"Searching {sel_s}…"):
                    cc = st.session_state.get("course_cache",{})
                    cc[sel_s] = search_course_links(sel_s)
                    st.session_state["course_cache"] = cc
                st.rerun()  # show results immediately

        cached = st.session_state.get("course_cache",{}).get(sel_s,[])
        if cached:
            for crs in cached:
                st.markdown(f'<div class="sf-search-result"><a class="sf-search-title" href="{crs["url"]}" target="_blank">{crs["icon"]} {crs["title"]}</a><div class="sf-search-url">{crs["platform"]}</div><div class="sf-search-body">{crs["snippet"]}</div></div>', unsafe_allow_html=True)
        elif sel_s in st.session_state.get("course_cache",{}):
            st.info(f"No links found for {sel_s}. Try searching manually above.")

# =============================================================================
#  TAB 4 — ATS & EXPORT
# =============================================================================
def render_tab_ats_export(res):
    c    = res["candidate"]; jd = res["jd"]
    gp   = res["gap_profile"]; roadmap = res["path"]
    im   = res["impact"];    ql = res.get("quality",{})
    iv   = res.get("interview",{})
    sm   = res.get("seniority",{}); cgm = res.get("career_months",0)

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    st.markdown('<div class="sf-sh">ATS audit</div>', unsafe_allow_html=True)
    st.markdown('<div class="sf-ss">Resume quality scores · improvement tips · ATS keyword gaps · interview talking points</div>', unsafe_allow_html=True)

    ats_pct = ql.get("ats_score",0) or 0
    st.markdown(f"""
    <div class="sf-ats-row">
      <div class="sf-ats-card"><div class="sf-ats-n">{ql.get('ats_score','–')}%</div><div class="sf-ats-l">ATS Score</div></div>
      <div class="sf-ats-card"><div class="sf-ats-n" style="color:var(--teal)">{ql.get('overall_grade','–')}</div><div class="sf-ats-l">Grade</div></div>
      <div class="sf-ats-card"><div class="sf-ats-n">{ql.get('completeness_score','–')}%</div><div class="sf-ats-l">Completeness</div></div>
      <div class="sf-ats-card"><div class="sf-ats-n">{ql.get('clarity_score','–')}%</div><div class="sf-ats-l">Clarity</div></div>
    </div>
    <div class="sf-prog"><div class="sf-prog-fill" style="width:{ats_pct}%"></div></div>
    """, unsafe_allow_html=True)

    left, right = st.columns(2, gap="large")
    with left:
        st.markdown('<div style="font-size:0.88rem;font-weight:600;color:var(--t1);margin-bottom:10px">Improvement tips</div>', unsafe_allow_html=True)
        for i, tip in enumerate((ql.get("improvement_tips") or [])[:6]):
            st.markdown(f'<div class="sf-tip"><span class="sf-tip-n">{str(i+1).zfill(2)}</span><span>{tip}</span></div>', unsafe_allow_html=True)

        st.markdown('<div style="font-size:0.88rem;font-weight:600;color:var(--t1);margin:16px 0 8px">Interview talking points</div>', unsafe_allow_html=True)
        for talking_pt in (ql.get("interview_talking_points") or [])[:4]:
            st.markdown(f'<div class="sf-talk">→ {talking_pt}</div>', unsafe_allow_html=True)

    with right:
        st.markdown('<div style="font-size:0.88rem;font-weight:600;color:var(--t1);margin-bottom:8px">ATS issues</div>', unsafe_allow_html=True)
        issues = ql.get("ats_issues") or []
        if issues:
            for iss in issues[:5]: st.warning(iss)
        else:
            st.success("No critical ATS issues found")

        st.markdown('<div style="font-size:0.88rem;font-weight:600;color:var(--t1);margin:12px 0 6px">Missing keywords</div>', unsafe_allow_html=True)
        kws = ql.get("missing_keywords") or []
        if kws:
            st.markdown("".join(f'<span class="sf-kw">{k}</span>' for k in kws), unsafe_allow_html=True)
        else:
            st.markdown('<span style="font-family:var(--mono);font-size:0.72rem;color:var(--t3)">None identified</span>', unsafe_allow_html=True)

        st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)
        c1,c2,c3 = st.columns(3)
        c1.metric("Interview",  f"{iv.get('score',0)}%",  iv.get("label","–"))
        c2.metric("Seniority gap", f"{sm.get('gap_levels',0)} lvl")
        c3.metric("Career est.", f"~{cgm}mo" if cgm else "–")

    st.markdown('<div class="sf-divider"></div>', unsafe_allow_html=True)
    st.markdown('<div style="font-size:0.88rem;font-weight:600;color:var(--t1);margin-bottom:4px">AI resume rewrite</div>', unsafe_allow_html=True)
    st.markdown('<div style="font-family:var(--mono);font-size:0.68rem;color:var(--t3);margin-bottom:10px">ATS-optimized version with missing keywords added naturally</div>', unsafe_allow_html=True)

    rtxt = st.session_state.get("resume_text","")
    is_image_resume = bool(st.session_state.get("resume_image"))
    if not rtxt and is_image_resume:
        st.info("Resume rewrite works on extracted text. For image resumes, the extracted text from vision analysis is used.")
        # Attempt to use candidate name + skills as a text proxy
        extracted_info = res.get("candidate", {})
        if extracted_info.get("name"):
            rtxt = f"{extracted_info.get('name','')}\n{extracted_info.get('current_role','')}\n" + \
                   "\n".join([f"- {s['skill']} ({s['proficiency']}/10)" for s in extracted_info.get("skills",[])])
    elif not rtxt:
        st.info("Resume text required for rewrite.")

    if rtxt:
        if st.button("Generate rewrite →", key="gen_rw"):
            with st.spinner("Rewriting with LLaMA 3.3-70b…"):
                rw = rewrite_resume(rtxt, jd, kws)
            st.session_state["rw_result"] = rw

        rw = st.session_state.get("rw_result")
        if rw:
            rc1, rc2 = st.columns(2)
            with rc1:
                st.markdown('<div style="font-family:var(--mono);font-size:0.62rem;letter-spacing:0.1em;text-transform:uppercase;color:var(--t3);margin-bottom:6px">Original</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="sf-diff">{rtxt[:1400]}</div>', unsafe_allow_html=True)
            with rc2:
                st.markdown('<div style="font-family:var(--mono);font-size:0.62rem;letter-spacing:0.1em;text-transform:uppercase;color:var(--teal);margin-bottom:6px">Rewritten ✓</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="sf-diff">{rw[:1400]}</div>', unsafe_allow_html=True)
            # FIX Bug 2: All download buttons have unique key=
            st.download_button("⬇ Download rewritten resume", data=rw,
                               file_name="skillforge_rewritten.txt", mime="text/plain",
                               key="dl_rewrite")

    st.markdown('<div class="sf-divider"></div>', unsafe_allow_html=True)
    st.markdown('<div class="sf-sh">Export</div>', unsafe_allow_html=True)
    st.markdown('<div class="sf-ss">Download your personalized roadmap as PDF · JSON · CSV</div>', unsafe_allow_html=True)

    ex1, ex2, ex3 = st.columns(3, gap="medium")
    with ex1:
        st.markdown('<div class="sf-export-card">', unsafe_allow_html=True)
        st.markdown('<div class="sf-export-hd">PDF report</div><div class="sf-export-sub">Full roadmap · AI reasoning · ATS audit · scores</div>', unsafe_allow_html=True)
        for k,v in [("Candidate",c.get("name","–")),("Role",jd.get("role_title","–")),
                    ("ATS score",f"{ql.get('ats_score','–')}%"),("Modules",im.get("modules_count",0)),
                    ("Training",f"{im['roadmap_hours']}h")]:
            st.markdown(f'<div class="sf-export-row"><span class="sf-ek">{k}</span><span class="sf-ev">{v}</span></div>', unsafe_allow_html=True)
        st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
        if REPORTLAB:
            pdf_buf = build_pdf(c, jd, gp, roadmap, im, ql, iv)
            nm = (c.get("name","candidate") or "candidate").replace(" ","_")
            # FIX Bug 2: key="dl_pdf"
            st.download_button("⬇ Download PDF", data=pdf_buf,
                               file_name=f"skillforge_{nm}_{datetime.now().strftime('%Y%m%d')}.pdf",
                               mime="application/pdf", use_container_width=True,
                               key="dl_pdf")
        else:
            st.caption("`pip install reportlab` for PDF export")
        st.markdown('</div>', unsafe_allow_html=True)

    with ex2:
        st.markdown('<div class="sf-export-card">', unsafe_allow_html=True)
        st.markdown('<div class="sf-export-hd">JSON export</div><div class="sf-export-sub">Complete structured result for integrations and downstream tools</div>', unsafe_allow_html=True)
        export_data = {
            "candidate": c, "jd": jd, "impact": im, "interview": iv,
            "gap_profile": [{k2:v2 for k2,v2 in g.items() if k2!="catalog_course"} for g in gp],
            "roadmap": [{"id":m["id"],"title":m["title"],"skill":m["skill"],"level":m["level"],
                         "duration_hrs":m["duration_hrs"],"is_critical":m.get("is_critical",False),
                         "reasoning":m.get("reasoning","")} for m in roadmap],
            "generated_at": datetime.now().isoformat(),
        }
        st.markdown("<div style='height:60px'></div>", unsafe_allow_html=True)
        # FIX Bug 2: key="dl_json"
        st.download_button("⬇ Download JSON",
                           data=json.dumps(export_data, indent=2, default=str),
                           file_name=f"skillforge_{datetime.now().strftime('%Y%m%d')}.json",
                           mime="application/json", use_container_width=True,
                           key="dl_json")
        st.markdown('</div>', unsafe_allow_html=True)

    with ex3:
        st.markdown('<div class="sf-export-card">', unsafe_allow_html=True)
        st.markdown('<div class="sf-export-hd">Skill gap CSV</div><div class="sf-export-sub">Tabular gap data for spreadsheet analysis and HR reporting</div>', unsafe_allow_html=True)
        rows = ["Skill,Status,Proficiency,Required,Demand,Decayed"]
        for g in gp:
            rows.append(f'"{g["skill"]}",{g["status"]},{g["proficiency"]},{g["is_required"]},{g.get("demand",1)},{g.get("decayed",False)}')
        st.markdown("<div style='height:60px'></div>", unsafe_allow_html=True)
        # FIX Bug 2: key="dl_csv"
        st.download_button("⬇ Download CSV", data="\n".join(rows),
                           file_name=f"skillforge_gap_{datetime.now().strftime('%Y%m%d')}.csv",
                           mime="text/csv", use_container_width=True,
                           key="dl_csv")
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    total_cost = sum(e.get("cost",0) for e in _audit_log)
    with st.expander(f"API log — {len(_audit_log)} calls · ${total_cost:.5f}"):
        for e in reversed(_audit_log[-20:]):
            ok = e.get("status")=="ok"
            sc = "#4ade80" if ok else "#ef4444"
            st.markdown(f'<div class="sf-log"><span style="color:{sc}">{"●" if ok else "✕"}</span><span>{e.get("ts","")}</span><span style="color:var(--teal)">{e.get("model","")}</span><span>in:{e.get("in",0)} out:{e.get("out",0)}</span><span>{e.get("ms",0)}ms</span><span>${e.get("cost",0):.6f}</span></div>', unsafe_allow_html=True)

# =============================================================================
#  SIDEBAR
# =============================================================================
def render_sidebar(res):
    im = res["impact"]; iv = res["interview"]; sm = res.get("seniority",{})
    with st.sidebar:
        st.markdown('<div style="padding:14px 10px 6px"><div style="font-family:\'DM Sans\',sans-serif;font-size:1rem;font-weight:700;color:#f1f5f9;letter-spacing:-0.02em">Skill<span style="color:#2dd4bf">Forge</span></div></div>', unsafe_allow_html=True)
        sections = [("overview","#2dd4bf","Overview"),("roadmap","#f59e0b","Roadmap"),
                    ("research","#4ade80","Research"),("ats","#a78bfa","ATS & Export")]
        for anc, color, lbl in sections:
            st.markdown(f'<a href="#{anc}" class="sf-nav-item"><div class="sf-nav-dot" style="background:{color}"></div>{lbl}</a>', unsafe_allow_html=True)
        st.markdown('<div style="height:1px;background:rgba(255,255,255,0.07);margin:10px 0"></div>', unsafe_allow_html=True)
        st.markdown(f"""<div style="padding:4px 12px;font-family:'DM Mono',monospace;font-size:0.68rem;color:#3d4d66">
          <div>fit &nbsp;<span style="color:#94a3b8">+{im.get('fit_delta',0)}%</span></div>
          <div>modules &nbsp;<span style="color:#94a3b8">{im.get('modules_count',0)}</span></div>
          <div>hours &nbsp;<span style="color:#94a3b8">{im.get('roadmap_hours',0)}h</span></div>
          <div>interview &nbsp;<span style="color:{iv.get('color','#4ade80')}">{iv.get('score',0)}% {iv.get('label','')}</span></div>
        </div>""", unsafe_allow_html=True)
        st.markdown('<div style="height:1px;background:rgba(255,255,255,0.07);margin:10px 0"></div>', unsafe_allow_html=True)
        if sm.get("has_mismatch"):
            st.markdown(f'<div style="margin:0 8px 10px;background:rgba(245,158,11,0.07);border:1px solid rgba(245,158,11,0.2);border-radius:6px;padding:9px 12px;font-size:0.72rem;color:#f59e0b">⚠ {sm["candidate"]} → {sm["required"]}</div>', unsafe_allow_html=True)
        st.markdown('<div class="sf-ghost" style="padding:0 8px">', unsafe_allow_html=True)
        if st.button("↩ Start over", key="sb_reset", use_container_width=True):
            _full_reset()
        st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
#  RESULTS PAGE
# =============================================================================
def render_results():
    res = st.session_state.get("result")
    if not res:
        st.session_state["step"] = "input"
        st.rerun()
        return

    render_sidebar(res)

    st.markdown('<div class="sf-page">', unsafe_allow_html=True)

    _, rc = st.columns([12,1])
    with rc:
        st.markdown('<div class="sf-ghost">', unsafe_allow_html=True)
        if st.button("Reset", key="top_rst"):
            _full_reset()
        st.markdown('</div>', unsafe_allow_html=True)

    render_banner(res)

    st.markdown('<div id="overview"></div><div id="roadmap"></div><div id="research"></div><div id="ats"></div>', unsafe_allow_html=True)

    t1, t2, t3, t4 = st.tabs(["🗺️ Roadmap", "📊 Gap Analysis", "✅ ATS & Export", "🌐 Research"])
    with t1: render_tab_roadmap(res)
    with t2: render_tab_overview(res)
    with t3: render_tab_ats_export(res)
    with t4: render_tab_research(res)

    st.markdown("</div>", unsafe_allow_html=True)

# =============================================================================
#  FOOTER
# =============================================================================
def render_footer():
    st.markdown("""
    <div class="sf-foot">
      <span><span class="sf-fdot" style="background:#2dd4bf"></span>Powered by Groq LLaMA 3.3-70b</span>
      <span><span class="sf-fdot" style="background:#2dd4bf"></span>Llama 4 Scout Vision</span>
      <span><span class="sf-fdot" style="background:#2dd4bf"></span>NetworkX DAG</span>
      <span><span class="sf-fdot" style="background:#2dd4bf"></span>sentence-transformers</span>
      <span><span class="sf-fdot" style="background:#2dd4bf"></span>DuckDuckGo Search</span>
      <span class="sf-fr">ARTPARK CodeForge Hackathon 2025 · SkillForge</span>
    </div>""", unsafe_allow_html=True)

# =============================================================================

# =============================================================================
#  MAIN
# =============================================================================
def main():
    try:
        import streamlit.components.v1 as components
        components.html("""<script>
(function(){
  var S=['[data-testid="stToolbar"]','[data-testid="stDecoration"]','[data-testid="stStatusWidget"]',
         '[data-testid="stAppDeployButton"]','[data-testid="stBaseButton-header"]',
         '[data-testid="stBaseButton-minimal"]','button[kind="header"]','button[kind="minimal"]',
         '#stDecoration','#MainMenu'];
  function k(){S.forEach(function(s){document.querySelectorAll(s).forEach(function(e){e.remove();});});
    var h=document.querySelector('header[data-testid="stHeader"]');
    if(h){h.style.cssText='display:none!important;height:0!important;';}}
  k();new MutationObserver(k).observe(document.documentElement,{childList:true,subtree:true});
  [500,1500,3000].forEach(function(t){setTimeout(k,t);});
})();
</script>""", height=0, scrolling=False)
    except: pass
    st.markdown(CSS, unsafe_allow_html=True)
    _init_state()
    step = st.session_state.get("step","input")
    render_topbar()
    if   step == "input":     render_input()
    elif step == "analyzing": render_loading()
    elif step == "results":   render_results()
    render_footer()

#  CLI MODE
# =============================================================================
def cli_analyze(scenario_key):
    if scenario_key not in SAMPLES:
        print(f"Unknown: {list(SAMPLES.keys())}"); sys.exit(1)
    s = SAMPLES[scenario_key]; t0 = time.time()
    print(f"\n  SkillForge v9 CLI · {s['label']}\n  {'='*50}")
    result = run_analysis(s["resume"], s["jd"])
    print(f"  Done in {round(time.time()-t0,2)}s")
    if "error" in result: print(f"  Error: {result}"); return
    c=result["candidate"]; im=result["impact"]; iv=result["interview"]; path=result["path"]
    print(f"  Candidate : {c.get('name','–')} ({c.get('seniority','–')})")
    print(f"  Role      : {result['jd'].get('role_title','–')}")
    print(f"  Fit       : {im['current_fit']}% → {im['projected_fit']}% (+{im['fit_delta']}%)")
    print(f"  Interview : {iv['score']}% ({iv['label']})")
    print(f"  Roadmap   : {im['modules_count']} modules / {im['roadmap_hours']}h / {im['critical_count']} critical")
    for i,m in enumerate(path):
        print(f"    {'★' if m.get('is_critical') else ' '} #{i+1:02d} [{m['level'][:3]}] {m['title']} ({m['duration_hrs']}h)")
    print(f"\n  Hours saved vs generic 60h: ~{im['hours_saved']}h\n")

# =============================================================================
#  ENTRY POINT
# =============================================================================
if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--analyze", metavar="SCENARIO")
    args, _ = parser.parse_known_args()
    if args.analyze: cli_analyze(args.analyze)
    else:
        threading.Thread(target=_load_semantic_bg, daemon=True).start()
        main()